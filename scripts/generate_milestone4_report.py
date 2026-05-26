"""
Generate the fourth bi-weekly report (May 18 – May 31, 2026) for the HeartCompass project.
WBS/RBS-anchored with visual progress bars, effort charts, comprehensive detail.
Outputs: DOCX + PDF in output/report/ — exactly 5 pages, each completely filled.

Format matches Third Biweekly Report exactly. Content calibrated to ~20,000 chars for 5-page fit.
"""
from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "report"
DOCX_PATH = OUTPUT_DIR / "Fourth Biweekly Report-HeartCompass-BuTian.docx"
PDF_PATH = OUTPUT_DIR / "Fourth Biweekly Report-HeartCompass-BuTian.pdf"

PALETTE = {
    "navy": "1F4E79", "slate": "456B84", "teal": "5D8AA8",
    "pale": "DDE9F2", "soft": "F5F8FB", "text": "23313D",
    "muted": "6B7C8C",
    "green": "2E7D32", "blue": "1565C0", "amber": "E65100", "red": "B71C1C",
    "light_green": "E8F5E9", "light_blue": "E3F2FD", "light_amber": "FFF3E0",
    "white": "FFFFFF",
}


# ===================================================================
# DOCX helpers (identical to third report)
# ===================================================================
def _shd(cell, fill: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), fill)
    tc.append(s)


def _run(p, text, bold=False, size=None, color=None, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return r


def _set_cell(cell, text, bold=False, size=8, color=None, align="left", bg=None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if color:
        r.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    if bg:
        _shd(cell, bg)


def _section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    _run(p, text, bold=True, size=11, color=PALETTE["navy"])


def _subhead(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    _run(p, text, bold=True, size=9, color=PALETTE["text"])


def _body(doc, text, size=7.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    _run(p, text, size=size, color=PALETTE["text"])


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C0C8D0")
        borders.append(el)
    tblPr.append(borders)


def _pct_bg(pct_str: str) -> str:
    try: v = int(pct_str.replace("%", ""))
    except ValueError: return "F5F5F5"
    if v >= 100: return "C8E6C9"
    if v >= 75:  return "BBDEFB"
    if v >= 40:  return "FFE0B2"
    return "F5F5F5"


def _pct_fg(pct_str: str) -> str:
    try: v = int(pct_str.replace("%", ""))
    except ValueError: return "616161"
    if v >= 100: return "1B5E20"
    if v >= 75:  return "0D47A1"
    if v >= 40:  return "BF360C"
    return "616161"


def _status_bg(status: str) -> str:
    if status in ("Complete", "Achieved"): return "C8E6C9"
    if "Progress" in status: return "BBDEFB"
    if status in ("Ongoing", "On Track"): return "E3F2FD"
    if status == "Started": return "FFF3E0"
    return "F5F5F5"


def _progress_cell(cell, pct_str: str):
    _shd(cell, _pct_bg(pct_str))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(pct_str)
    r.bold = True
    r.font.size = Pt(7)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.color.rgb = RGBColor.from_string(_pct_fg(pct_str))


def _status_cell(cell, status: str):
    _shd(cell, _status_bg(status))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(status)
    r.bold = True
    r.font.size = Pt(6.5)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _make_wbs_table(doc, rows_data):
    table = doc.add_table(rows=len(rows_data) + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    headers = ["WBS Task", "Description", "Owner", "Hours", "Compl.", "Status"]
    widths = [0.66, 1.78, 1.0, 0.42, 0.48, 0.72]
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, size=6.5, color="FFFFFF", align="center", bg=PALETTE["navy"])
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            align = "center" if c_idx in (0, 3, 4, 5) else "left"
            if c_idx == 4:
                _progress_cell(table.rows[r_idx + 1].cells[c_idx], val)
            elif c_idx == 5:
                _status_cell(table.rows[r_idx + 1].cells[c_idx], val)
            else:
                _set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=7, align=align)
                if r_idx % 2 == 1:
                    _shd(table.rows[r_idx + 1].cells[c_idx], "FAFAFA")
    _set_table_borders(table)
    return table


def _make_person_task_table(doc, rows_data):
    table = doc.add_table(rows=len(rows_data) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    headers = ["WBS Task", "Description", "Hours", "Compl.", "Status"]
    widths = [0.7, 1.85, 0.48, 0.48, 0.65]
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, size=7.5, color="FFFFFF", align="center", bg=PALETTE["navy"])
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            align = "center" if c_idx in (0, 2, 3, 4) else "left"
            if c_idx == 3:
                _progress_cell(table.rows[r_idx + 1].cells[c_idx], val)
            elif c_idx == 4:
                _status_cell(table.rows[r_idx + 1].cells[c_idx], val)
            else:
                _set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=7.5, align=align)
                if r_idx % 2 == 1:
                    _shd(table.rows[r_idx + 1].cells[c_idx], "FAFAFA")
    _set_table_borders(table)
    return table


def _bullet_item(doc, wbs_tag, title, body, hours, pct):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.1)
    tag = f"▸ [{wbs_tag} · {hours} · {pct}] "
    r = p.add_run(tag)
    r.bold = True
    r.font.size = Pt(7.5)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.color.rgb = RGBColor.from_string(PALETTE["teal"].replace("#", ""))
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(7.5)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r2.font.color.rgb = RGBColor.from_string(PALETTE["text"].replace("#", ""))
    r3 = p.add_run(f"  {body}")
    r3.font.size = Pt(7.5)
    r3.font.name = "Times New Roman"
    r3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r3.font.color.rgb = RGBColor.from_string(PALETTE["text"].replace("#", ""))


def _next_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.12)
    _run(p, f"→ {text}", size=7.5, color=PALETTE["muted"])


def _challenge_item(doc, issue, solution):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.1)
    _run(p, "Issue: ", bold=True, size=7.5, color=PALETTE["red"])
    _run(p, issue, size=7.5, color=PALETTE["text"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.left_indent = Inches(0.2)
    _run(p2, f"→ {solution}", size=7.5, color=PALETTE["slate"])


def _effort_bar(doc, label, hours, max_hours, color_hex):
    pct = min(hours / max_hours, 1.0) if max_hours > 0 else 0
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    _set_cell(table.rows[0].cells[0], label, bold=True, size=7.5, align="right")
    table.rows[0].cells[0].width = Inches(1.1)
    _shd(table.rows[0].cells[1], "F0F0F0")
    table.rows[0].cells[1].width = Inches(3.0)
    p = table.rows[0].cells[1].paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    filled = int(pct * 20)
    bar = "█" * filled + "░" * (20 - filled) + f"  {hours}h"
    r = p.add_run(bar)
    r.font.size = Pt(7)
    r.font.name = "Courier New"
    r.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
    _set_cell(table.rows[0].cells[2], f"{hours}h", size=7.5, align="center")
    table.rows[0].cells[2].width = Inches(0.5)
    return table


# ===================================================================
# REPORT DATA — May 18–31, 2026 (Complete)
# ===================================================================
REPORT_TITLE = "Fourth Biweekly Report  (May 18 – May 31, 2026)"
PERIOD = "May 18 – May 31, 2026 (Weeks 10–11)  |  Submission: May 31, 2026"

WBS_SUMMARY_ROWS = [
    ["WP2.1-2.5", "Persona Building Pipeline (Stable)",             "Leishiyu/Haojie", "—",  "100%", "Complete"],
    ["WP3.1-3.2", "ConversationGraph Core & Memory Trimming",        "Yihan Liu",       "38", "100%", "Complete"],
    ["WP3.3",     "LLM Call with Retry Strategy",                    "Yihan Liu",       "14", "100%", "Complete"],
    ["WP3.4",     "Feed Sync — Conversation-to-Persona Loop",      "Yihan Liu",       "12", "100%", "Complete"],
    ["WP3.5-3.6", "Coherence Testing & Prompt Optimization",         "Yihan Liu",       "32", "90%",  "In Progress"],
    ["WP4.1",     "Lark WebSocket Connection & Message Routing",     "Leishiyu Chen",   "8",  "100%", "Complete"],
    ["WP4.2",     "Bot Command Menu & Persona Interaction",          "Tian/Leishiyu",   "22", "100%", "Complete"],
    ["WP4.3",     "Message Batching & Buffer Management",            "Tian/Leishiyu",   "14", "100%", "Complete"],
    ["WP4.5-4.6", "WebSocket Testing & Card Template System",        "Leishiyu Chen",   "22", "95%",  "In Progress"],
    ["WP5.1",     "CLI Framework: Doctor, Setup, Logs Commands",     "Tian Bu",         "16", "100%", "Complete"],
    ["WP5.2",     "CLI Auth Login & Session Persistence",            "Tian/Haojie",     "16", "100%", "Complete"],
    ["WP5.3-5.4", "Docker, Alembic, Deployment & Integration Tests", "Haojie Hu",       "26", "84%",  "In Progress"],
]

SUMMARY_INTRO = (
    "This report covers Weeks 10–11 (May 18–31, 2026), during which HeartCompass transitioned "
    "from core development into integration, validation, and delivery readiness. The team "
    "completed all remaining WP3 tasks (WP3.3–WP3.4 at 100%), finalized WP4 Channel Integration "
    "(WP4.1–WP4.3 at 100%, WP4.5-4.6 at 95%), and brought WP5 CLI/Deployment to full operational "
    "capability (WP5.1–WP5.2 at 100%, WP5.3 at 95%). WP3.5 coherence testing reached 100% pass "
    "rate; WP3.6 prompt optimization achieved 80% with measurable quality gains. Total: 220 "
    "hours across 12 tracked WBS tasks, weighted completion 94.8%. All milestones on track."
)

SUMMARY_METRICS = (
    "Overall: WP2 Persona Building maintained stable operation across all 4 dimensions with zero "
    "regressions. WP3 Conversation System achieved full feature completeness: LLM calls with "
    "exponential-backoff retry (WP3.3, 100%), 3-tier confidence-scored feed sync closing the "
    "conversation-to-persona loop (WP3.4, 100%), comprehensive 20-turn coherence test suite at "
    "100% pass rate across all 5 FigureRole types (WP3.5, 100%). Prompt optimization hit 80% "
    "with 35% repetition reduction and 28% naturalness gain. WP4 Channel Integration is production-"
    "ready: WebSocket with heartbeat, graceful shutdown, and live-sandbox-validated reconnection; "
    "complete 4-command bot menu with rate limiting; message batching with 50-msg cap and cross-"
    "device dedup; card templates with null-safe rendering. WP5 CLI at 100% (doctor/setup/logs, "
    "auth login with JWT persistence + transparent refresh). Docker Compose with healthcheck-gated "
    "orchestration and Alembic migration 003 validated staging-ready (95%). WP5.4 integration test "
    "plan approved for Week 12. Risk: LOW. Fully on schedule."
)

RISK_ITEMS = [
    ("Buffer overflow in message batching (resolved, May 22)",
     "Per-user buffer grew unbounded at 50+ msg/15s. Fixed with 50-message hard cap + auto force-flush."),
    ("LLM API rate limiting under concurrency (resolved)",
     "3+ concurrent ConversationGraph instances triggered 429 errors. Mitigated by WP3.3 "
     "exponential-backoff retry (max 3 attempts, 1s/2s/4s delays). Request queue planned."),
    ("Card template empty-field display regression (resolved, May 23)",
     "WP4.2 fix caused 'None' literal rendering for empty fields. Corrected via null-safe "
     "accessors with per-dimension DEFAULT_PLACEHOLDERS. ~0.5 day rework."),
]

WEEKLY_MILESTONES = [
    ["Week 10\n(May 18–24)", "WP3.3–WP3.5 + WP4.2–WP4.5 + WP5.1–WP5.3",
     "LLM retry complete (100%), feed sync 100%, persona commands interactive with null-safe "
     "cards, buffer overflow protected, WebSocket reconnection tested, hybrid search integrated, "
     "Docker Compose with healthcheck gating, CLI auth login 95%, 20-turn coherence scenarios.",
     "Achieved"],
    ["Week 11\n(May 25–31)", "WP3.5–WP3.6 + WP4.5–WP4.6 + WP5.2–WP5.4",
     "Coherence suite 100% pass rate after topic-segmented summary fix, prompt optimization 80% "
     "(35% repetition reduction, 28% naturalness gain), WebSocket live sandbox validated, card "
     "templates at 90%, CLI auth 100%, Alembic migration validated, test plan approved.",
     "On Track"],
]

# ===================================================================
# Individual Data — calibrated for ~20,000 total chars
# ===================================================================
MEMBER_DATA = [
    # ===== Leishiyu Chen (Page 2) =====
    {
        "name": "Leishiyu Chen",
        "role": "Data Layer & Channel Integration (RBS R1.2.1, R2.2.1)",
        "hours": 48,
        "task_table": [
            ["WP4.1", "WebSocket Connection & Message Routing",       "8",  "100%", "Complete"],
            ["WP4.2", "Bot Command Menu — Persona Interaction",    "12", "100%", "Complete"],
            ["WP4.3", "Message Batching & Buffer Management",         "6",  "100%", "Complete"],
            ["WP4.5", "WebSocket Reconnection & Error Recovery Tests","14", "100%", "Complete"],
            ["WP4.6", "Card Template System & Display Refinement",    "8",  "90%",  "In Progress"],
        ],
        "summary": (
            "My primary focus in Weeks 10–11 was completing the WP4 Channel Integration pipeline "
            "end-to-end. WP4.1 WebSocket connectivity reached 100% with open_id-based user-to-FR "
            "mapping, 30s heartbeat keepalive, and graceful shutdown. WP4.2 delivered the complete "
            "bot command suite at 100%: /menu, /list_available_persons, /show_persona with null-safe "
            "accessors, and /build_persona with async pipeline triggering. WP4.3 message batching "
            "was hardened with 50-message cap, force-flush, and cross-device content-hash dedup "
            "(100%). WP4.5 WebSocket reconnection testing reached 100% after live Lark sandbox "
            "validation. WP4.6 card template refinement at 90%: null-safe rendering, text truncation, "
            "error-state templates, and 4-stage build-progress indicators. Total: 48h."
        ),
        "progress": [
            ("WP4.2", "Bot Command Menu — Complete Persona Interaction Suite",
             "Completed all 4 Lark Bot commands at 100%. /menu displays available commands with "
             "usage hints. /list_available_persons queries user FR records with ID, name, type, "
             "and last-updated timestamp. /show_persona renders structured persona cards across "
             "4 core dimensions with null-safe accessors resolving the 'None' display regression "
             "via per-dimension Chinese DEFAULT_PLACEHOLDERS. /build_persona triggers the async "
             "FRBuildingGraph pipeline with progress-tracking card. All commands include input "
             "validation, Chinese error messages, and rate limiting (5 req/30s per user).",
             "12h", "100%"),
            ("WP4.3", "Message Batching — Overflow Protection & Hardening",
             "Hardened batching with production safeguards: 50-message hard cap per user buffer "
             "(exceeded → force-flush + WARNING log), suppressed duplicate timer firings during "
             "in-progress flush, cross-device deduplication via SHA256(open_id + text + timestamp/"
             "15s bucket) with TTL-limited hash set. Correctly handles rapid-fire bursts, single-"
             "message turns, timer/send races, and multi-device overlap across desktop and mobile.",
             "6h", "100%"),
            ("WP4.5-4.6", "WebSocket Testing & Card Template Refinement",
             "Built comprehensive WebSocket test suite: connect/disconnect cycle validation, "
             "exponential-backoff reconnection (1s/2s/4s, max 3 attempts), message queuing with "
             "replay on reconnect, invalid auth rejection, and 30s idle timeout proactive "
             "reconnect. Week 11 milestone: completed live Lark sandbox integration validation "
             "(WP4.5 at 100%). Card templates refined with null-safe accessors using contextual "
             "Chinese placeholders, 500-char truncation with 'Show More' expansion, 4-stage "
             "build-progress indicator, and 3 error-state templates. Dark-mode palette pending.",
             "22h", "95%"),
        ],
        "challenges": [
            ("Empty-field card regression & cross-device message duplication (both resolved)",
             "WP4.2 fix caused 'None' rendering for empty MEMORY fields → added null-coalescing "
             "with per-dimension DEFAULT_PLACEHOLDERS. Users switching desktop/mobile within 15s "
             "batching window produced near-duplicate messages → content-hash dedup with 15s-"
             "bucketed timestamps. Verified across 50+ cross-device test sessions."),
        ],
        "deliverables": [
            "Production-ready WebSocket layer with heartbeat and graceful shutdown (100%)",
            "Complete bot command menu: /menu, /list, /show_persona, /build_persona (100%)",
            "Hardened batching: 50-msg cap, force-flush, timer race protection, cross-device dedup (100%)",
            "Card template system: null-safe accessors, 4-stage progress, error-state templates (90%)",
        ],
        "next": [
            "WP4.6: Complete dark-mode palette; add /delete_persona and /update_persona commands",
            "Support WP5.4 channel-layer regression testing with automated WebSocket message replay",
            "Prepare WP4 handoff: WebSocket lifecycle, bot command API, card template schema",
        ],
    },

    # ===== Yihan Liu (Page 3) =====
    {
        "name": "Yihan Liu",
        "role": "Conversation & Dialogue System (RBS R2.1)",
        "hours": 52,
        "task_table": [
            ["WP3.3", "LLM Call with Retry Strategy",                  "6",  "100%", "Complete"],
            ["WP3.4", "Feed Sync — Conversation-to-Persona Loop",    "8",  "100%", "Complete"],
            ["WP3.5", "Multi-Turn Coherence Testing & Optimization",    "20", "100%", "Complete"],
            ["WP3.6", "Prompt Template Refinement & Tone Tuning",       "12", "80%",  "In Progress"],
            ["WP3.2f","Memory Trim Boundary Condition Fix",            "6",  "100%", "Complete"],
        ],
        "summary": (
            "My work in Weeks 10–11 completed the WP3 Conversation System to full feature readiness "
            "and advanced into systematic testing and quality optimization. WP3.3 LLM call reached "
            "100% with production-grade exponential-backoff retry and structured output validation. "
            "WP3.4 feed sync achieved 100% with 3-tier confidence-scored update logic — closing "
            "the critical conversation-to-persona feedback loop. WP3.5 delivered a comprehensive "
            "20-turn coherence test suite covering all 5 FigureRole types, reaching 100% pass rate "
            "(20/20 scenarios) after implementing topic-segmented conversation summaries. WP3.6 "
            "prompt tone optimization reached 80% with 35% repetition reduction and 28% naturalness "
            "gain. Total: 52h."
        ),
        "progress": [
            ("WP3.4", "Feed Sync — Conversation-to-Persona Closed Loop",
             "Completed feed sync node closing the critical feedback loop. 3-tier confidence "
             "threshold: (1) ≥0.8 → direct FR core field update with full audit trail; (2) "
             "0.5–0.79 → upsert as FineGrainedFeed with PENDING_REVIEW status; (3) <0.5 → "
             "observation log for manual curation. Integrated with WP2.3 conflict detection: "
             "contradictory high-confidence insights automatically create CONFLICT_PENDING "
             "entries preserving both values. Sync latency <2s per turn. Validated across "
             "12 synthetic scenarios covering all 5 FigureRole types.",
             "8h", "100%"),
            ("WP3.5", "Multi-Turn Coherence Testing — 100% Pass Rate Achieved",
             "Built comprehensive test suite: parametrized 5/10/20-turn scenarios across all "
             "5 FigureRole types (family/friend/mentor/colleague/partner). Validates response "
             "relevance, persona consistency, context retention, and memory trim correctness. "
             "Edge cases: rapid topic switching, 500+ word monologue, Chinese/English code-"
             "switching. Reached 100% pass rate (20/20 scenarios) in Week 11 after implementing "
             "topic-segmented conversation summaries that prevent emotional-framing carryover "
             "between unrelated topics — resolving the partner-role context bleed bug.",
             "20h", "100%"),
            ("WP3.6", "Prompt Template Refinement — Tone & Anti-Repetition",
             "Analyzed 250+ LLM responses identifying 3 core issues: over-formal language in "
             "friend/partner roles, repetitive phrases (3–4x per 10 turns), and underuse of "
             "persona-specific speech patterns. Phase-1 delivered: role-specific formality "
             "parameters, anti-repetition instructions with phrase-history tracker, and 3–5 "
             "concrete speech examples per persona. Quantitative results: 35% reduction in "
             "repetition frequency (3.2→2.1 avg. repeats/10 turns), 28% improvement in "
             "perceived naturalness (blind A/B test, n=40). Phase 2 (emotion-aware tone "
             "modulation, cultural context adaptation) scoped for Week 12.",
             "12h", "80%"),
        ],
        "challenges": [
            ("Rapid topic switching causing emotional-framing context bleed (resolved, Week 11)",
             "Partner-role 20-turn test with 5 rapid topic switches showed emotional carryover "
             "from 'work stress' into 'weekend plans'. Root cause: rolling summary condensed all "
             "topics into one undifferentiated paragraph. → Implemented topic-segmented summary "
             "with embedding-based semantic shift detection — each topic gets labeled paragraph, "
             "prompt attends primarily to most recent segment. 100% pass rate achieved."),
        ],
        "deliverables": [
            "LLM call node with exponential-backoff retry (max 3) and structured output validation (100%)",
            "Feed sync node: 3-tier confidence threshold integrated with WP2.3 conflict system (100%)",
            "20-turn coherence test suite: 20 scenarios, 100% pass rate, all 5 FigureRole types (100%)",
            "Phase-1 prompt optimization: anti-repetition (35% reduction), formality tuning (80%)",
        ],
        "next": [
            "WP3.6: Phase-2 — emotion-aware tone modulation, cultural context adaptation (Week 12)",
            "Add 30/50-turn extended scenarios for long-session stability testing",
            "Prepare WP3 handoff: architecture doc, prompt engineering guide, test coverage report",
        ],
    },

    # ===== Haojie Hu (Page 4) =====
    {
        "name": "Haojie Hu",
        "role": "Knowledge Base, Vector Recall & Deployment (RBS R1.2.2, R2.1.1, R3.1.2)",
        "hours": 48,
        "task_table": [
            ["WP2.4e","Hybrid Search — Semantic + Keyword Recall",       "10", "100%", "Complete"],
            ["WP3.1o","Recall Dimension Weight Tuning for Conversation",   "6",  "100%", "Complete"],
            ["WP5.2", "CLI Auth Login & Session Persistence",             "10", "100%", "Complete"],
            ["WP5.3", "Docker Compose, Alembic & Deployment Config",      "16", "95%",  "In Progress"],
            ["WP5.4", "System Integration Testing — Test Plan Design",   "6",  "50%",  "In Progress"],
        ],
        "summary": (
            "My work in Weeks 10–11 spanned recall enhancements, deployment infrastructure, and "
            "system testing preparation. WP2.4 hybrid search combined semantic vector recall with "
            "BM25 keyword scoring via Reciprocal Rank Fusion — 18–22% recall improvement for named "
            "entities and mixed-language queries (10h, 100%). WP3.1 optimization tuned per-dimension "
            "recall weights (MEMORY 1.2, PROCEDURAL_INFO 1.1, PERSONALITY 0.9, INTERACTION_STYLE "
            "0.8), reducing LLM input tokens ~12% (6h, 100%). WP5.2 CLI auth login completed to "
            "100% with JWT persistence, transparent refresh, and token revocation. WP5.3 delivered "
            "Docker Compose with custom healthcheck-gated pgvector+app orchestration and Alembic "
            "migration 003 with verified rollback (16h, 95%). WP5.4 integration test plan with 8 "
            "critical-path scenarios approved for Week 12 (6h, 50%). Total: 48h."
        ),
        "progress": [
            ("WP2.4-enh", "Hybrid Search — Semantic + Keyword Recall Fusion",
             "Implemented hybrid search engine combining semantic vector recall with BM25 "
             "keyword scoring via Reciprocal Rank Fusion (k=60). Pipeline: query vectorization "
             "→ top-K×2 semantic recall via pgvector HNSW index → parallel BM25 keyword search "
             "(PostgreSQL to_tsvector/to_tsquery with language stemming) → RRF merge. "
             "Evaluation results: 18% recall@5 improvement for mixed Chinese-English queries, "
             "22% for named entities, 15% for rare terms. RRF merge latency <50ms for top-20 "
             "candidate pools. Validated on 200-query benchmark covering 5 semantic categories.",
             "10h", "100%"),
            ("WP5.2", "CLI Auth Login — JWT Flow & Session Management",
             "Completed CLI auth subsystem to 100%. Full implementation: session file storage "
             "(~/.immortality/session.json, 0600 permissions, atomic temp-file write with "
             "O_CREAT|O_EXCL), transparent token refresh (401 response → refresh → retry "
             "original request), `immortality whoami` command (user identity + token expiry "
             "display), logout (secure local deletion + server-side token revocation). Security "
             "guarantees: tokens never logged, file permissions enforced at OS level, refresh "
             "rotation verified across 3 consecutive cycles. Co-developed with Tian Bu.",
             "10h", "100%"),
            ("WP5.3-5.4", "Docker Deployment & Integration Test Planning",
             "Docker Compose infrastructure: pgvector/pgvector:pg16 with custom healthcheck "
             "(verifies both pg_isready AND vector type availability in pg_available_extensions "
             "→ resolves ~8s pgvector compilation loading race) + app container with multi-stage "
             "Dockerfile and depends_on service_healthy. Alembic migration 003 adds source_type "
             "and trigger_event columns to FROverallUpdateLog; auto-runs on container startup; "
             "downgrade path verified. Staging validation: 30 cold-start cycles, zero failures, "
             "~10s average startup time. WP5.4 test plan identifies 8 critical-path scenarios "
             "(onboarding→persona→conversation, multi-session evolution, concurrent graph "
             "execution, WS disconnect/reconnect, DB pool under 20 concurrent, LLM failure "
             "recovery, Docker restart state recovery, cross-platform delivery). Approved for "
             "Week 12 execution with detailed specifications.",
             "22h", "84%"),
        ],
        "challenges": [
            ("Docker pgvector loading delay causing healthcheck false negatives (resolved)",
             "pgvector compilation took ~8s during first run; default pg_isready passed before "
             "extension availability, causing CREATE EXTENSION failures. → Custom healthcheck "
             "verifying both pg_isready AND vector type in pg_available_extensions. Extended "
             "start_period to 15s. 30 cold-start cycles in staging — zero failures."),
        ],
        "deliverables": [
            "Hybrid search engine: semantic + BM25 keyword recall with RRF fusion (18-22% recall gain, 100%)",
            "Dimension weight tuning: MEMORY 1.2, PROCEDURAL 1.1, PERSONALITY 0.9, STYLE 0.8 (100%)",
            "CLI auth login: JWT persistence, transparent refresh, session status, token revocation (100%)",
            "Docker Compose: pgvector/pg16 + app, custom healthcheck, Alembic 003 with rollback (95%)",
        ],
        "next": [
            "WP5.3: Create docker-compose.prod.yml with resource limits and restart policies",
            "WP5.4: Execute critical-path scenarios 1-4 manually; automate 5-8 with test scripts",
            "Prepare WP5 handoff: deployment runbook, migration guide, healthcheck configuration",
        ],
    },

    # ===== Tian Bu (Page 5) =====
    {
        "name": "Tian Bu  (Team Leader)",
        "role": "CLI, Bot Commands & Project Governance (RBS R2.2.2, R3.1.1)",
        "hours": 44,
        "task_table": [
            ["WP5.1", "CLI Framework — Doctor, Setup, Logs Commands",   "10", "100%", "Complete"],
            ["WP5.2", "CLI Auth Login — Co-development",                "6",  "100%", "Complete"],
            ["WP4.2", "Bot Commands — Co-development & Integration",      "8",  "100%", "Complete"],
            ["WP4.3", "Message Batching — Co-development & Review",       "4",  "100%", "Complete"],
            ["WP5.3", "Docker & Deployment — Oversight & Validation",      "4",  "95%",  "In Progress"],
            ["—",   "Project Governance & Biweekly Report",            "12", "—",   "Ongoing"],
        ],
        "summary": (
            "My work in Weeks 10–11 spanned CLI completion, bot command co-development, deployment "
            "oversight, and project governance. WP5.1 CLI framework reached 100%: doctor validates "
            "Python version, env vars, dependencies, DB connectivity, and Lark credentials; setup "
            "provides interactive .env generation with masked password input; logs supports "
            "--level/--module/--tail filtering with --json output. WP5.2 auth login co-implemented "
            "to 100% — I designed the CLI UX (secure getpass, session file 0600, automatic "
            "Authorization header injection). WP4.2/4.3 bot commands and batching co-developed "
            "with Leishiyu Chen. WP5.3 deployment oversight: validated healthcheck gating, "
            "approved migration 003, verified staging. Governance (12h): maintained WBS tracking "
            "through Weeks 8–9, normalized logging, coordinated 14-merge Week 10 with zero "
            "conflicts, tagged v0.4.0-dev (May 24), compiled this report. Total: 44h."
        ),
        "progress": [
            ("WP5.1", "CLI Framework — Full Feature Completion (75% → 100%)",
             "Brought immortality CLI from 75% to 100% feature completion. Doctor validates "
             "Python 3.11–3.13, all required env vars (DATABASE_URL, LARK_APP_ID/SECRET, "
             "ARK_API_KEY/BASE_URL), dependency imports, DB connectivity, and Lark API "
             "credentials with clear pass/fail indicators. Setup provides interactive .env "
             "generation with masked getpass input, SQLAlchemy table creation, and pgvector "
             "extension verification. Logs supports --level/--module/--tail/--follow filters "
             "plus --json structured output for pipeline consumption. Consistent UX throughout: "
             "colored status indicators, progress spinners, explicit exit codes (0/1/2) for "
             "CI/CD scripting integration.",
             "10h", "100%"),
            ("WP4.2/4.3", "Bot Commands & Message Batching — Co-development to Production",
             "Co-developed the complete Lark Bot interaction layer with Leishiyu Chen. WP4.2 "
             "contributions: command routing dispatcher with fuzzy Chinese/English matching, "
             "/list_available_persons implementation, and error response templates for invalid "
             "or malformed commands. WP4.3 contributions: 15s configurable wait timer with "
             "per-user buffer isolation (keyed by open_id), timer-reset-on-new-message logic, "
             "and design review of the 50-message hard cap and cross-device deduplication "
             "mechanism. All 4 bot commands plus the batching pipeline verified end-to-end on "
             "both Lark desktop client and mobile app. Both subsystems at 100%, production-"
             "hardened with comprehensive error handling.",
             "12h", "100%"),
            ("Governance", "Project Governance, Release Management & Biweekly Report",
             "Maintained WBS milestone tracking across Weeks 8–9: Week 8 checkpoint (May 16) "
             "recorded WP2.1-2.5 and WP3.1-3.2 at 100%, WP3.3 at 80%, WP3.4 at 60%, WP4.1 "
             "at 100%, WP4.2 at 85%, WP4.3 at 70%, WP5.1 at 85%, WP5.2 at 50%, WP5.3 at "
             "40%. Week 9 checkpoint (May 23): WP3.3-3.4 reached 100%, WP4.2-4.3 hardened "
             "to 100%, WP5.2 advanced to 60%, WP5.3 to 50%. Normalized logging format across "
             "all ConversationGraph nodes for consistency. Coordinated Week 10 merge cadence: "
             "4 parallel feature branches, 14 --no-ff merges, morning-merge discipline (all "
             "merges before 11:00 AM CST) achieving zero unresolved conflicts reaching main. "
             "Tagged v0.4.0-dev on May 24 marking Weeks 8–9 deliverable completion. Compiled "
             "this fourth biweekly report: 12 WBS tasks, 220 total hours, 94.8% weighted "
             "completion across all work packages. Risk: LOW, all milestones on schedule.",
             "12h", "Ongoing"),
        ],
        "challenges": [
            ("Coordinating 4 parallel branches with 14 merges across 7 days (managed)",
             "Week 10 saw peak velocity: all 4 members committing simultaneously → elevated merge "
             "conflict risk. → Morning-merge cadence: branches merged before 11:00 AM CST, "
             "afternoon for resolution and integration testing. All merge commits carry team "
             "leader signature for auditable integration. Zero unresolved conflicts reached main."),
        ],
        "deliverables": [
            "immortality CLI v1.0: doctor (5 checks), setup (interactive .env + DB), logs (filtered + --json)",
            "CLI auth login: secure getpass, JWT persistence, auto-refresh, logout with revocation (100%)",
            "Lark Bot production suite: 4 commands with validation, Chinese errors, rate limiting (100%)",
            "v0.4.0-dev release (May 24); Fourth biweekly report: 12 tasks, 220h, 94.8% completion",
        ],
        "next": [
            "WP5.3: Finalize docker-compose.prod.yml; production deployment dry-run (Week 12)",
            "Coordinate WP5.4 test execution: assign scenarios, track pass/fail",
            "Prepare final presentation: architecture, WBS milestones, key metrics, demo outline",
        ],
    },
]


# ===================================================================
# Build DOCX
# ===================================================================
def build_docx() -> None:
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(0.35)
    sec.bottom_margin = Inches(0.25)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(hp, "HeartCompass  |  Fourth Biweekly Report  |  Confidential", size=6.5, color=PALETTE["muted"])

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, "Generated May 31, 2026  |  HeartCompass Project", size=6.5, color=PALETTE["muted"])

    # ================================================================
    # PAGE 1 — TEAM SUMMARY
    # ================================================================

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(1)
    _run(tp, REPORT_TITLE, bold=True, size=13, color=PALETTE["navy"])

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(4)
    _run(sp, f"Reporting Period: {PERIOD}", size=7.5, color=PALETTE["muted"])

    _section(doc, "1. Team Summary — WBS Progress Overview")
    _body(doc, SUMMARY_INTRO, size=7.5)

    _subhead(doc, "1.1  WBS Task Progress Matrix (Weeks 10–11)")
    _make_wbs_table(doc, WBS_SUMMARY_ROWS)

    _subhead(doc, "1.2  Overall Assessment & Outlook")
    _body(doc, SUMMARY_METRICS, size=7.5)

    _subhead(doc, "1.3  Effort Distribution by WBS Area")
    _body(doc,
        "WP2 (Persona + Recall Enh.): 16h (7.3%)  |  WP3 (Conversation): 96h (43.6%)  |  "
        "WP4 (Channel): 70h (31.8%)  |  WP5 (CLI + Deploy + Gov): 38h (17.3%)", size=7.5)
    max_h = 96
    _effort_bar(doc, "WP3 Conversation System", 96, 96, PALETTE["blue"])
    _effort_bar(doc, "WP4 Channel Integration", 70, 96, PALETTE["teal"])
    _effort_bar(doc, "WP5 CLI + Deploy + Gov", 38, 96, PALETTE["green"])
    _effort_bar(doc, "WP2 Persona + Recall", 16, 96, PALETTE["muted"])

    _subhead(doc, "1.4  Weekly Milestone Achievement")
    mt = doc.add_table(rows=len(WEEKLY_MILESTONES) + 1, cols=4)
    mt.alignment = WD_TABLE_ALIGNMENT.LEFT
    mt.autofit = True
    mh = ["Week", "Key Deliverables (WBS)", "Description", "Status"]
    for i, h in enumerate(mh):
        _set_cell(mt.rows[0].cells[i], h, bold=True, size=7, color="FFFFFF", align="center", bg=PALETTE["navy"])
    col_ws = [0.85, 1.5, 2.5, 0.7]
    for i, w in enumerate(col_ws):
        for row in mt.rows:
            row.cells[i].width = Inches(w)
    for r_idx, row_data in enumerate(WEEKLY_MILESTONES):
        for c_idx, val in enumerate(row_data):
            align = "center" if c_idx in (0, 3) else "left"
            if c_idx == 3:
                _set_cell(mt.rows[r_idx + 1].cells[c_idx], val, bold=True, size=7, align="center",
                          bg=_status_bg(val))
            else:
                _set_cell(mt.rows[r_idx + 1].cells[c_idx], val, size=7, align=align)
    _set_table_borders(mt)

    _subhead(doc, "1.5  Key Risks & Issues Log")
    for issue, detail in RISK_ITEMS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.02
        _run(p, f"▸ {issue}: ", bold=True, size=7.5, color=PALETTE["text"])
        _run(p, detail, size=7.5, color=PALETTE["text"])

    # ================================================================
    # PAGES 2-5 — INDIVIDUAL REPORTS
    # ================================================================
    for m in MEMBER_DATA:
        doc.add_page_break()

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _run(p, m["name"], bold=True, size=11, color=PALETTE["navy"])

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _run(p, f"RBS Area: {m['role']}    |    Total Hours: {m['hours']}h",
             bold=True, size=8, color=PALETTE["teal"])

        _subhead(doc, "I. WBS Task Breakdown")
        _make_person_task_table(doc, m["task_table"])

        _subhead(doc, "II. Period Summary")
        _body(doc, m["summary"], size=7.5)

        _subhead(doc, "III. Key Progress by WBS Task")
        for wbs_tag, title, body, hours, pct in m["progress"]:
            _bullet_item(doc, wbs_tag, title, body, hours, pct)

        _subhead(doc, "IV. Challenges & Solutions")
        for issue, solution in m["challenges"]:
            _challenge_item(doc, issue, solution)

        _subhead(doc, "V. Key Deliverables")
        for d in m["deliverables"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.12)
            _run(p, f"✓ {d}", size=7.5, color=PALETTE["text"])

        _subhead(doc, "VI. Next Steps")
        for s in m["next"]:
            _next_item(doc, s)

    doc.save(str(DOCX_PATH))
    print(f"Generated: {DOCX_PATH}")


def build_pdf() -> None:
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(OUTPUT_DIR), str(DOCX_PATH)],
        check=True, timeout=60,
    )
    print(f"Generated: {PDF_PATH}")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    try:
        build_pdf()
    except (FileNotFoundError, subprocess.CalledProcessError, Exception) as e:
        print(f"Note: PDF conversion skipped ({e}). DOCX ready at: {DOCX_PATH}")


if __name__ == "__main__":
    main()
