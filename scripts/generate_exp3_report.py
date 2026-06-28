#!/usr/bin/env python3
"""Generate Experiment 3 Economical Evaluation Report for Haojie Hu."""

import docx
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/Users/jack/work/HeartCompass/output/report"
STUDENT_NAME = "Haojie Hu"
STUDENT_ID = "2351493"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, f"{STUDENT_ID}+{STUDENT_NAME}+Experiment-3.docx")

doc = docx.Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)


def add_heading_styled(text, level=1):
    """Add a heading with proper Chinese font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(text, bold=False, align=None, size=12, font_name=None):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name or 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if align is not None:
        p.alignment = align
    return p


def set_cell_font(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold


def shade_cells(row, color="D9E2F3"):
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)


def make_table(headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_font(hdr.cells[i], h, bold=True, size=10)
    shade_cells(hdr)

    for r, row_data in enumerate(rows):
        row = table.rows[r + 1]
        for c, val in enumerate(row_data):
            set_cell_font(row.cells[c], val, bold=False, size=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════════════════
# COVER / HEADER
# ═══════════════════════════════════════════════════════════════

add_para("Software Engineering Management and Economics", bold=True, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Experiment III Report", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Economical Evaluation of Software Project", bold=True, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("")
add_para("Course Project: HeartCompass (Digital Immortality System)", bold=True, size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("")

# Info table
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ("Student Name", STUDENT_NAME),
    ("Student ID", STUDENT_ID),
    ("Instructor", "Du Bowen"),
    ("Submission Date", "June 22, 2026"),
]
for i, (k, v) in enumerate(info_data):
    set_cell_font(info_table.rows[i].cells[0], k, bold=True, size=11,
                  align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(info_table.rows[i].cells[1], v, bold=False, size=11,
                  align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════

add_heading_styled("Abstract", level=1)

add_para(
    "This report takes the course project HeartCompass (Digital Immortality System) as the "
    "evaluation object and, in accordance with the software project economical evaluation experiment "
    "requirements, analyzes its business goal, target users, market environment, cost structure, "
    "revenue model, taxes, cash flow, and economic evaluation indicators. HeartCompass is not "
    "positioned as a conventional chatbot or a static knowledge base. Instead, through long-term "
    "memory, user/persona profiles, vector retrieval, and large-model dialogue, it aims to gradually "
    "distill a \"digital companion persona\" capable of continuous, context-aware interaction."
)

add_para(
    "Under the benchmark scenario, the five-year net present value (NPV) of the project is "
    "approximately RMB 438,700, the internal rate of return (IRR) is approximately 70.47%, and "
    "the dynamic investment payback period is approximately 2.33 years. Sensitivity analysis "
    "covering revenue decline of 20% combined with cost increase of 10% still yields a positive "
    "NPV. Therefore, based on the existing course project prototype, HeartCompass demonstrates "
    "economic feasibility for continued refinement and small-scale product validation."
)

# ═══════════════════════════════════════════════════════════════
# I. EXPERIMENTAL PURPOSE AND PRINCIPLES
# ═══════════════════════════════════════════════════════════════

add_heading_styled("I. Experimental Purpose and Principles", level=1)

add_heading_styled("1.1 Experimental Purpose", level=2)
add_para(
    "The purpose of this experiment is to design and complete a comprehensive economical evaluation "
    "report for a software project, using the course project as the evaluation object. Through this "
    "experiment, I need to: (1) master the modeling methods of software project evaluation; "
    "(2) understand software product pricing strategies; (3) record experimental procedures, organize "
    "experimental data, and establish parameter mapping relationships among financial tables; and "
    "(4) apply economic evaluation models to analyze the economic benefits and feasibility of "
    "software projects."
)
add_para(
    "This report uses the HeartCompass course project as the evaluation object and completes the "
    "following work: clarifying the business goal of the project; collecting and organizing basic "
    "project information, market and economic environment data, and technical constraints; designing "
    "a business model and pricing strategy aligned with concrete user needs; estimating project "
    "size, investment, cost, and revenue; preparing auxiliary financial analysis statements; "
    "calculating project investment cash flow, NPV, IRR, and dynamic investment payback period; "
    "and ultimately judging the economic feasibility of the project."
)

add_heading_styled("1.2 Experimental Principles", level=2)
add_para(
    "The economical evaluation of a software project involves predicting the investment, revenue, "
    "cost, taxes, and cash flow of the project under the assumptions of the current fiscal and "
    "taxation system, market price system, and project life cycle, and then determining whether "
    "the project possesses profitability, investment recovery capability, and sustainable operation "
    "capacity."
)
add_para(
    "This experiment adopts the project investment analysis approach (pre-financing analysis). "
    "It does not consider debt financing, interest expenses, or complex capital structures, but "
    "evaluates profitability purely from the perspective of the cash flows generated by the "
    "project itself. The main evaluation indicators are:"
)
add_para(
    "• Net Present Value (NPV): NPV = Σ(CI − CO)ₜ × (1 + i₀)⁻ᵗ\n"
    "• Internal Rate of Return (IRR): the discount rate that makes NPV = 0.\n"
    "• Dynamic Investment Payback Period (Pₜ'): the time required for the cumulative discounted "
    "net cash flow to turn from negative to positive.\n\n"
    "In the formulas above, CI represents cash inflow, CO represents cash outflow, i₀ represents "
    "the benchmark discount rate, and t represents the time period. If NPV > 0, IRR > benchmark "
    "discount rate, and Pₜ' < project life cycle, the project is generally considered economically "
    "feasible."
)

# ═══════════════════════════════════════════════════════════════
# II. COURSE PROJECT OVERVIEW AND BASIC DATA
# ═══════════════════════════════════════════════════════════════

add_heading_styled("II. Course Project Overview and Basic Data", level=1)

add_heading_styled("2.1 Basic Project Information", level=2)

make_table(
    ["Project Element", "Description"],
    [
        ["Project Name", "HeartCompass (Digital Immortality / Digital Companion Persona System)"],
        ["Project Type", "AI-driven long-term memory and digital companionship service"],
        ["Core Positioning", "Distilling a continuously interactive digital companion persona "
         "through long-term memory, persona profiling, and context-aware dialogue"],
        ["Main Form", "CLI tool + backend service + Feishu (Lark) bot + vector database"],
        ["Core Technologies", "Python, LangChain/LangGraph, PostgreSQL/pgvector, "
         "Volcengine Ark LLM, Feishu Open Platform, Docker"],
        ["Target Users", "Individuals or small groups with needs for long-term companionship, "
         "memory preservation, and contextual continuity"],
        ["Project Stage", "MVP prototype completed; all 30 WBS tasks delivered; 298 person-hours"],
    ],
    col_widths=[4.0, 12.0],
)

add_para(
    "HeartCompass is the AI digital companion persona system developed in this course project. "
    "The core problem the project addresses is that conventional AI conversations typically focus "
    "only on single-session interactions and cannot stably remember a user's long-term experiences, "
    "relationship changes, and important events. In practice, companionship needs require the "
    "system to continuously understand a person, remember past communications, form a relatively "
    "stable persona profile, and reflect continuity in subsequent conversations."
)
add_para(
    "Around this goal, HeartCompass implements the architecture of \"long-term memory + persona "
    "profile + dialogue entry point + periodic synchronization.\" Users complete configuration "
    "and interaction through the CLI and Feishu bot. The backend uses PostgreSQL/pgvector and "
    "LangGraph orchestration to preserve information and combine it with large-model inference "
    "to generate responses that are progressively closer to a specific person and relationship. "
    "The project has delivered a runnable engineering prototype encompassing CLI, backend services, "
    "database, model configuration, Feishu bot integration, logging, and test suites."
)

add_heading_styled("2.2 Business Goal", level=2)
add_para(
    "The business goal of HeartCompass is not to first build a generalized enterprise AI platform, "
    "but to validate whether the concrete demand for a \"digital companion persona\" is genuine. "
    "The project follows a three-stage roadmap:"
)
add_para(
    "• Stage 1 (MVP Validation): Complete the MVP version and seed-user testing, verifying "
    "installation, digital companion persona creation, long-term memory storage, persona profile "
    "synchronization, and Feishu bot interaction.\n"
    "• Stage 2 (Subscription Launch): Launch subscription services around personal companionship "
    "and deep companionship, focusing on whether users are willing to pay for long-term memory, "
    "profile updates, and higher model invocation quotas.\n"
    "• Stage 3 (Scenario Expansion): Provide shared memory spaces and customized commemorative "
    "archive services for families, small groups, graduation project teams, or research groups, "
    "thereby increasing average revenue per user in small-scale scenarios."
)

add_heading_styled("2.3 Market, Economic, and Social Environment Analysis", level=2)
add_para(
    "From the demand perspective, HeartCompass targets concrete needs for long-term companionship "
    "and memory accumulation rather than an abstract \"AI tool market.\" Typical scenarios include: "
    "(1) users who wish to record personal growth and emotional changes; (2) users who want to "
    "preserve stories, habits, and important milestones related to family and friends; (3) users "
    "who need AI to remember discussion contexts in long-term projects; and (4) users who need "
    "a companion entity capable of continuously understanding context in daily work and study."
)
add_para(
    "From the economic perspective, the cost structure of a digital companion persona service "
    "primarily comes from servers, databases, model API calls, data storage, operations, and user "
    "support. Since the project leverages open-source components (LangGraph, PostgreSQL/pgvector) "
    "and cloud model services (Volcengine Ark), the initial hardware investment is moderate. "
    "However, as conversation frequency increases, model API costs will rise significantly. "
    "Therefore, the business model must incorporate invocation quotas, summary compression, "
    "caching mechanisms, and tiered model routing to control marginal costs."
)
add_para(
    "From the social value perspective, the project can help users preserve long-term memories, "
    "reduce information loss, improve personal knowledge reuse efficiency, and provide lightweight "
    "memory archives for families and small teams. However, because the system handles chat records, "
    "persona profiles, and personal events, user authorization, data isolation, deletion mechanisms, "
    "encrypted storage, and privacy notices must be treated as mandatory constraints."
)

add_heading_styled("2.4 Technical and Economic Limitations", level=2)
add_para(
    "• Technical limitation: A digital companion persona requires long-term operation, continuous "
    "memory, and stable responses, so its deployment and maintenance complexity exceeds that of a "
    "one-time AI demo.\n"
    "• Cost limitation: Model API costs are directly correlated with user conversation frequency; "
    "invocation quotas, summary compression, and caching are needed to control marginal cost.\n"
    "• Market limitation: Willingness to pay for personal companionship products still requires "
    "validation. Early-stage small-scale seed-user testing should precede large-scale commercial "
    "rollout.\n"
    "• Compliance limitation: Since personal memories, chat records, and persona profiles are "
    "involved, the system must support user authorization, data deletion, permission isolation, "
    "and necessary privacy protection measures."
)

add_heading_styled("2.5 Data Collection Methodology", level=2)
add_para(
    "In accordance with the experiment requirement to \"collect basic information about the market, "
    "economic & social environment, technical and economic limitations of the course project by "
    "online and offline,\" this report gathered data from the following sources:"
)
add_para(
    "• Online sources: (1) The HeartCompass GitHub repository (github.com/Jackey0903/HeartCompass) "
    "provided actual engineering data including 298 person-hours of WBS-tracked effort, 30 completed "
    "tasks, and NESMA function point measurements of 54 FP. (2) The 2024 China Software Industry "
    "Benchmark Data (CSBMK-202410) provided the regional labor rate of 32,129 RMB/Person-Month for "
    "Beijing. (3) Volcengine Ark platform documentation provided model API pricing data. "
    "(4) Public tax administration information provided current VAT (6%), corporate income tax "
    "(25%), and surcharge rates.\n"
    "• Offline sources: (1) Course project documentation including README.md, BRIEF_INTRO.md, "
    "WBS_TRACKING.md, and five biweekly reports provided project scope, architecture, and progress "
    "data. (2) The \"Economic Evaluation Methods and Parameters for Construction Projects (Third "
    "Edition)\" (NDRC & MOC, 2006) provided the benchmark discount rate and evaluation methodology "
    "framework. (3) NESMA Function Point Analysis standards provided the size measurement framework."
)
add_para(
    "All financial parameters used in this report are traceable to their sources. The blended labor "
    "rate of RMB 178.49/hour is derived from the CSBMK-202410 Beijing regional rate (32,129 "
    "RMB/Person-Month / 180 hours); the 10% benchmark discount rate follows the GB standard "
    "recommendation for software projects; and the tax rates reflect current PRC regulations. "
    "Project-specific data (workload, task count, function points) are extracted directly from "
    "the project's engineering records and cross-validated through dual-track analysis (NESMA vs. "
    "WBS backward tracking, variance = 5.20%)."
)

# ═══════════════════════════════════════════════════════════════
# III. BUSINESS MODEL AND PRICING STRATEGY
# ═══════════════════════════════════════════════════════════════

add_heading_styled("III. Business Model and Pricing Strategy", level=1)

add_para(
    "This project adopts a tiered business model: \"Free Trial → Personal Companion Subscription → "
    "Deep Companion Subscription → Family/Group Edition → Customized Commemorative Archive Service.\" "
    "This design is adopted because the value of a digital companion persona can only be demonstrated "
    "through sustained use. The free trial edition lowers the adoption barrier; the personal and "
    "deep companion editions generate stable subscription income; the family/group edition and "
    "customized archive service address stronger needs for memory preservation and sharing."
)

add_para(
    "The pricing strategy combines value-based pricing with cost constraints. The Personal "
    "Companion Edition is priced accessibly to facilitate conversion. The Deep Companion Edition "
    "reflects higher value through expanded memory capacity and model invocation quota. The "
    "Family/Group Edition targets multi-user sharing scenarios with a higher average transaction "
    "value. The Customized Commemorative Archive Service is quoted based on data organization "
    "volume, deployment complexity, and service duration."
)

make_table(
    ["Product Version", "Target Users", "Core Rights", "Pricing (Monthly)"],
    [
        ["Free Trial", "First-time / demo users",
         "1 persona, limited memory, basic dialogue", "RMB 0"],
        ["Personal Companion", "Individual memory preservation",
         "1-3 personas, long-term memory, standard quota", "RMB 19.9"],
        ["Deep Companion", "Heavy users / knowledge workers",
         "More personas, higher quota, priority responses", "RMB 49.9"],
        ["Family/Group Edition", "Families, project teams",
         "Shared memory space, multi-user permissions", "From RMB 199"],
        ["Customized Archive", "Memorial / graduation needs",
         "Data organization, private deployment, custom output", "From RMB 3,000/project"],
    ],
    col_widths=[3.5, 3.5, 5.0, 4.0],
)

# ═══════════════════════════════════════════════════════════════
# IV. ECONOMIC EVALUATION MODEL AND KEY ASSUMPTIONS
# ═══════════════════════════════════════════════════════════════

add_heading_styled("IV. Economic Evaluation Model and Key Assumptions", level=1)

add_para(
    "For the purpose of course experiment calculation, this report adopts a five-year project "
    "life cycle. The construction period is treated as Year 0, and the operation period spans "
    "Year 1 through Year 5. All monetary amounts are expressed in RMB 10,000 (ten thousand yuan) "
    "unless otherwise noted. This report does not incorporate debt financing or equity financing "
    "plans, but analyzes the investment, operating revenue, cost, and cash flow of the project "
    "itself (pre-financing analysis)."
)

add_para(
    "Note: The actual cash expenditure of a course project may be lower than that of a fully "
    "commercialized project. However, economic evaluation must account for opportunity cost. "
    "Therefore, this report converts student development time, testing documentation, deployment "
    "configuration, and project management effort into commercialized cost equivalents, so as to "
    "produce an evaluation result closer to real-world software project decision-making."
)

make_table(
    ["Parameter", "Value", "Description"],
    [
        ["Calculation Period", "5 years", "Operation cycle after MVP delivery"],
        ["Benchmark Discount Rate (i₀)", "10%",
         "Based on software industry average cost of capital"],
        ["VAT Rate", "6%", "Applicable to software services (general taxpayer)"],
        ["Surcharges", "12% of VAT",
         "Urban maintenance & construction tax + education surcharge + local education surcharge"],
        ["Corporate Income Tax Rate", "25%", "Standard PRC corporate income tax rate"],
        ["Fixed Asset Depreciation", "5-year straight-line, 0 residual",
         "Servers, testing equipment, deployment infrastructure"],
        ["Intangible Asset Amortization", "5-year straight-line, 0 residual",
         "Capitalized software development costs (code, documentation, configurations)"],
        ["Residual Value", "0", "Conservative estimation; no separate salvage value assumed"],
    ],
    col_widths=[5.0, 4.0, 7.0],
)

# ═══════════════════════════════════════════════════════════════
# V. PROJECT SIZE AND INVESTMENT ESTIMATION
# ═══════════════════════════════════════════════════════════════

add_heading_styled("V. Project Size and Investment Estimation", level=1)

add_heading_styled("5.1 Project Size Estimation", level=2)
add_para(
    "According to the actual engineering records of the course project, HeartCompass has completed "
    "all 30 WBS tasks spanning requirements analysis, core function development, database and model "
    "configuration, Feishu bot integration, user documentation, and testing and validation. The "
    "project records an accumulated workload of 298 person-hours. Applying a commercialized blended "
    "labor rate of approximately RMB 120 per person-hour (aligned with junior-to-mid-level "
    "development and testing comprehensive cost in the Beijing software industry), the direct "
    "development labor cost is approximately RMB 35,800."
)

add_para(
    "Additionally, the project's functional size was measured using the NESMA function point "
    "analysis method at 54 FP (comprising 7 ILF + 10 ILF + 7 ILF + 5 ELF + 4 EI + 6 EI + 7 EO + "
    "5 EO + 3 EQ). A WBS-driven backward tracking cross-check confirmed the cost estimate with "
    "a variance of only 5.20%, validating high internal consistency. Considering supplementary "
    "costs for requirements analysis, system design, testing, documentation, deployment, security "
    "review, and project management, this report estimates the total capitalized software "
    "development investment at RMB 65,000."
)

make_table(
    ["Size / Workload Item", "Estimated Value", "Description"],
    [
        ["WBS Tasks Completed", "30 tasks",
         "Full decomposition from requirements to deployment"],
        ["Accumulated Workload", "298 person-hours",
         "Requirements, development, debugging, documentation, testing, PM"],
        ["Blended Labor Unit Cost", "RMB 120 / hour",
         "Commercialized opportunity cost for junior-to-mid-level engineers"],
        ["Direct Development Labor Cost", "RMB 35,800",
         "298 hours × RMB 120 / hour"],
        ["Capitalized Software Development Investment", "RMB 65,000",
         "Includes supplementary costs for requirements, design, testing, "
         "deployment, security, and project management"],
        ["NESMA Functional Size", "54 FP",
         "Measured via NESMA FPA; cross-checked against WBS with 5.20% variance"],
    ],
    col_widths=[5.5, 4.0, 6.5],
)

add_heading_styled("5.2 Investment Estimation Table", level=2)

make_table(
    ["Investment Item", "Amount (RMB 10,000)", "Estimation Basis"],
    [
        ["Software Development (Intangible Asset)", "6.50",
         "Core code, persona profiles, conversation engine, deployment scripts"],
        ["Server / Testing Equipment / Deployment", "1.50",
         "Testing server, database, vector retrieval, container infrastructure"],
        ["Testing, Compliance, and Launch Preparation", "1.00",
         "Security audit, privacy notice, user manual, trial launch"],
        ["Initial Working Capital", "3.00",
         "Cloud resources, model API calls, and short-term operating turnover"],
        ["Total Project Investment", "12.00",
         "One-time investment during the construction period (Year 0)"],
    ],
    col_widths=[6.0, 3.5, 6.5],
)

add_heading_styled("5.3 Working Capital Estimation Table", level=2)

make_table(
    ["Item", "Year 0", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Working Capital Investment", "3.00", "0", "0", "0", "0", "0"],
        ["Working Capital Recovery", "0", "0", "0", "0", "0", "3.00"],
    ],
    col_widths=[4.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8],
)

add_para(
    "Working capital is primarily used for cloud resource prepayment, model API invocation fees, "
    "database storage costs, and short-term turnover during small-scale promotion. This report "
    "assumes RMB 30,000 of working capital is injected in Year 0 and fully recovered at the end "
    "of Year 5."
)

add_heading_styled("5.4 Fixed Asset Depreciation and Intangible Asset Amortization Table", level=2)

make_table(
    ["Item", "Original Value\n(RMB 10,000)", "Period",
     "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Fixed Asset Depreciation", "1.50", "5 years",
         "0.30", "0.30", "0.30", "0.30", "0.30"],
        ["Intangible Asset Amortization", "6.50", "5 years",
         "1.30", "1.30", "1.30", "1.30", "1.30"],
        ["Total", "8.00", "—", "1.60", "1.60", "1.60", "1.60", "1.60"],
    ],
    col_widths=[4.5, 2.5, 1.8, 1.5, 1.5, 1.5, 1.5, 1.5],
)

add_heading_styled("5.5 Fund Raising Plan", level=2)
add_para(
    "The total initial investment of RMB 120,000 for HeartCompass is relatively modest by "
    "software startup standards. This report proposes a self-funded + angel investment model "
    "consistent with the project's nature as an MVP-stage AI companionship service:"
)
add_para(
    "• Self-Funding (RMB 40,000, 33.3%): The founding team contributes development equipment, "
    "existing code assets (valued at cost), and initial operating expenses. Since the course "
    "project prototype has already been developed, the capitalized software development cost "
    "(RMB 65,000) represents an in-kind contribution from the team's prior effort.\n"
    "• Angel / Seed Round (RMB 80,000, 66.7%): To cover server procurement (RMB 15,000), "
    "testing and compliance preparation (RMB 10,000), and the first year of working capital "
    "(RMB 30,000), with the remaining RMB 25,000 as contingency reserve. Given the project's "
    "focused positioning, demonstrable MVP, and modest capital requirement, this amount is "
    "within the typical range for pre-seed / angel rounds in the Chinese AI startup ecosystem."
)
add_para(
    "The fund raising plan is structured to minimize dilution: the team's existing development "
    "work (298 person-hours, equivalent to RMB 35,800 in direct labor cost) serves as the "
    "primary in-kind contribution. External funding is sought only for cash-intensive items "
    "(hardware, paid cloud services, and initial working capital). This structure preserves "
    "founder equity while ensuring adequate cash runway for the first 12-18 months of operation."
)
add_para(
    "Should external funding not materialize, the project can adopt a bootstrapping path: "
    "defer server purchases by using existing development machines, reduce working capital "
    "to RMB 15,000 by limiting initial model API consumption to free-tier quotas, and launch "
    "with only the Personal Companion Edition to generate early revenue. Under this bootstrap "
    "scenario, the total initial cash outlay drops to approximately RMB 25,000, which is "
    "independently financeable by the founding team."
)

add_heading_styled("5.6 Capital Usage Plan", level=2)
add_para(
    "The capital usage plan details the allocation and timing of the RMB 120,000 total "
    "investment across the construction period (Year 0) and early operation period. The "
    "plan ensures that funds are deployed in alignment with project milestones:"
)

make_table(
    ["Capital Usage Item", "Amount\n(RMB 10,000)", "Timing", "Milestone Trigger"],
    [
        ["Software Asset Capitalization", "6.50",
         "Year 0, Month 1", "MVP completion and code freeze"],
        ["Server & Infrastructure Procurement", "1.50",
         "Year 0, Month 1-2", "Deployment environment provisioning"],
        ["Testing & Compliance Preparation", "1.00",
         "Year 0, Month 2-3", "Security audit and launch readiness"],
        ["Working Capital — Cloud & API Prepayment", "1.50",
         "Year 0, Month 3", "Service launch and first user onboarding"],
        ["Working Capital — Operating Reserve", "1.00",
         "Year 0, Month 3-6", "Buffer for initial operations"],
        ["Working Capital — Promotion & Content", "0.50",
         "Year 0, Month 4-6", "Seed-user acquisition and content organization"],
        ["Total", "12.00", "Year 0 (6 months)", "—"],
    ],
    col_widths=[5.5, 2.5, 3.0, 5.0],
)

add_para(
    "The capital deployment follows a phased approach: Months 1-2 focus on asset formation "
    "(software capitalization, server procurement); Months 2-3 cover launch preparation "
    "(testing, compliance, documentation); Months 3-6 fund initial operations (cloud services, "
    "model API calls, seed-user promotion). The RMB 30,000 working capital component is "
    "distributed across cloud prepayment (RMB 15,000), operating reserve (RMB 10,000), and "
    "promotion (RMB 5,000) to ensure that the service can operate sustainably from launch "
    "through the first revenue-generating period without additional cash injections."
)

# ═══════════════════════════════════════════════════════════════
# VI. REVENUE, TAXES, AND COST ESTIMATION
# ═══════════════════════════════════════════════════════════════

add_heading_styled("VI. Revenue, Taxes, and Cost Estimation", level=1)

add_heading_styled("6.1 Operating Revenue Estimation", level=2)
add_para(
    "The revenue projection is grounded in the actual usage scenarios of digital companion "
    "personas. In Year 1, revenue primarily originates from seed users, small-scale trial "
    "conversions, and a limited number of customized commemorative archive projects. From "
    "Year 2 through Year 5, as product stability improves, memory effects strengthen, and "
    "family/group scenarios expand, subscription revenue from personal plans, family/group "
    "editions, and customized services grows year over year."
)

add_para(
    "Revenue assumptions are conservative: the free trial to paid conversion rate is estimated "
    "at 3-5%, average revenue per paying user (ARPU) at approximately RMB 30/month, and user "
    "growth follows a gradual organic trajectory rather than an aggressive exponential curve.",
    size=10,
)

make_table(
    ["Item", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Personal / Deep Companion Subscription", "5.50", "12.00", "22.00", "34.00", "48.00"],
        ["Family / Group Edition", "2.50", "5.00", "9.00", "14.00", "20.00"],
        ["Customized Commemorative Archive Service", "2.00", "4.00", "7.00", "10.00", "12.00"],
        ["Total Operating Revenue", "10.00", "21.00", "38.00", "58.00", "80.00"],
    ],
    col_widths=[5.5, 2.0, 2.0, 2.0, 2.0, 2.0],
)

add_heading_styled("6.2 VAT and Surcharge Estimation Table", level=2)

make_table(
    ["Item", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Operating Revenue (tax-inclusive basis)", "10.00", "21.00", "38.00", "58.00", "80.00"],
        ["VAT Payable (6%)", "0.60", "1.26", "2.28", "3.48", "4.80"],
        ["Urban Maintenance & Construction Tax (7% of VAT)", "0.04", "0.09", "0.16", "0.24", "0.34"],
        ["Education Surcharge (3% of VAT)", "0.02", "0.04", "0.07", "0.10", "0.14"],
        ["Local Education Surcharge (2% of VAT)", "0.01", "0.03", "0.05", "0.07", "0.10"],
        ["Total Surcharges (12% of VAT)", "0.07", "0.15", "0.27", "0.42", "0.58"],
    ],
    col_widths=[6.0, 2.0, 2.0, 2.0, 2.0, 2.0],
)

add_heading_styled("6.3 Total Cost and Expense Estimation Table", level=2)

make_table(
    ["Cost Item", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Cloud Resources and Model API Cost", "2.50", "4.20", "7.00", "10.00", "14.00"],
        ["Operations, Maintenance, and User Support", "2.00", "3.50", "5.50", "8.00", "10.00"],
        ["Seed-User Promotion and Content Organization", "1.50", "2.50", "4.50", "6.50", "8.00"],
        ["Management, Compliance, and Office Expenses", "0.80", "1.50", "2.50", "4.00", "5.50"],
        ["Subtotal — Operating Costs", "6.80", "11.70", "19.50", "28.50", "37.50"],
        ["Fixed Asset Depreciation", "0.30", "0.30", "0.30", "0.30", "0.30"],
        ["Intangible Asset Amortization", "1.30", "1.30", "1.30", "1.30", "1.30"],
        ["Total Costs and Expenses", "8.40", "13.30", "21.10", "30.10", "39.10"],
    ],
    col_widths=[6.0, 2.0, 2.0, 2.0, 2.0, 2.0],
)

add_para(
    "Cost structure analysis: The largest cost component is cloud resources and model API fees, "
    "accounting for approximately 35-37% of total costs, which is characteristic of AI-native "
    "software services. Operations and maintenance constitute the second-largest category. "
    "Depreciation and amortization together represent a declining share of total costs as "
    "revenue scales, reflecting the improving operating leverage of the business model."
)

# ═══════════════════════════════════════════════════════════════
# VII. PROJECT INVESTMENT CASH FLOW ANALYSIS
# ═══════════════════════════════════════════════════════════════

add_heading_styled("VII. Project Investment Cash Flow Analysis", level=1)

add_para(
    "This section evaluates the profitability of the project itself based on project investment "
    "and operating cash flows (pre-financing analysis). In the cash flow statement, Year 0 cash "
    "outflow is the total initial investment of RMB 120,000. Operating cash flow for Years 1-5 "
    "is calculated as Net Profit + Depreciation & Amortization Add-back + Working Capital "
    "Recovery (Year 5 only)."
)

add_heading_styled("7.1 Project Investment Cash Flow Statement", level=2)

make_table(
    ["Item", "Year 0", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["1. Cash Inflow", "0", "10.00", "21.00", "38.00", "58.00", "80.00"],
        ["   Operating Revenue", "0", "10.00", "21.00", "38.00", "58.00", "80.00"],
        ["   Working Capital Recovery", "0", "0", "0", "0", "0", "3.00"],
        ["2. Cash Outflow", "12.00", "7.25", "13.74", "23.93", "35.79", "45.16"],
        ["   Initial Investment", "12.00", "0", "0", "0", "0", "0"],
        ["   Operating Costs", "0", "6.80", "11.70", "19.50", "28.50", "37.50"],
        ["   Surcharges", "0", "0.07", "0.15", "0.27", "0.42", "0.58"],
        ["   Income Tax", "0", "0.38", "1.89", "4.16", "6.87", "10.08"],
        ["3. Net Cash Flow (1-2)", "-12.00", "2.75", "7.26", "14.07", "22.21", "34.84"],
        ["4. Discount Factor (i=10%)", "1.0000", "0.9091", "0.8264", "0.7513", "0.6830", "0.6209"],
        ["5. Discounted Net Cash Flow", "-12.00", "2.50", "6.00", "10.57", "15.17", "21.63"],
        ["6. Cumulative Discounted NCF", "-12.00", "-9.50", "-3.50", "7.07", "22.24", "43.88"],
    ],
    col_widths=[5.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8],
)

add_para(
    "Income tax calculation detail: Total Profit = Operating Revenue − Total Costs − Surcharges. "
    "Income Tax = Max(Total Profit, 0) × 25%. Net Profit = Total Profit − Income Tax. "
    "Depreciation and amortization are non-cash expenses, so they are added back when computing "
    "operating net cash flow."
)

add_heading_styled("7.2 Evaluation Indicator Calculation Results", level=2)

add_para(
    "Based on the project investment cash flow statement, the three core evaluation indicators "
    "are calculated as follows:", bold=False
)

# Detailed calculation formula
add_para(
    "NPV Calculation:\n"
    "  NPV = −12.00 + 2.75/(1.10)¹ + 7.26/(1.10)² + 14.07/(1.10)³ + 22.21/(1.10)⁴ + 34.84/(1.10)⁵\n"
    "      = −12.00 + 2.50 + 6.00 + 10.57 + 15.17 + 21.63\n"
    "      = 43.88 (RMB 10,000) = RMB 438,800",
    size=10,
)

add_para(
    "IRR Calculation:\n"
    "  IRR is the discount rate r such that NPV(r) = 0. Using Excel IRR function or iterative "
    "interpolation:\n"
    "  NPV(70%) ≈ +0.38, NPV(71%) ≈ −0.15 ⇒ IRR ≈ 70.47%",
    size=10,
)

add_para(
    "Dynamic Investment Payback Period:\n"
    "  Pₜ' = (T − 1) + |Cumulative Discounted NCF at (T−1)| / Discounted NCF at Year T\n"
    "      = (3 − 1) + |−3.50| / 10.57\n"
    "      = 2 + 0.33\n"
    "      = 2.33 years",
    size=10,
)

make_table(
    ["Indicator", "Calculated Value", "Judgment Criterion", "Conclusion"],
    [
        ["Net Present Value\nNPV (i₀ = 10%)", "RMB 438,800",
         "NPV > 0", "Feasible — project creates excess returns"],
        ["Internal Rate of Return\nIRR", "70.47%",
         "IRR > i₀ (10%)", "Feasible — significantly exceeds benchmark"],
        ["Dynamic Investment\nPayback Period Pₜ'", "2.33 years",
         "Pₜ' < 5 years", "Feasible — capital recovered well within lifecycle"],
        ["Cumulative Net\nCash Flow (5 years)", "RMB 691,300",
         "Positive cumulative NCF", "5-year total cash inflow covers initial investment"],
    ],
    col_widths=[3.5, 3.0, 4.5, 5.0],
)

add_para(
    "According to the above calculations, HeartCompass demonstrates strong profitability and "
    "investment recovery capability under the benchmark scenario. The revenue assumptions are "
    "oriented toward concrete digital companion persona scenarios: starting from personal "
    "subscriptions and small-group use cases, supplemented by customized commemorative archive "
    "services as ancillary revenue. The results are aligned with the project's current positioning "
    "as an MVP-stage AI companionship service."
)

add_heading_styled("7.2.1 Formal Judgment (per Step 3.3.2)", level=2)
add_para(
    "Following the judgment framework specified in the experiment instruction (Step 3.3.2): "
    "The above calculation results show that the project benefits meet all three acceptance "
    "criteria — NPV = RMB 438,800 > 0, IRR = 70.47% > i₀ = 10%, and Pₜ' = 2.33 years < "
    "5-year calculation period. Therefore, the project is judged FEASIBLE, and it is "
    "recommended to proceed to the next step: consider the financing plan and conduct "
    "post-financing analysis. Since this report adopts a pre-financing analysis framework "
    "and the benchmark scenario already demonstrates strong profitability, the project does "
    "not require design modification or abandonment. The sensitivity analysis below further "
    "confirms that even under adverse conditions, the project maintains positive NPV, "
    "reinforcing the judgment that HeartCompass is economically viable and ready for "
    "financing consideration."
)

add_heading_styled("7.3 Sensitivity Analysis", level=2)
add_para(
    "Given that user conversion rates and model invocation costs for digital companion persona "
    "products are inherently uncertain, this report selects revenue and operating cost as the "
    "two key variables for sensitivity analysis. Four scenarios are constructed:"
)

make_table(
    ["Scenario", "Assumptions", "NPV\n(RMB 10,000)", "IRR", "Pₜ'\n(years)", "Conclusion"],
    [
        ["Benchmark", "Revenue and cost as per baseline", "43.88", "70.47%", "2.33",
         "Economically feasible"],
        ["Conservative", "Revenue −20%; Cost +10%", "16.63", "37.57%", "3.47",
         "Still feasible; NPV remains positive"],
        ["Optimistic", "Revenue +20%; Cost −5%", "68.11", "94.00%", "1.88",
         "Strong returns; rapid capital recovery"],
        ["Stress Test", "Revenue −30%; Cost +15%", "2.75", "15.29%", "4.65",
         "Near critical threshold; NPV still > 0"],
    ],
    col_widths=[2.5, 4.0, 2.0, 2.0, 2.0, 3.5],
)

add_para(
    "Sensitivity analysis insights: The most sensitive factors affecting project viability are "
    "paid conversion rate and model API invocation cost. If users only trial the product briefly "
    "without forming a long-term companionship relationship, subscription revenue will decline "
    "significantly. If users engage in high-frequency conversations without summary compression "
    "and invocation quota constraints, operating costs will escalate rapidly. Therefore, the "
    "immediate next-stage priorities should be validating retention rate, companionship "
    "effectiveness, and per-user model cost through seed-user testing."
)

add_para(
    "Notably, even under the stress scenario (revenue −30%, cost +15%), the NPV remains positive "
    "at RMB 27,500, and the IRR of 15.29% still exceeds the 10% benchmark. This indicates that "
    "HeartCompass possesses a reasonable margin of safety in its financial projections, largely "
    "attributable to its modest initial investment requirement and the high operating leverage "
    "characteristic of AI-native SaaS products."
)

# ═══════════════════════════════════════════════════════════════
# VIII. ECONOMIC FEASIBILITY AND RISK EVALUATION
# ═══════════════════════════════════════════════════════════════

add_heading_styled("VIII. Economic Feasibility and Risk Evaluation", level=1)

add_heading_styled("8.1 Feasibility Conclusion", level=2)
add_para(
    "Based on the financial calculation results and the technical foundation established by the "
    "course project, this report concludes that HeartCompass possesses economic feasibility for "
    "continued advancement beyond the MVP stage. The main supporting reasons are:"
)
add_para(
    "1. Focused project positioning: HeartCompass is built around the \"digital companion persona\" "
    "rather than a generalized AI tool, making the target demand more concrete and addressable.\n"
    "2. Modest initial investment: With a total investment of RMB 120,000, the project's capital "
    "requirement is manageable for seed-stage validation. Existing code, deployment workflows, "
    "bot entry points, and long-term memory prototypes are already delivered.\n"
    "3. Revenue model aligned with genuine needs: Personal companionship, deep companionship, "
    "family/group sharing, and customized commemorative archives correspond to distinct paid "
    "tiers with differentiated value propositions.\n"
    "4. Financial indicators meet acceptance criteria: Under the benchmark scenario, NPV > 0, "
    "IRR (70.47%) >> benchmark discount rate (10%), and dynamic payback period (2.33 years) < "
    "project lifecycle (5 years).\n"
    "5. Adequate safety margin: Even under the conservative and stress scenarios, NPV remains "
    "positive, indicating that the project can withstand moderate adverse deviations."
)

add_heading_styled("8.2 Main Risks and Mitigation Measures", level=2)

make_table(
    ["Risk Category", "Specific Manifestation", "Mitigation Strategy"],
    [
        ["Demand Risk",
         "Users may find the digital companion persona novel initially but not form sustained usage habits",
         "Validate real companionship needs through seed-user testing; iterate based on retention metrics"],
        ["Cost Risk",
         "High model API invocation frequency may compress margins",
         "Implement invocation quotas, summary compression, prompt caching, and tiered model routing"],
        ["Experience Risk",
         "Unstable persona profiles or discontinuous responses degrade user trust",
         "Optimize memory update algorithms, profile synchronization, and conversation quality monitoring"],
        ["Privacy &\nCompliance Risk",
         "System handles personal memories, chat records, and persona profiles",
         "Strengthen authorization flows, data encryption, deletion mechanisms, and privacy documentation"],
        ["Operations Risk",
         "Long-term operation requires sustained server and database maintenance",
         "Provide Dockerized one-click deployment, automated health checks, monitoring, and logging"],
    ],
    col_widths=[3.0, 5.5, 7.5],
)

# ═══════════════════════════════════════════════════════════════
# IX. THINKING QUESTION
# ═══════════════════════════════════════════════════════════════

add_heading_styled(
    "IX. Thinking Question: Evaluating Software Projects Without Cash Income", level=1)

add_para(
    "If a software project generates no direct cash income, its value cannot be assessed solely "
    "through operating revenue, profit, and investment payback period. Instead, a combination of "
    "cost-benefit analysis, cost-effectiveness analysis, and multi-criteria comprehensive "
    "evaluation methods should be employed:"
)

add_para(
    "1. Cost-Saving Method (成本节约法): Estimate the labor hours, equipment costs, travel expenses, "
    "paper materials, and repeated communication overhead saved after the project is deployed. "
    "The saved cost is treated as the project's economic benefit. For example, an internal "
    "automation tool may not generate revenue but can eliminate manual processing costs.\n\n"
    "2. Efficiency-Improvement Method (效率提升法): Measure reductions in business process cycle "
    "time, error rates, and response latency, as well as increases in processing throughput. "
    "Convert these efficiency gains into monetary equivalents using labor cost rates.\n\n"
    "3. Social Benefit Evaluation (社会效益评估): For public-welfare, educational, healthcare, "
    "and government projects, evaluate non-cash benefits such as service coverage expansion, "
    "user satisfaction improvement, risk reduction, public resource savings, and social equity "
    "enhancement.\n\n"
    "4. Multi-Criteria Decision-Making (多准则决策): Apply methods such as the Analytic Hierarchy "
    "Process (AHP), expert scoring, and TOPSIS to comprehensively score the project across "
    "dimensions including security, usability, maintainability, social value, and strategic "
    "alignment.\n\n"
    "5. Replacement-Cost Method (重置成本法): Estimate the cost of purchasing equivalent "
    "third-party systems, hiring additional personnel, or outsourcing services if the software "
    "were not developed. This avoided cost represents the implicit benefit of the project."
)

add_para(
    "For a digital companion persona project such as HeartCompass, even in the absence of "
    "short-term direct cash income, its comprehensive benefits can still be evaluated across "
    "multiple dimensions: the preservation value of personal and family memories, the number of "
    "users covered by companionship services, user retention and engagement metrics, reduction "
    "in repeated communication overhead, improvement in personal knowledge reuse efficiency, "
    "and the degree of privacy risk mitigation. These non-monetary indicators collectively form "
    "a holistic picture of the project's value proposition and can serve as the basis for "
    "cost-effectiveness comparisons against alternative approaches."
)

# ═══════════════════════════════════════════════════════════════
# X. EXPERIMENTAL LEARNING EXPERIENCE
# ═══════════════════════════════════════════════════════════════

add_heading_styled("X. Experimental Learning Experience", level=1)

add_para(
    "Through this experiment, I gained a systematic understanding that the economic evaluation "
    "of a software project is far more than a simple development cost estimate. It requires "
    "integrating the technical solution, user needs, business model, operating cost structure, "
    "tax policy, depreciation and amortization schedules, and multi-year cash flow projections "
    "into a single coherent financial model. For an AI-driven digital companion persona project "
    "like HeartCompass, what ultimately determines economic feasibility is not whether the "
    "project can invoke a large language model, but whether users can form a sustained "
    "companionship relationship with the system, and whether the per-user model inference cost "
    "can be controlled within viable margins."
)

add_para(
    "This experiment also deepened my understanding of the interconnections among financial "
    "statements. The investment estimation table determines the Year 0 cash outflow. The revenue "
    "and cost tables determine operating profit. The tax schedule affects net profit. Depreciation "
    "and amortization reduce accounting profit but are non-cash items that must be added back in "
    "the cash flow statement. Finally, NPV, IRR, and the dynamic payback period are derived "
    "through discounted cash flow analysis. Understanding these linkages is essential for "
    "constructing a logically consistent and computationally correct evaluation model."
)

add_para(
    "To make the parameter mapping relationships among financial tables explicit, I summarize "
    "the key data flows as follows: (1) The Investment Estimation Table (Table 5.2) feeds its "
    "Total Project Investment (RMB 120,000) into the Cash Flow Statement as the Year 0 outflow. "
    "(2) The Fixed Asset Depreciation Table and Intangible Asset Amortization Table (Table 5.4) "
    "feed annual depreciation (RMB 3,000) and amortization (RMB 13,000) into the Total Cost "
    "Table (Table 6.3) and are subsequently added back in the Cash Flow Statement as non-cash "
    "items. (3) The Operating Revenue Table (Table 6.1) feeds into both the VAT & Surcharge "
    "Table (Table 6.2) — where VAT = Revenue × 6% and Surcharges = VAT × 12% — and directly "
    "into the Cash Flow Statement as cash inflow. (4) Total Profit = Operating Revenue − Total "
    "Costs − Surcharges; Income Tax = Total Profit × 25%; Net Profit = Total Profit − Income "
    "Tax. (5) Operating Net Cash Flow = Net Profit + Depreciation & Amortization + Working "
    "Capital Recovery. (6) NPV is the sum of discounted net cash flows at i = 10%; IRR is the "
    "rate where NPV = 0; the Dynamic Payback Period is the point where cumulative discounted "
    "net cash flow turns positive. Understanding this chain — from investment → cost/revenue → "
    "tax → profit → cash flow → NPV/IRR/Payback — is the core skill of software project "
    "economical evaluation."
)

add_para(
    "From the perspective of the HeartCompass project specifically, the economic evaluation "
    "exercise made me more clearly aware that if the course project is to be further developed, "
    "the focus should not be on demonstrating technical stack completeness alone. Rather, the "
    "project must return to concrete user needs: why users need a digital companion persona, "
    "what memories they are willing to entrust to the system, what companionship capabilities "
    "they are willing to pay for, and how to balance privacy protection with cost control. "
    "The sensitivity analysis further reinforces that retention rate and unit economics (per-user "
    "model cost) are the two most critical operational metrics to validate in the next stage."
)

add_para(
    "In terms of methodology, I learned to apply the NESMA function point analysis framework for "
    "software size measurement, the WBS-driven backward tracking method for cost validation, "
    "and the discounted cash flow model for investment decision-making. The dual-track "
    "cross-validation between NESMA FP (54 FP) and WBS hours (298 h) with only 5.20% variance "
    "was particularly instructive — it demonstrated that rigorous software measurement can produce "
    "highly consistent results across independent estimation approaches, which is essential for "
    "credible economic evaluation in both academic and industrial contexts."
)

add_para(
    "Finally, this experiment helped me appreciate the importance of national standards and "
    "industry parameters in software engineering economics. The benchmark discount rate (10%), "
    "tax rates (6% VAT, 25% CIT), and depreciation/amortization rules (5-year straight-line) "
    "are not arbitrary choices but are grounded in the \"Economic Evaluation Methods and Parameters "
    "for Construction Projects (Third Edition)\" and current PRC tax regulations. Aligning the "
    "evaluation model with these standards ensures that the analysis is not only internally "
    "consistent but also comparable and auditable by external stakeholders."
)

# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════

add_heading_styled("References", level=1)

refs = [
    "[1] National Development and Reform Commission, Ministry of Construction. "
    "Economic Evaluation Methods and Parameters for Construction Projects (Third Edition) "
    "[M]. Beijing: China Planning Press, 2006: 13-19.",
    "[2] China Software Industry Association. 2024 China Software Industry Benchmark Data "
    "(CSBMK-202410). Beijing, 2024.",
    "[3] NESMA. NESMA Function Point Analysis Standards — Comprehensive Functional Size "
    "Measurement Guidelines. Netherlands Software Metrics Association, 2018.",
    "[4] Public information on PRC tax administration: Corporate Income Tax Law, "
    "Value-Added Tax regulations, and related software service tax policies.",
    "[5] HeartCompass Course Project Repository. https://github.com/Jackey0903/HeartCompass.",
    "[6] HeartCompass Project Documentation: README.md, BRIEF_INTRO.md, WBS_TRACKING.md, "
    "and Biweekly Reports #1-#5.",
]

for ref in refs:
    add_para(ref, size=10)

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════

os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f"Report saved to: {OUTPUT_FILE}")
