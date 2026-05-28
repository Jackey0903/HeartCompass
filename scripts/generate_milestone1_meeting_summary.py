from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from PIL import ImageChops
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "meeting"
DOCX_PATH = OUTPUT_DIR / "Milestone 1 Review Meeting Summary - HeartCompass.docx"
PDF_PATH = OUTPUT_DIR / "Milestone 1 Review Meeting Summary - HeartCompass.pdf"
SIGNATURE_ASSET_DIR = OUTPUT_DIR / "signature_assets"
CHART_ASSET_DIR = OUTPUT_DIR / "chart_assets"

TITLE = "Milestone 1 Review Meeting Summary"
SUBTITLE = "HeartCompass / Digital Immortality"
DOCUMENT_DATE = "May 9, 2026"
MEETING_TIME = "May 9, 2026, 4:00 PM"
MEETING_DURATION = "44 minutes 42 seconds"
TEAM = "Leishiyu Chen, Haojie Hu, Tian Bu, Yihan Liu"
PRESENTER = "Tian Bu"

PALETTE = {
    "navy": "#1F4E79",
    "slate": "#456B84",
    "teal": "#5D8AA8",
    "pale": "#DDE9F2",
    "soft": "#F5F8FB",
    "grid": "#D3DCE6",
    "text": "#23313D",
    "muted": "#6B7C8C",
    "green": "#7BAF8D",
    "amber": "#D9A441",
}


ORIGINAL_RBS = [
    ("R1", "User and Relationship Data Foundation", "User registration, profile management, MBTI collection, Crush and RelationChain setup, ownership and validation rules."),
    ("R2", "Context Acquisition and Knowledge Management", "Natural language narratives, screenshots, curated knowledge, Event, ChatTopic, DerivedInsight, Knowledge, and ContextEmbedding records."),
    ("R3", "Relationship Analysis Engine", "Communication suggestions, risk prompts, narrative analysis, conversation analysis, and interaction signals."),
    ("R4", "Virtual Figure Interaction System", "Persona context loading, memory recall, WebSocket interaction, timed message batching, and context refresh."),
    ("R5", "Client Experience, Integration, and Deployment", "Lightweight client, product integration, testing, performance optimization, packaging, and final delivery."),
]

UPDATED_RBS = [
    ("R1.1", "User and Figure-Relation Data Foundation", "User identity, authentication, profile management, MBTI collection, FigureAndRelation records, FigureRole assignment, and ownership validation."),
    ("R1.2", "Persona Building and Knowledge Management", "FRBuildingGraph preprocessing, OriginalSource persistence, FineGrainedFeed extraction, conflict tracking, update logs, embeddings, and semantic recall."),
    ("R2.1", "Conversation and Dialogue System", "ConversationGraph persona loading, service-backed recall, prompt assembly, response generation, short-term memory trimming, rolling summary, and feed sync."),
    ("R2.2", "Channel Integration and Lark Bot", "Lark WebSocket access, open_id mapping, timed message batching, bot commands, card output, FR switching, and busy-state feedback."),
    ("R3.1", "CLI, Harness, Packaging, and Deployment", "immortality CLI, doctor/setup/auth/fr/logs/lark-service commands, JSON contracts, Docker/Alembic automation, PyPI packaging, Harness skills, testing, and deployment."),
]

ORIGINAL_WBS = [
    ("WP1", "User and Relationship Data Foundation", "Data model and user/relationship baseline", "Mar 30-Apr 26", 84),
    ("WP2", "Context Acquisition and Knowledge Management", "Context ingestion, screenshot/topic extraction, knowledge and embeddings", "Apr 13-Apr 26", 120),
    ("WP3", "Relationship Analysis Engine", "Narrative analysis, conversation analysis, scoring and guidance", "Apr 20-May 10", 132),
    ("WP4", "Virtual Figure Interaction System", "Persona loading, memory recall, WebSocket interaction and virtual figure dialogue", "Apr 27-May 17", 126),
    ("WP5", "Client Experience, Integration, and Deployment", "Lightweight client, integration, testing, optimization, and deployment", "Apr 27-Jun 14", 162),
]

UPDATED_WBS = [
    ("WP1.1", "User registration and authentication", "R1.1", "W3", 20),
    ("WP1.2", "Profile management and MBTI collection", "R1.1", "W3", 22),
    ("WP1.3", "FigureAndRelation CRUD and FigureRole", "R1.1", "W4", 24),
    ("WP1.4", "Data foundation review and test", "R1.1", "W4", 18),
    ("WP2.1", "FRBuildingGraph preprocessing pipeline", "R1.2", "W4-W5", 28),
    ("WP2.2", "Fine-grained feed extraction and storage", "R1.2", "W5", 26),
    ("WP2.3", "Conflict detection and update logging", "R1.2", "W5-W6", 22),
    ("WP2.4", "Embedding generation and vector recall", "R1.2", "W6", 26),
    ("WP2.5", "Persona building review and test", "R1.2", "W6", 18),
    ("WP3.1", "Service-backed ConversationGraph recall", "R2.1", "W5-W6", 28),
    ("WP3.2", "Short-term memory trim and summary", "R2.1", "W6-W7", 30),
    ("WP3.3", "LLM call and response generation", "R2.1", "W7", 28),
    ("WP3.4", "Feed sync and prompt alignment", "R2.1", "W7-W8", 26),
    ("WP3.5", "Conversation system review and test", "R2.1", "W8", 20),
    ("WP4.1", "Lark WebSocket and message handling", "R2.2", "W6-W7", 30),
    ("WP4.2", "Bot menu and persona build commands", "R2.2", "W7", 26),
    ("WP4.3", "Message batching and card templates", "R2.2", "W7-W8", 24),
    ("WP4.4", "Service-only data access alignment", "R2.2", "W8", 26),
    ("WP4.5", "Busy-state cards and integration test", "R2.2", "W8", 20),
    ("WP5.1", "CLI framework, doctor/setup, and logs", "R3.1", "W7-W8", 32),
    ("WP5.2", "CLI auth, FR management, and JSON contracts", "R3.1", "W8-W9", 30),
    ("WP5.3", "Docker support and database automation", "R3.1", "W9", 30),
    ("WP5.4", "Graph/channel/CLI regression testing", "R3.1", "W9-W10", 30),
    ("WP5.5", "PyPI packaging and publishing", "R3.1", "W10", 22),
    ("WP5.6", "Final docs, Harness assets, and deployment", "R3.1", "W11", 18),
]

UPDATED_WP_SUMMARY = [
    ("WP1", "User and Figure-Relation Data Foundation", "R1.1", 84),
    ("WP2", "Persona Building and Knowledge Management", "R1.2", 120),
    ("WP3", "Conversation and Dialogue System", "R2.1", 132),
    ("WP4", "Channel Integration and Lark Bot", "R2.2", 126),
    ("WP5", "CLI, Harness, Packaging, and Deployment", "R3.1", 162),
]

SCHEDULE_UPDATE = [
    ("W1", "Mar 30-Apr 05", "P0: Planning and Requirements"),
    ("W2", "Apr 06-Apr 12", "P1: System Design"),
    ("W3", "Apr 13-Apr 19", "WP1.1, WP1.2"),
    ("W4", "Apr 20-Apr 26", "WP1.3, WP1.4, WP2.1 begins"),
    ("W5", "Apr 27-May 03", "WP2.1, WP2.2, WP3.1 begins"),
    ("W6", "May 04-May 10", "WP2.3, WP2.4, WP2.5, WP3.1, WP4.1 begins"),
    ("W7", "May 11-May 17", "WP3.2, WP3.3, WP4.1, WP4.2"),
    ("W8", "May 18-May 24", "WP3.4, WP3.5, WP4.3, WP4.4, WP4.5, WP5.1"),
    ("W9", "May 25-May 31", "WP5.2, WP5.3"),
    ("W10", "Jun 01-Jun 07", "WP5.4, WP5.5"),
    ("W11", "Jun 08-Jun 14", "WP5.6"),
]

REVIEWER_COMMENTS = [
    ("Ding Zhenyao", "Information retention and iteration", "Asked whether the system can preserve original persona information during later iterations and avoid losing foundational facts.", "The team explained that stored persona data will remain in the database, while corrections or missing information can be handled through the persona-building update flow."),
    ("Chen Yiming", "Differentiation from existing persona skills", "Asked how the project differs from existing colleague/ex-partner/persona-distillation skills.", "The team clarified that existing skills mainly produce files for another agent to consume, while this project provides an end-to-end engineered product with database storage, vector recall, CLI/Lark access, and no dependency on a general-purpose external agent workflow."),
    ("Li Ranxi", "Project management artifacts", "Praised the project architecture and delivery thinking, but noted that the presentation spent more time on technical design than on RBS, WBS, planned hours, and schedule progress.", "The team accepted this as a key improvement area and updated the RBS, WBS, schedule, and task-hour allocation for the next submission."),
    ("Huang Yihe", "Persona consistency and memory realism", "Suggested using an LLM to periodically review indirect logical conflicts or self-contradictions in persona items. Also suggested making memory more realistic by considering time decay and event importance.", "The team identified this as a next-milestone action item for persona health review, conflict inspection, and memory-decay design."),
    ("Qian Baoqiang", "Message buffering responsiveness", "Pointed out that the 15-second Lark message buffer may reduce responsiveness, because users may want a timely reply even while sending a series of short messages.", "The team will evaluate adaptive batching, early response triggers, and configurable waiting time instead of relying only on a fixed 15-second reset window."),
    ("Shangguan Siyang", "Ethical boundary and abuse prevention", "Asked whether the system has usage boundaries for simulating real people and how misuse can be prevented.", "The team discussed role-based extraction limits, especially stricter boundaries for colleagues, mentors, and public figures, and the need to restrict full persona export after the system becomes more mature."),
    ("Zhao Zhuobing", "Next-step roadmap", "Asked about the next planned work after the current milestone.", "The team listed single-Lark-bot limitations, shared database support, lower setup barriers, built-in public figures, prompt optimization, Harness work, new correction/analysis/decision graphs, log governance, and external long-term memory integration."),
]

SIGNATURES = [
    ("Ding Zhenyao", ROOT / "output" / "name" / "丁桢垚.png"),
    ("Chen Yiming", ROOT / "output" / "name" / "陈奕名.jpg"),
    ("Li Ranxi", ROOT / "output" / "name" / "李冉曦.jpg"),
    ("Huang Yihe", ROOT / "output" / "name" / "黄一和.jpg"),
    ("Qian Baoqiang", ROOT / "output" / "name" / "钱宝强.jpg"),
    ("Shangguan Siyang", ROOT / "output" / "name" / "上官思洋.png"),
    ("Zhao Zhuobing", ROOT / "output" / "name" / "赵卓冰.jpg"),
]

ACTION_ITEMS = [
    ("A1", "Update RBS, WBS, schedule, and task-hour allocation", "Revise project-management artifacts to show original and updated scope, task-level decomposition, planned hours, and total hours. This action aligns with WP5.6 final documentation and governance assets.", "Tian Bu (team leader)", "W11 / before Canvas submission", "Completed"),
    ("A2", "Add persona consistency review", "Design a periodic LLM-based review flow to detect indirect conflicts, contradictions, stale facts, and unresolved FineGrainedFeedConflict records. This action aligns with WP2.3 conflict detection and WP2.5 persona-building review.", "Leishiyu Chen", "W6 / Milestone 2", "Planned"),
    ("A3", "Improve memory realism", "Introduce memory-decay thinking, event importance, and time-aware recall behavior so long-term memories can feel more realistic. This action aligns with WP3.1 service-backed recall and WP3.2 short-term memory trim and summary.", "Haojie Hu", "W6-W7 / Milestone 2", "Planned"),
    ("A4", "Optimize Lark message buffering", "Evaluate adaptive batching, shorter waits, manual reply triggers, and configuration options for WAITING_SECONDS_FOR_CONVERSATION. This action aligns with WP4.1 Lark WebSocket handling and WP4.3 message batching.", "Yihan Liu", "W6-W8 / Milestone 2", "Planned"),
    ("A5", "Define ethical and role-based boundaries", "Document what can be extracted, recalled, displayed, or exported for self, family, friends, colleagues, mentors, public figures, and strangers. This action aligns with WP1.3 FigureRole management and WP4.2 bot menu permissions.", "Tian Bu (team leader)", "W4-W7 / Milestone 2", "Planned"),
    ("A6", "Lower setup barriers with shared database design", "Investigate encrypted shared database support, gateway-based access, and protection against exposing database URIs to users. This action aligns with WP5.3 Docker/database automation and WP5.4 regression testing.", "Leishiyu Chen", "W9-W10 / Milestone 2-Milestone 3", "Planned"),
    ("A7", "Address single Lark Bot limitation", "Design multi-bot or multi-service support so a single installed package is not limited to one Lark Bot proxy. This action aligns with WP4.4 service-only data access alignment and WP4.5 channel integration testing.", "Yihan Liu", "W8 / Milestone 2", "Planned"),
    ("A8", "Strengthen log governance", "Reduce heavy ConversationGraph logging, preserve necessary checkpoint observability, and avoid latency from excessive log payloads. This action aligns with WP5.1 CLI logs and WP5.4 graph/channel/CLI regression testing.", "Haojie Hu", "W8-W10 / Milestone 2", "Planned"),
]

MAIN_CONTENT = [
    "The presenter introduced Digital Immortality as a persona-centric AI project built to solve the consumption problem left by file-based persona-distillation skills. Existing skills can extract personality, interaction, procedural, and memory materials, but they tend to grow into large files that must be read in full or nearly in full. The project therefore aims to engineer persona construction and persona consumption into a maintainable system.",
    "The team explained the current product form: a CLI-first management tool and a Lark Bot for daily use. The team intentionally does not prioritize a traditional web front end at this stage, because the main challenge is the AI workflow, memory structure, data layer, and low-friction interaction through an IM channel.",
    "The core architecture was presented as layered storage plus vectorized recall. Lightweight intrinsic fields are stored directly in FigureAndRelation, while large, high-context persona facts are stored as FineGrainedFeed records by dimension and retrieved semantically according to the current conversation context.",
    "Two main graph workflows were discussed. FRBuildingGraph receives raw user materials, cleans and preprocesses them, extracts intrinsic fields and fine-grained feeds, compares new values against stored values, records conflicts, and writes update logs. ConversationGraph receives batched Lark messages, loads persona context, recalls relevant feeds, manages short-term memory, trims and summarizes older rounds, and generates persona-grounded replies.",
    "The meeting also covered the rationale for LangGraph rather than a pure ReAct agent. The team emphasized that persona construction and dialogue generation require controllable, observable, and repeatable steps instead of a black-box autonomous agent.",
    "The team introduced Harness work as a way to convert repeated engineering practices into reusable skills and rules. Examples included code-quality review, commit-update writing, document generation, document optimization, graph documentation, and session-practice learning.",
]


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    SIGNATURE_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    CHART_ASSET_DIR.mkdir(parents=True, exist_ok=True)


def total_hours(rows: list[tuple]) -> int:
    return sum(int(row[-1]) for row in rows)


def build_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleCenter", fontName="Helvetica-Bold", fontSize=21, leading=26, alignment=TA_CENTER, textColor=colors.HexColor(PALETTE["navy"]), spaceAfter=8))
    styles.add(ParagraphStyle("SubtitleCenter", fontName="Helvetica", fontSize=10.5, leading=14, alignment=TA_CENTER, textColor=colors.HexColor(PALETTE["muted"]), spaceAfter=8))
    styles.add(ParagraphStyle("Section", fontName="Helvetica-Bold", fontSize=15, leading=19, textColor=colors.HexColor(PALETTE["navy"]), spaceBefore=10, spaceAfter=7))
    styles.add(ParagraphStyle("Subsection", fontName="Helvetica-Bold", fontSize=11.5, leading=15, textColor=colors.HexColor(PALETTE["text"]), spaceBefore=7, spaceAfter=4))
    styles.add(ParagraphStyle("Body", fontName="Helvetica", fontSize=9.8, leading=14, textColor=colors.HexColor(PALETTE["text"]), spaceAfter=5))
    styles.add(ParagraphStyle("Small", fontName="Helvetica", fontSize=8.2, leading=10.5, textColor=colors.HexColor(PALETTE["text"])))
    styles.add(ParagraphStyle("TableHead", fontName="Helvetica-Bold", fontSize=8, leading=9.5, textColor=colors.white, alignment=TA_CENTER))
    styles.add(ParagraphStyle("TableCell", fontName="Helvetica", fontSize=7.7, leading=9.5, textColor=colors.HexColor(PALETTE["text"]), alignment=TA_LEFT))
    styles.add(ParagraphStyle("TableCellSmall", fontName="Helvetica", fontSize=7.2, leading=8.8, textColor=colors.HexColor(PALETTE["text"]), alignment=TA_LEFT))
    return styles


def normalize_signature_assets() -> list[tuple[str, Path]]:
    normalized = []
    for idx, (name, src) in enumerate(SIGNATURES, 1):
        image = PILImage.open(src).convert("RGB")
        background = PILImage.new("RGB", image.size, "white")
        diff = ImageChops.difference(image, background)
        bbox = diff.getbbox()
        if bbox:
            pad = 24
            left = max(0, bbox[0] - pad)
            top = max(0, bbox[1] - pad)
            right = min(image.width, bbox[2] + pad)
            bottom = min(image.height, bbox[3] + pad)
            image = image.crop((left, top, right, bottom))

        out = SIGNATURE_ASSET_DIR / f"signature_{idx:02d}.png"
        image.save(out)
        normalized.append((name, out))
    return normalized


def fitted_reportlab_image(path: Path, max_width_mm: float, max_height_mm: float) -> RLImage:
    with PILImage.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    max_width = max_width_mm * mm
    max_height = max_height_mm * mm
    width = max_width
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    result = RLImage(str(path), width=width, height=height)
    result.hAlign = "CENTER"
    return result


def trim_white_borders(path: Path, padding: int = 12) -> None:
    image = PILImage.open(path).convert("RGB")
    background = PILImage.new("RGB", image.size, "white")
    diff = ImageChops.difference(image, background)
    bbox = diff.getbbox()
    if not bbox:
        return
    left = max(0, bbox[0] - padding)
    top = max(0, bbox[1] - padding)
    right = min(image.width, bbox[2] + padding)
    bottom = min(image.height, bbox[3] + padding)
    image.crop((left, top, right, bottom)).save(path)


def dot_label(text: str) -> str:
    return text.replace("\\", "\\\\").replace('"', '\\"')


def create_updated_rbs_chart(path: Path) -> None:
    dot_path = path.with_suffix(".dot")
    dot_path.write_text(
        f"""
digraph G {{
  graph [rankdir=TB, splines=polyline, nodesep=0.50, ranksep=0.85, bgcolor="white", pad=0.28];
  node [shape=box, style="rounded,filled", fontname="Helvetica", fontsize=14, color="{PALETTE["navy"]}", fillcolor="white", margin="0.20,0.12"];
  edge [color="{PALETTE["slate"]}", penwidth=1.25];

  Root [label="HeartCompass\\nUpdated RBS", fillcolor="{PALETTE["pale"]}", fontsize=16];
  R1 [label="R1: Foundation and\\nPersona Input", fillcolor="{PALETTE["soft"]}"];
  R2 [label="R2: Dialogue and\\nChannel", fillcolor="{PALETTE["soft"]}"];
  R3 [label="R3: CLI, Harness,\\nand Delivery", fillcolor="{PALETTE["soft"]}"];

  R11 [label="R1.1: User and Figure-Relation\\nData Foundation", fillcolor="{PALETTE["soft"]}"];
  R12 [label="R1.2: Persona Building\\nand Knowledge Mgmt", fillcolor="{PALETTE["soft"]}"];
  R21 [label="R2.1: Conversation and\\nDialogue System", fillcolor="{PALETTE["soft"]}"];
  R22 [label="R2.2: Channel Integration\\nand Lark Bot", fillcolor="{PALETTE["soft"]}"];
  R31 [label="R3.1: CLI, Harness,\\nPackaging, and Delivery", fillcolor="{PALETTE["soft"]}"];

  R111 [label="R1.1.1: User identity\\nand authentication"];
  R112 [label="R1.1.2: FigureAndRelation\\nand FigureRole mgmt"];
  R121 [label="R1.2.1: FRBuildingGraph\\nand feed extraction"];
  R122 [label="R1.2.2: Conflict detection\\nand embedding storage"];
  R211 [label="R2.1.1: ConversationGraph\\nand feed recall"];
  R212 [label="R2.1.2: Memory trim\\nand feed sync"];
  R221 [label="R2.2.1: Lark WebSocket\\nand message handling"];
  R222 [label="R2.2.2: Bot commands\\nand card templates"];
  R311 [label="R3.1.1: CLI framework,\\nlogs, and commands"];
  R312 [label="R3.1.2: Docker, PyPI,\\nHarness, and deployment"];

  Root -> R1; Root -> R2; Root -> R3;
  R1 -> R11; R1 -> R12;
  R2 -> R21; R2 -> R22;
  R3 -> R31;

  R11 -> R111; R11 -> R112;
  R12 -> R121; R12 -> R122;
  R21 -> R211; R21 -> R212;
  R22 -> R221; R22 -> R222;
  R31 -> R311; R31 -> R312;

  {{ rank=same; R1; R2; R3; }}
  {{ rank=same; R11; R12; R21; R22; R31; }}
}}
""".strip(),
        encoding="utf-8",
    )
    subprocess.run(["dot", "-Tpng", "-Gdpi=220", str(dot_path), "-o", str(path)], check=True)
    trim_white_borders(path, padding=8)


def create_updated_wbs_chart(path: Path) -> None:
    dot_path = path.with_suffix(".dot")
    tasks_by_wp: dict[str, list[tuple[str, str, str, str, int]]] = {}
    for task_id, task_name, rbs, weeks, hours in UPDATED_WBS:
        wp_id = task_id.split(".")[0]
        tasks_by_wp.setdefault(wp_id, []).append((task_id, task_name, rbs, weeks, hours))

    cluster_blocks = []
    for wp_id, wp_name, rbs_ref, hours in UPDATED_WP_SUMMARY:
        task_lines = []
        for task_id, task_name, _rbs, weeks, task_hours in tasks_by_wp[wp_id]:
            node_id = task_id.replace(".", "_")
            short_name = task_name if len(task_name) <= 34 else task_name[:31] + "..."
            label = f"{task_id}\\n{dot_label(short_name)}\\n{weeks}, {task_hours}h"
            task_lines.append(f'    {node_id} [label="{label}"];')
            task_lines.append(f"    {wp_id} -> {node_id};")
        cluster_blocks.append(
            f'  subgraph cluster_{wp_id} {{\n'
            f'    style=rounded; color="{PALETTE["grid"]}"; bgcolor="{PALETTE["soft"]}";\n'
            f'    label="";\n'
            f'    {wp_id} [label="{wp_id}: {dot_label(wp_name[:32])}\\n[{rbs_ref}] {hours}h", fillcolor="{PALETTE["pale"]}", fontsize=12];\n'
            + "\n".join(task_lines)
            + "\n  }"
        )

    dot_path.write_text(
        f"""
digraph G {{
  graph [rankdir=LR, splines=polyline, nodesep=0.25, ranksep=0.70, bgcolor="white", pad=0.25];
  node [shape=box, style="rounded,filled", fontname="Helvetica", fontsize=10, color="{PALETTE["navy"]}", fillcolor="white", margin="0.10,0.07"];
  edge [color="{PALETTE["slate"]}", penwidth=1.0];

  Root [label="HeartCompass\\nUpdated WBS\\n624h", fillcolor="{PALETTE["pale"]}", fontsize=14];

{chr(10).join(cluster_blocks)}

  Root -> WP1;
  Root -> WP2;
  Root -> WP3;
  Root -> WP4;
  Root -> WP5;
}}
""".strip(),
        encoding="utf-8",
    )
    subprocess.run(["dot", "-Tpng", "-Gdpi=200", str(dot_path), "-o", str(path)], check=True)
    trim_white_borders(path, padding=8)


def create_updated_charts() -> tuple[Path, Path]:
    rbs_chart = CHART_ASSET_DIR / "updated_rbs_chart.png"
    wbs_chart = CHART_ASSET_DIR / "updated_wbs_chart.png"
    create_updated_rbs_chart(rbs_chart)
    create_updated_wbs_chart(wbs_chart)
    return rbs_chart, wbs_chart


def p(text: str, style) -> Paragraph:
    return Paragraph(text, style)


def footer(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor(PALETTE["grid"]))
    canvas.line(doc.leftMargin, 17 * mm, A4[0] - doc.rightMargin, 17 * mm)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor(PALETTE["muted"]))
    canvas.drawString(doc.leftMargin, 11 * mm, TITLE)
    canvas.drawRightString(A4[0] - doc.rightMargin, 11 * mm, f"Page {doc.page}")
    canvas.restoreState()


def make_table(rows, col_widths, styles, header=True, font_small=False):
    cell_style = styles["TableCellSmall"] if font_small else styles["TableCell"]
    data = []
    for r_idx, row in enumerate(rows):
        style = styles["TableHead"] if header and r_idx == 0 else cell_style
        data.append([p(str(value), style) for value in row])
    table = Table(data, colWidths=col_widths, repeatRows=1 if header else 0)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(PALETTE["navy"]) if header else colors.white),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white if header else colors.HexColor(PALETTE["text"])),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(PALETTE["grid"])),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(PALETTE["grid"])),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FBFCFE")]),
            ]
        )
    )
    return table


def build_pdf() -> None:
    styles = build_styles()
    signatures = normalize_signature_assets()
    updated_rbs_chart, updated_wbs_chart = create_updated_charts()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=16 * mm, bottomMargin=22 * mm)
    story = []

    story.append(p(TITLE, styles["TitleCenter"]))
    story.append(p(SUBTITLE, styles["SubtitleCenter"]))
    story.append(HRFlowable(width="100%", color=colors.HexColor(PALETTE["grid"]), thickness=1))
    story.append(Spacer(1, 7))

    meta = [
        ["Project", "HeartCompass / Digital Immortality"],
        ["Meeting Time", MEETING_TIME],
        ["Duration", MEETING_DURATION],
        ["Presenter", PRESENTER],
        ["Prepared By", TEAM],
        ["Document Date", DOCUMENT_DATE],
        ["Submission Purpose", "Canvas submission for Milestone 1 Review Meeting Summary"],
    ]
    story.append(make_table(meta, [36 * mm, 138 * mm], styles, header=False))

    story.append(p("1. Main Content of the Meeting", styles["Section"]))
    for item in MAIN_CONTENT:
        story.append(p(item, styles["Body"]))

    story.append(PageBreak())
    story.append(p("2. Reviewer Comments and Advice", styles["Section"]))
    rows = [["Reviewer", "Focus", "Comment / Advice", "Team Response"]]
    rows += REVIEWER_COMMENTS
    story.append(make_table(rows, [25 * mm, 33 * mm, 61 * mm, 55 * mm], styles, font_small=True))

    story.append(PageBreak())
    story.append(p("3. Action Plan for Current Problems in the Next Milestone", styles["Section"]))
    story.append(p("The next milestone will focus on converting reviewer feedback into concrete management and engineering tasks. The first action has already been reflected in the updated RBS, WBS, schedule, and planned working hours below.", styles["Body"]))
    rows = [["ID", "Action", "Description", "Owner", "Target", "Status"]]
    rows += ACTION_ITEMS
    story.append(make_table(rows, [12 * mm, 34 * mm, 57 * mm, 26 * mm, 26 * mm, 19 * mm], styles, font_small=True))

    story.append(PageBreak())
    story.append(p("4. Original and Updated RBS", styles["Section"]))
    story.append(p("The original RBS described five broad responsibility streams. After the review, the RBS was updated into a numbered hierarchy that better matches the current main-branch implementation and gives traceability to the updated WBS.", styles["Body"]))
    story.append(p("Original RBS", styles["Subsection"]))
    rows = [["ID", "Original Responsibility Area", "Scope"]]
    rows += ORIGINAL_RBS
    story.append(make_table(rows, [16 * mm, 48 * mm, 110 * mm], styles, font_small=True))
    story.append(p("Updated RBS", styles["Subsection"]))
    rows = [["ID", "Updated Responsibility Area", "Scope"]]
    rows += UPDATED_RBS
    story.append(make_table(rows, [16 * mm, 50 * mm, 108 * mm], styles, font_small=True))
    story.append(p("Updated RBS Chart", styles["Subsection"]))
    story.append(fitted_reportlab_image(updated_rbs_chart, 170, 82))
    story.append(p("Figure 1. Updated Responsibility Breakdown Structure with numbered hierarchy.", styles["Small"]))

    story.append(p("5. Original WBS and Planned Working Hours", styles["Section"]))
    story.append(p("The original WBS used five major work packages with package-level planned hours. It did not yet provide sub-task-level decomposition, which was one of the most important project-management issues raised during the review.", styles["Body"]))
    rows = [["WP", "Original Work Package", "Original Scope", "Original Schedule", "Hours"]]
    rows += ORIGINAL_WBS
    rows.append(["Total", "", "", "", str(total_hours(ORIGINAL_WBS))])
    story.append(make_table(rows, [15 * mm, 42 * mm, 75 * mm, 28 * mm, 14 * mm], styles, font_small=True))

    story.append(p("6. Updated WBS, Planned Working Hours, and Total Hours", styles["Section"]))
    story.append(p("The updated WBS preserves the same total planned workload of 624 hours, but it replaces the broad original work packages with traceable task IDs, RBS references, week assignments, and planned hours for each task.", styles["Body"]))
    rows = [["WP", "Updated Work Package", "RBS", "Hours"]]
    rows += UPDATED_WP_SUMMARY
    rows.append(["Total", "", "", str(total_hours(UPDATED_WP_SUMMARY))])
    story.append(make_table(rows, [18 * mm, 108 * mm, 24 * mm, 24 * mm], styles))
    story.append(p("Detailed Updated Task-Level WBS", styles["Subsection"]))
    rows = [["Task ID", "Task Name", "RBS", "Weeks", "Hours"]]
    rows += UPDATED_WBS
    rows.append(["Total", "", "", "", str(total_hours(UPDATED_WBS))])
    story.append(make_table(rows, [18 * mm, 91 * mm, 20 * mm, 25 * mm, 20 * mm], styles, font_small=True))
    story.append(PageBreak())
    story.append(p("Updated WBS Chart", styles["Subsection"]))
    story.append(fitted_reportlab_image(updated_wbs_chart, 170, 210))
    story.append(p("Figure 2. Updated Work Breakdown Structure with task-level work packages and planned hours.", styles["Small"]))

    story.append(PageBreak())
    story.append(p("7. Updated Schedule for the Assignment 'Update RBS, WBS and Schedule'", styles["Section"]))
    story.append(p("The schedule was updated from broad phase-level planning to week-by-week task completion. The new schedule allows the team to report concrete task IDs at weekly checkpoints and directly connects progress tracking to the WBS.", styles["Body"]))
    rows = [["Week", "Dates", "Planned Completion / Milestone"]]
    rows += SCHEDULE_UPDATE
    story.append(make_table(rows, [18 * mm, 38 * mm, 118 * mm], styles))

    story.append(p("8. Summary of Changes from Original Plan to Updated Plan", styles["Section"]))
    changes = [
        "The project scope moved away from relationship-analysis APIs and a lightweight client as the primary delivery surface, and toward persona construction, ConversationGraph dialogue, CLI management, Lark Bot interaction, and Harness assets.",
        "Original entities such as Crush, RelationChain, Event, ChatTopic, DerivedInsight, and ContextEmbedding were replaced or narrowed into the current main-branch model centered on User, FigureAndRelation, OriginalSource, FineGrainedFeed, FineGrainedFeedConflict, FROverallUpdateLog, FRBuildingGraphReport, Knowledge, and Analysis.",
        "The updated RBS is numbered and traceable, while the original RBS was descriptive but not task-linked.",
        "The updated WBS keeps the same total 624 planned hours but introduces 25 task-level work items, week assignments, RBS references, and explicit review/test tasks.",
        "The updated schedule reflects current implementation realities: service-only data access, FRBuildingGraph concurrency control, CLI logs, Lark cards and commands, PyPI packaging, Docker/Alembic deployment support, and Harness documentation.",
    ]
    for item in changes:
        story.append(p(f"- {item}", styles["Body"]))

    story.append(PageBreak())
    story.append(p("9. Reviewer Signatures", styles["Section"]))
    story.append(p("The following signatures confirm that the reviewers participated in the Milestone 1 review discussion and provided feedback recorded in this meeting summary.", styles["Body"]))
    signature_cells = []
    for name, path in signatures:
        image = fitted_reportlab_image(path, 68, 30)
        signature_cells.append([p(name, styles["Subsection"]), image])
    while len(signature_cells) % 2:
        signature_cells.append([p("", styles["Body"])])
    grid_rows = []
    for i in range(0, len(signature_cells), 2):
        grid_rows.append([signature_cells[i], signature_cells[i + 1]])
    sig_table = Table(grid_rows, colWidths=[86 * mm, 86 * mm])
    sig_table.setStyle(
        TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(PALETTE["grid"])),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(PALETTE["grid"])),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(sig_table)

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def add_run(paragraph, text: str, bold: bool = False, size: float = 10.0, color: str = PALETTE["text"]):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return run


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    para.paragraph_format.space_after = Pt(4)
    add_run(para, text, bold=True, size=15 if level == 1 else 12, color=PALETTE["navy"] if level == 1 else PALETTE["text"])


def add_body(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.18
    add_run(para, text, size=10.2)


def add_docx_table(doc: Document, headers: list[str], rows: list[tuple], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, PALETTE["navy"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p0, header, bold=True, size=8.4, color="#FFFFFF")
    for r_idx, row in enumerate(rows, 1):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            if r_idx % 2 == 0:
                set_cell_shading(cells[c_idx], "FBFCFE")
            cells[c_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p0 = cells[c_idx].paragraphs[0]
            p0.paragraph_format.space_after = Pt(0)
            add_run(p0, str(value), size=8.2)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def build_docx() -> None:
    signatures = normalize_signature_assets()
    updated_rbs_chart, updated_wbs_chart = create_updated_charts()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.orientation = WD_ORIENT.PORTRAIT

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.2)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, TITLE, bold=True, size=20, color=PALETTE["navy"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, SUBTITLE, size=10, color=PALETTE["muted"])

    add_docx_table(
        doc,
        ["Field", "Value"],
        [
            ("Project", "HeartCompass / Digital Immortality"),
            ("Meeting Time", MEETING_TIME),
            ("Duration", MEETING_DURATION),
            ("Presenter", PRESENTER),
            ("Prepared By", TEAM),
            ("Document Date", DOCUMENT_DATE),
            ("Submission Purpose", "Canvas submission for Milestone 1 Review Meeting Summary"),
        ],
        [1.55, 5.6],
    )

    add_heading(doc, "1. Main Content of the Meeting")
    for item in MAIN_CONTENT:
        add_body(doc, item)

    add_heading(doc, "2. Reviewer Comments and Advice")
    add_docx_table(doc, ["Reviewer", "Focus", "Comment / Advice", "Team Response"], REVIEWER_COMMENTS, [1.05, 1.45, 2.45, 2.35])

    add_heading(doc, "3. Action Plan for Current Problems in the Next Milestone")
    add_body(doc, "The next milestone will focus on converting reviewer feedback into concrete management and engineering tasks. The first action has already been reflected in the updated RBS, WBS, schedule, and planned working hours below.")
    add_docx_table(doc, ["ID", "Action", "Description", "Owner", "Target", "Status"], ACTION_ITEMS, [0.45, 1.45, 2.45, 1.05, 1.0, 0.8])

    add_heading(doc, "4. Original and Updated RBS")
    add_body(doc, "The original RBS described five broad responsibility streams. After the review, the RBS was updated into a numbered hierarchy that better matches the current main-branch implementation and gives traceability to the updated WBS.")
    add_heading(doc, "Original RBS", level=2)
    add_docx_table(doc, ["ID", "Original Responsibility Area", "Scope"], ORIGINAL_RBS, [0.55, 2.1, 4.6])
    add_heading(doc, "Updated RBS", level=2)
    add_docx_table(doc, ["ID", "Updated Responsibility Area", "Scope"], UPDATED_RBS, [0.65, 2.15, 4.45])
    add_heading(doc, "Updated RBS Chart", level=2)
    doc.add_picture(str(updated_rbs_chart), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Figure 1. Updated Responsibility Breakdown Structure with numbered hierarchy.")

    add_heading(doc, "5. Original WBS and Planned Working Hours")
    add_body(doc, "The original WBS used five major work packages with package-level planned hours. It did not yet provide sub-task-level decomposition, which was one of the most important project-management issues raised during the review.")
    original_rows = ORIGINAL_WBS + [("Total", "", "", "", total_hours(ORIGINAL_WBS))]
    add_docx_table(doc, ["WP", "Original Work Package", "Original Scope", "Original Schedule", "Hours"], original_rows, [0.5, 1.75, 3.0, 1.15, 0.65])

    add_heading(doc, "6. Updated WBS, Planned Working Hours, and Total Hours")
    add_body(doc, "The updated WBS preserves the same total planned workload of 624 hours, but it replaces the broad original work packages with traceable task IDs, RBS references, week assignments, and planned hours for each task.")
    summary_rows = UPDATED_WP_SUMMARY + [("Total", "", "", total_hours(UPDATED_WP_SUMMARY))]
    add_docx_table(doc, ["WP", "Updated Work Package", "RBS", "Hours"], summary_rows, [0.65, 4.5, 0.75, 0.65])
    add_heading(doc, "Detailed Updated Task-Level WBS", level=2)
    updated_rows = UPDATED_WBS + [("Total", "", "", "", total_hours(UPDATED_WBS))]
    add_docx_table(doc, ["Task ID", "Task Name", "RBS", "Weeks", "Hours"], updated_rows, [0.75, 3.7, 0.75, 0.85, 0.65])
    doc.add_page_break()
    add_heading(doc, "Updated WBS Chart", level=2)
    doc.add_picture(str(updated_wbs_chart), width=Inches(6.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Figure 2. Updated Work Breakdown Structure with task-level work packages and planned hours.")

    add_heading(doc, "7. Updated Schedule for the Assignment 'Update RBS, WBS and Schedule'")
    add_body(doc, "The schedule was updated from broad phase-level planning to week-by-week task completion. The new schedule allows the team to report concrete task IDs at weekly checkpoints and directly connects progress tracking to the WBS.")
    add_docx_table(doc, ["Week", "Dates", "Planned Completion / Milestone"], SCHEDULE_UPDATE, [0.65, 1.45, 5.2])

    add_heading(doc, "8. Summary of Changes from Original Plan to Updated Plan")
    for item in [
        "The project scope moved away from relationship-analysis APIs and a lightweight client as the primary delivery surface, and toward persona construction, ConversationGraph dialogue, CLI management, Lark Bot interaction, and Harness assets.",
        "Original entities such as Crush, RelationChain, Event, ChatTopic, DerivedInsight, and ContextEmbedding were replaced or narrowed into the current main-branch model centered on User, FigureAndRelation, OriginalSource, FineGrainedFeed, FineGrainedFeedConflict, FROverallUpdateLog, FRBuildingGraphReport, Knowledge, and Analysis.",
        "The updated RBS is numbered and traceable, while the original RBS was descriptive but not task-linked.",
        "The updated WBS keeps the same total 624 planned hours but introduces 25 task-level work items, week assignments, RBS references, and explicit review/test tasks.",
        "The updated schedule reflects current implementation realities: service-only data access, FRBuildingGraph concurrency control, CLI logs, Lark cards and commands, PyPI packaging, Docker/Alembic deployment support, and Harness documentation.",
    ]:
        add_body(doc, f"- {item}")

    add_heading(doc, "9. Reviewer Signatures")
    add_body(doc, "The following signatures confirm that the reviewers participated in the Milestone 1 review discussion and provided feedback recorded in this meeting summary.")
    sig_table = doc.add_table(rows=0, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.style = "Table Grid"
    for row_start in range(0, len(signatures), 2):
        row = sig_table.add_row()
        row.height = Pt(64)
        for col in range(2):
            idx = row_start + col
            cell = row.cells[col]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx >= len(signatures):
                cell.text = ""
                continue
            name, path = signatures[idx]
            p_name = cell.paragraphs[0]
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_name.paragraph_format.space_after = Pt(2)
            add_run(p_name, name, bold=True, size=9.5, color=PALETTE["navy"])
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(0)
            run = p_img.add_run()
            with PILImage.open(path) as image:
                ratio = image.width / image.height
            if ratio >= 2.6:
                run.add_picture(str(path), width=Inches(2.6))
            else:
                run.add_picture(str(path), height=Inches(0.85))

    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    build_pdf()
    build_docx()
    print(f"Generated: {PDF_PATH}")
    print(f"Generated: {DOCX_PATH}")


if __name__ == "__main__":
    main()
