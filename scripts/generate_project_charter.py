from __future__ import annotations

import math
import subprocess
from datetime import date, datetime, timedelta
from html import escape
from pathlib import Path

from PIL import Image as PILImage
from PIL import ImageChops
from PIL import ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    Image,
    PageBreak,
    Paragraph,
    KeepTogether,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOC_DIR = ROOT / "output" / "doc"
OUTPUT_PDF_DIR = ROOT / "output" / "pdf"
TMP_DIR = ROOT / "tmp" / "docs" / "charter_assets"
ASSET_DIR = TMP_DIR / "images"
DOCX_PATH = OUTPUT_DOC_DIR / "HeartCompass+Project Charter+卜天.docx"
EDITABLE_DOCX_PATH = OUTPUT_DOC_DIR / "HeartCompass+Project Charter+卜天-editable.docx"
PDF_PATH = OUTPUT_PDF_DIR / "HeartCompass+Project Charter+卜天.pdf"
PDF_RENDER_PREFIX = TMP_DIR / "pdf_render"

PROJECT_TITLE = "HeartCompass Project Charter"
PROJECT_PERIOD = "March 30, 2026 to June 14, 2026"
TEAM_MEMBERS = ["Leishiyu Chen", "Haojie Hu", "Tian Bu", "Yihan Liu"]
TEAM_LEADER = "Tian Bu"
APPROVED_BY = "Tian Bu"
PROJECT_TOPIC = "HeartCompass (Digital Immortality) v1.2.2"
TOTAL_HOURS = 624

PALETTE = {
    "navy": "#1F4E79",
    "slate": "#456B84",
    "teal": "#5D8AA8",
    "pale": "#DDE9F2",
    "soft": "#F5F8FB",
    "grid": "#D3DCE6",
    "text": "#23313D",
    "muted": "#6B7C8C",
    "green": "#7BAF8D",
}


def ensure_dirs() -> None:
    for path in [OUTPUT_DOC_DIR, OUTPUT_PDF_DIR, TMP_DIR, ASSET_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def parse_date(s: str) -> date:
    return datetime.strptime(s, "%Y-%m-%d").date()


# ---------------------------------------------------------------------------
# Week definitions (W1-W11)
# ---------------------------------------------------------------------------
WEEKS = [
    {"id": "W1", "start": "2026-03-30", "end": "2026-04-05"},
    {"id": "W2", "start": "2026-04-06", "end": "2026-04-12"},
    {"id": "W3", "start": "2026-04-13", "end": "2026-04-19"},
    {"id": "W4", "start": "2026-04-20", "end": "2026-04-26"},
    {"id": "W5", "start": "2026-04-27", "end": "2026-05-03"},
    {"id": "W6", "start": "2026-05-04", "end": "2026-05-10"},
    {"id": "W7", "start": "2026-05-11", "end": "2026-05-17"},
    {"id": "W8", "start": "2026-05-18", "end": "2026-05-24"},
    {"id": "W9", "start": "2026-05-25", "end": "2026-05-31"},
    {"id": "W10", "start": "2026-06-01", "end": "2026-06-07"},
    {"id": "W11", "start": "2026-06-08", "end": "2026-06-14"},
]
WEEK_INDEX = {w["id"]: i for i, w in enumerate(WEEKS)}

# ---------------------------------------------------------------------------
# Work packages with numbered sub-tasks, RBS references, and week assignments
# ---------------------------------------------------------------------------
WORK_PACKAGES = [
    {
        "id": "WP1",
        "name": "User and Figure-Relation Data Foundation",
        "rbs_ref": "R1.1",
        "hours": 84,
        "color": PALETTE["teal"],
        "tasks": [
            {"id": "WP1.1", "name": "User registration and authentication", "hours": 20, "weeks": ["W3"]},
            {"id": "WP1.2", "name": "Profile management and MBTI collection", "hours": 22, "weeks": ["W3"]},
            {"id": "WP1.3", "name": "FigureAndRelation CRUD and FigureRole", "hours": 24, "weeks": ["W4"]},
            {"id": "WP1.4", "name": "Data foundation review and test", "hours": 18, "weeks": ["W4"]},
        ],
    },
    {
        "id": "WP2",
        "name": "Persona Building and Knowledge Management",
        "rbs_ref": "R1.2",
        "hours": 120,
        "color": PALETTE["teal"],
        "tasks": [
            {"id": "WP2.1", "name": "FRBuildingGraph preprocessing pipeline", "hours": 28, "weeks": ["W4", "W5"]},
            {"id": "WP2.2", "name": "Fine-grained feed extraction and storage", "hours": 26, "weeks": ["W5"]},
            {"id": "WP2.3", "name": "Conflict detection and update logging", "hours": 22, "weeks": ["W5", "W6"]},
            {"id": "WP2.4", "name": "Embedding generation and vector recall", "hours": 26, "weeks": ["W6"]},
            {"id": "WP2.5", "name": "Persona building review and test", "hours": 18, "weeks": ["W6"]},
        ],
    },
    {
        "id": "WP3",
        "name": "Conversation and Dialogue System",
        "rbs_ref": "R2.1",
        "hours": 132,
        "color": PALETTE["slate"],
        "tasks": [
            {"id": "WP3.1", "name": "Service-backed ConversationGraph recall", "hours": 28, "weeks": ["W5", "W6"]},
            {"id": "WP3.2", "name": "Short-term memory trim and summary", "hours": 30, "weeks": ["W6", "W7"]},
            {"id": "WP3.3", "name": "LLM call and response generation", "hours": 28, "weeks": ["W7"]},
            {"id": "WP3.4", "name": "Feed sync and prompt alignment", "hours": 26, "weeks": ["W7", "W8"]},
            {"id": "WP3.5", "name": "Conversation system review and test", "hours": 20, "weeks": ["W8"]},
        ],
    },
    {
        "id": "WP4",
        "name": "Channel Integration and Lark Bot",
        "rbs_ref": "R2.2",
        "hours": 126,
        "color": PALETTE["green"],
        "tasks": [
            {"id": "WP4.1", "name": "Lark WebSocket and message handling", "hours": 30, "weeks": ["W6", "W7"]},
            {"id": "WP4.2", "name": "Bot menu and persona build commands", "hours": 26, "weeks": ["W7"]},
            {"id": "WP4.3", "name": "Message batching and card templates", "hours": 24, "weeks": ["W7", "W8"]},
            {"id": "WP4.4", "name": "Service-only data access alignment", "hours": 26, "weeks": ["W8"]},
            {"id": "WP4.5", "name": "Busy-state cards and integration test", "hours": 20, "weeks": ["W8"]},
        ],
    },
    {
        "id": "WP5",
        "name": "CLI, Harness, Packaging, and Deployment",
        "rbs_ref": "R3.1",
        "hours": 162,
        "color": PALETTE["navy"],
        "tasks": [
            {"id": "WP5.1", "name": "CLI framework, doctor/setup, and logs", "hours": 32, "weeks": ["W7", "W8"]},
            {"id": "WP5.2", "name": "CLI auth, FR management, and JSON contracts", "hours": 30, "weeks": ["W8", "W9"]},
            {"id": "WP5.3", "name": "Docker support and database automation", "hours": 30, "weeks": ["W9"]},
            {"id": "WP5.4", "name": "Graph/channel/CLI regression testing", "hours": 30, "weeks": ["W9", "W10"]},
            {"id": "WP5.5", "name": "PyPI packaging and publishing", "hours": 22, "weeks": ["W10"]},
            {"id": "WP5.6", "name": "Final docs, Harness assets, and deployment", "hours": 18, "weeks": ["W11"]},
        ],
    },
]

PLANNING_PHASES = [
    {"id": "P0", "name": "Planning and Requirements", "weeks": ["W1"], "color": PALETTE["navy"]},
    {"id": "P1", "name": "System Design", "weeks": ["W2"], "color": PALETTE["navy"]},
]

WORKLOAD_TABLE = [(wp["id"], wp["name"], wp["hours"]) for wp in WORK_PACKAGES]

# Flatten all tasks for easy iteration
ALL_TASKS = []
for wp in WORK_PACKAGES:
    for task in wp["tasks"]:
        ALL_TASKS.append({**task, "wp_id": wp["id"], "wp_name": wp["name"], "rbs_ref": wp["rbs_ref"], "color": wp["color"]})

# ---------------------------------------------------------------------------
# RBS content (numbered hierarchy)
# ---------------------------------------------------------------------------
RBS_CONTENT = [
    (
        "R1.1 \u2013 User and Figure-Relation Data Foundation",
        "This responsibility stream will establish the identity and relationship baseline of the platform. "
        "It covers two sub-areas: (R1.1.1) user identity and authentication, including registration, login, "
        "and profile management with MBTI collection; and (R1.1.2) FigureAndRelation entity management, including "
        "the creation, update, and querying of figure-relation records with FigureRole assignment (SELF, FAMILY, "
        "FRIEND, MENTOR, COLLEAGUE, PARTNER, PUBLIC_FIGURE, STRANGER) and ownership validation. "
        "All downstream modules depend on the data integrity guaranteed by this stream.",
    ),
    (
        "R1.2 \u2013 Persona Building and Knowledge Management",
        "This stream will convert fragmented interpersonal materials into structured persona data. "
        "It covers two sub-areas: (R1.2.1) the FRBuildingGraph pipeline, which preprocesses raw input, extracts "
        "FineGrainedFeed entries across four dimensions (PERSONALITY, INTERACTION_STYLE, PROCEDURAL_INFO, MEMORY), "
        "and updates FigureAndRelation core fields; and (R1.2.2) conflict detection, update logging, and embedding "
        "generation, including FineGrainedFeedConflict tracking, FROverallUpdateLog recording, and pgvector-based "
        "vector storage for semantic recall. "
        "The main responsibility is to ensure that stored persona data remains retrievable, well-typed, and suitable "
        "for downstream dialogue and recall.",
    ),
    (
        "R2.1 \u2013 Conversation and Dialogue System",
        "This responsibility stream will organize the dialogue core of HeartCompass. "
        "It covers two sub-areas: (R2.1.1) the ConversationGraph pipeline, which loads persona and recalled feeds, "
        "assembles system prompts with persona markdown, and generates context-grounded replies; and (R2.1.2) "
        "short-term memory management, including round-based message trimming, rolling conversation summary, "
        "service-backed persona loading, and feed synchronization back to FigureAndRelation core fields to maintain "
        "coherence over extended dialogue sessions.",
    ),
    (
        "R2.2 \u2013 Channel Integration and Lark Bot",
        "This stream will be responsible for the channel integration mainline. "
        "It covers two sub-areas: (R2.2.1) Lark WebSocket connection and message handling, including event "
        "reception, user-to-FR mapping via open_id, and timed message batching that groups incoming messages "
        "before dispatching to ConversationGraph; and (R2.2.2) bot command menu and card-based output, including "
        "persona viewing, persona building triggers, FR switching commands, structured card templates for "
        "delivering reports and conversation replies, and busy-state feedback when FRBuildingGraph is already running.",
    ),
    (
        "R3.1 \u2013 CLI, Harness, Packaging, and Deployment",
        "This stream will provide the management interface and prepare the system for delivery. "
        "It covers two sub-areas: (R3.1.1) CLI framework development, including the immortality command tree "
        "(doctor, setup, auth, fr, logs, lark-service), JSON output mode, and local session management under "
        "~/.immortality/; and (R3.1.2) system packaging, Harness documentation, and deployment, including Docker "
        "Compose automation for PostgreSQL, database migration via Alembic, PyPI publishing as digital-immortality, "
        "regression testing, and reusable skill assets for course delivery.",
    ),
]

# ---------------------------------------------------------------------------
# WBS content (task-level descriptions referencing task IDs)
# ---------------------------------------------------------------------------
WBS_CONTENT = [
    (
        "WP1: User and Figure-Relation Data Foundation [R1.1] \u2013 84 hours",
        "This work package builds the minimum data layer required by every other module. "
        "It contains four sub-tasks: WP1.1 implements user registration and authentication with session "
        "management (20h, W3); WP1.2 adds profile management with MBTI collection and open_id binding (22h, W3); "
        "WP1.3 delivers FigureAndRelation CRUD with FigureRole assignment and ownership validation (24h, W4); "
        "and WP1.4 conducts the review and test cycle covering CRUD reliability, permission isolation, "
        "and role-based data consistency (18h, W4).",
    ),
    (
        "WP2: Persona Building and Knowledge Management [R1.2] \u2013 120 hours",
        "This work package handles the transformation of raw interpersonal materials into structured persona data. "
        "WP2.1 implements the FRBuildingGraph preprocessing pipeline, including source validation, content "
        "cleaning, and OriginalSource persistence (28h, W4\u2013W5); "
        "WP2.2 adds FineGrainedFeed extraction across four dimensions (PERSONALITY, INTERACTION_STYLE, "
        "PROCEDURAL_INFO, MEMORY) with upsert logic (26h, W5); "
        "WP2.3 builds conflict detection, FineGrainedFeedConflict tracking, and FROverallUpdateLog "
        "recording (22h, W5\u2013W6); "
        "WP2.4 delivers embedding generation via pgvector and semantic vector recall (26h, W6); "
        "and WP2.5 runs the review and test cycle for extraction accuracy and retrieval behavior (18h, W6).",
    ),
    (
        "WP3: Conversation and Dialogue System [R2.1] \u2013 132 hours",
        "This work package constructs the dialogue workflow. "
        "WP3.1 builds the service-backed ConversationGraph persona loading and feed recall pipeline, including persona "
        "markdown assembly and MEMORY/PROCEDURAL_INFO semantic retrieval (28h, W5\u2013W6); "
        "WP3.2 implements short-term memory trimming with round-based cutoff and rolling conversation "
        "summary generation (30h, W6\u2013W7); "
        "WP3.3 implements LLM call and response generation with system prompt assembly (28h, W7); "
        "WP3.4 adds feed synchronization back to FigureAndRelation core fields and prompt alignment for current "
        "main-branch behavior (26h, W7\u2013W8); "
        "and WP3.5 conducts the review and test cycle for persona consistency and dialogue coherence (20h, W8).",
    ),
    (
        "WP4: Channel Integration and Lark Bot [R2.2] \u2013 126 hours",
        "This work package implements the channel integration mainline. "
        "WP4.1 builds Lark WebSocket connection and message handling with open_id user mapping (30h, W6\u2013W7); "
        "WP4.2 adds bot menu commands for persona viewing, building, and FR switching (26h, W7); "
        "WP4.3 implements timed message batching and structured card templates for reports and replies "
        "(24h, W7\u2013W8); "
        "WP4.4 delivers service-only data access alignment between channel, graph, and service layers (26h, W8); "
        "and WP4.5 runs the review and test cycle for end-to-end message flow, command reliability, and busy-state "
        "user feedback (20h, W8).",
    ),
    (
        "WP5: CLI, Harness, Packaging, and Deployment [R3.1] \u2013 162 hours",
        "This work package completes the product and prepares it for delivery. "
        "WP5.1 develops the CLI framework with doctor, setup, and logs commands, including Docker-based PostgreSQL "
        "provisioning and database initialization (32h, W7\u2013W8); "
        "WP5.2 builds CLI auth, FR management, and JSON output contracts (30h, W8\u2013W9); "
        "WP5.3 implements Docker Compose support and database automation via Alembic migrations (30h, W9); "
        "WP5.4 conducts graph, channel, and CLI regression testing across all pipelines (30h, W9\u2013W10); "
        "WP5.5 handles PyPI packaging and publishing as digital-immortality (22h, W10); "
        "and WP5.6 delivers final documentation, Harness skill assets, and deployment (18h, W11).",
    ),
]

KEY_APIS = [
    "immortality doctor",
    "immortality setup",
    "immortality auth login/logout/whoami",
    "immortality fr list/show/create",
    "immortality logs",
    "immortality lark-service start",
    "FRBuildingGraph (persona construction pipeline)",
    "ConversationGraph (dialogue generation pipeline)",
    "Lark Bot commands (/menu, /list_available_persons, /show_persona, /build_persona)",
]

BUSINESS_ENTITIES = [
    "User",
    "FigureAndRelation",
    "OriginalSource",
    "FineGrainedFeed",
    "FineGrainedFeedConflict",
    "FROverallUpdateLog",
    "FRBuildingGraphReport",
    "Knowledge",
    "Analysis",
]

# ---------------------------------------------------------------------------
# Text constants
# ---------------------------------------------------------------------------
OBJECTIVES_TEXT = "The project will pursue four closely related objectives. First, it will collect and structure user and figure-relation data behind service-layer contracts so that graph, channel, and CLI modules share a dependable foundation. Second, it will transform raw interpersonal materials into structured FineGrainedFeed entries organized by dimension (personality, interaction style, procedural info, memory) and store them with vector embeddings for semantic recall. Third, it will build a conversation pipeline (ConversationGraph) that generates context-grounded dialogue by combining persona fields, recalled feeds, and short-term memory management. Fourth, it will deliver the system through a CLI-first architecture, Lark Bot integration, and reusable Harness documentation assets, enabling users and maintainers to manage figures, build personas, and sustain dialogue."

SUCCESS_CRITERIA_TEXT = "The project will be considered successful if it can organize fragmented interpersonal materials into stable FigureAndRelation entities with structured FineGrainedFeed data, produce persona-building reports that accurately reflect source materials, and support a conversation pipeline whose replies remain consistent with stored persona and recalled memory. Success will also require the CLI and Lark Bot to demonstrate the complete workflow from persona construction to sustained dialogue, including logs access, busy-state feedback, and service-only data access in the main user path, within the planned schedule and estimated workload."

SCOPE_ANCHORS_TEXT = f"The planned system scope will cover the public capabilities exposed through {', '.join(KEY_APIS)}. At the business level, the platform will rely on the core entities {', '.join(BUSINESS_ENTITIES)}. These interfaces and entities define the functional boundary of the course project and will be treated as the main scope anchors in later planning, review, and acceptance."

ASSUMPTIONS_AND_RISKS_TEXT = "The charter assumes that users will provide enough narrative or documentary input for the FRBuildingGraph to construct meaningful persona data, and it also assumes that dimension-based FineGrainedFeed organization and pgvector semantic recall will improve dialogue grounding quality. The main risks lie in persona extraction accuracy when source materials are ambiguous, response latency caused by vector recall and LLM inference in the same request chain, conflict accumulation when contradictory information enters the system, single-instance FRBuildingGraph throughput limits, future shared-database and multi-environment configuration complexity, and the sensitivity of interpersonal data, which makes ownership validation and permission control especially important throughout the project."

WORKLOAD_INTERPRETATION_TEXT = "The workload estimate places the greatest emphasis on the persona building and conversation pipelines while reserving substantial effort for data foundation, service alignment, channel integration, CLI delivery, and Harness documentation. This distribution is intended to keep the project academically credible as a software management exercise while still reflecting the real technical center of gravity of HeartCompass."

WBS_INTRO_TEXT = "Each work package is derived from its corresponding RBS responsibility area and is further decomposed into numbered sub-tasks. Every sub-task carries an explicit hour estimate, a week assignment, and an RBS traceability reference, so that progress can be tracked quantitatively at any point during execution. Review and test activities are embedded within each work package rather than treated as separate phases, ensuring that verification happens close to the implementation it validates."

WBS_TABLE_LEADIN_TEXT = "The following summary table shows the effort distribution across the five major work packages. A detailed task schedule with sub-task-level assignments follows."

DELIVERABLES_TEXT = "By the end of the planned schedule, the team will deliver a documented data model covering FigureAndRelation and FineGrainedFeed entities, a persona building pipeline (FRBuildingGraph) with conflict detection, update logging, and explicit single-run concurrency control, a conversation pipeline (ConversationGraph) with short-term memory management, a Lark Bot integration for daily interaction, a CLI tool (immortality) for system management and log inspection, a PyPI-publishable package (digital-immortality), and a final integration package that includes testing evidence, Harness documentation, and Docker-based deployment."

GOVERNANCE_TEXT = "Progress will be reviewed at the end of each week against the task schedule. The team will track completed task IDs (e.g., WP1.1, WP2.3) and compare them against the planned weekly milestones. Any scope adjustment will be evaluated against its likely effect on timeline, integration complexity, service contract stability, and delivery quality before it is accepted into the project plan. Reusable documentation and review skills will be treated as governance assets rather than one-off notes."

DELIVERY_CRITERIA_TEXT = "The final submission will need to satisfy several quality conditions at the same time. The charter must remain fully in English, use consistent future tense, and keep all dates aligned to 2026. Its sections, charts, workload estimate, and narrative explanations must remain mutually consistent, and the final PDF must be printable, readable, and professionally formatted without unstable spacing, clipped graphics, or broken page flow."

WBS_TRACEABILITY_TEXT = "The WBS is derived directly from the RBS. Each Level-2 RBS responsibility area (R1.1, R1.2, R2.1, R2.2, R3.1) maps one-to-one to a work package (WP1\u2013WP5). Within each work package, the numbered sub-tasks correspond to the Level-3 RBS sub-areas. For example, WP1.1 and WP1.2 trace to R1.1.1 (user identity), while WP1.3 traces to R1.1.2 (FigureAndRelation management), and WP5.6 traces to R3.1.2 (delivery plus Harness documentation). This structure ensures that completing all sub-tasks within a work package fulfills the corresponding RBS responsibility, and that no responsibility area is left without concrete, trackable deliverables."

GANTT_INTRO_TEXT = "The project schedule is organized at the sub-task level across eleven weeks. Each numbered task (WP1.1 through WP5.6) appears as an individual bar on the Gantt chart, grouped under its parent work package. This level of detail allows the team to verify specific task completion at any weekly checkpoint and to identify schedule slippage before it compounds."

GANTT_INTERPRETATION_TEXT = "The schedule follows a deliberate progression: planning and design occupy Weeks 1\u20132, data foundation is built in Weeks 3\u20134, persona building and conversation pipelines develop in Weeks 4\u20138, channel integration with Lark Bot and service alignment progresses in parallel during Weeks 6\u20138, CLI, Harness, and packaging runs in Weeks 7\u20139, and system testing through final deployment closes out Weeks 9\u201311. Each work package includes its own review and test sub-task, so quality verification is distributed throughout the schedule rather than concentrated at the end."

GANTT_MILESTONE_TEXT = "The weekly milestone table below summarizes the expected task completions for each week. At the end of any given week, the team should be able to report which task IDs have been completed, enabling quantitative progress tracking against the baseline schedule."

CRITICAL_PATH_TEXT = "The critical path identifies the longest dependency chain that determines whether the project can finish on time. For HeartCompass, the critical path runs through: Planning (W1) \u2192 System Design (W2) \u2192 WP1 Data Foundation (W3\u2013W4) \u2192 WP2 Persona Building (W4\u2013W6) \u2192 WP4 Channel Integration and Service Alignment (W6\u2013W8) \u2192 WP5 CLI, Harness, and Packaging (W7\u2013W9) \u2192 WP5.4 System Testing (W9\u2013W10) \u2192 WP5.5 PyPI Publishing (W10) \u2192 WP5.6 Documentation and Deployment (W11). The conversation and dialogue system (WP3) develops in parallel and converges during system integration."

CRITICAL_PATH_INTERP_TEXT = "The project will begin with requirements gathering and system design because neither the data model nor the downstream AI modules can be stabilized before those decisions are made. Once the FigureAndRelation model and persona building pipeline are defined, the channel integration mainline becomes the controlling branch because it depends on persona construction, dialogue generation, and Lark Bot integration in sequence. The conversation and dialogue system remains a major parallel branch; however, its outputs converge with the CLI and packaging flow during system integration rather than controlling the overall finish date. This dependency structure makes system testing (WP5.4), PyPI publishing (WP5.5), and final deployment (WP5.6) the final schedule gates."


# ---------------------------------------------------------------------------
# Weekly milestone summary data
# ---------------------------------------------------------------------------
def build_weekly_milestones() -> list[tuple[str, str, str]]:
    """Return (week_id, date_range, completing_tasks) for each week."""
    week_completions: dict[str, list[str]] = {w["id"]: [] for w in WEEKS}
    for task in ALL_TASKS:
        last_week = task["weeks"][-1]
        week_completions[last_week].append(task["id"])
    for phase in PLANNING_PHASES:
        last_week = phase["weeks"][-1]
        week_completions[last_week].append(phase["id"])

    result = []
    for w in WEEKS:
        wid = w["id"]
        start = parse_date(w["start"]).strftime("%b %d")
        end = parse_date(w["end"]).strftime("%b %d")
        date_range = f"{start} \u2013 {end}"
        tasks = ", ".join(week_completions[wid]) if week_completions[wid] else "\u2014"
        result.append((wid, date_range, tasks))
    return result


# ---------------------------------------------------------------------------
# Font and image utilities
# ---------------------------------------------------------------------------
def get_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Helvetica.ttc",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def hex_rgb(color_hex: str) -> tuple[int, int, int]:
    color_hex = color_hex.lstrip("#")
    return tuple(int(color_hex[i : i + 2], 16) for i in (0, 2, 4))


# ---------------------------------------------------------------------------
# Gantt chart — task-level, grouped by work package, weekly columns
# ---------------------------------------------------------------------------
def create_gantt_chart(path: Path) -> None:
    num_weeks = len(WEEKS)
    left_margin = 440
    right_margin = 60
    top_margin = 170
    bottom_margin = 100
    width = 1880
    chart_width = width - left_margin - right_margin

    # Build row list
    rows = []
    # Planning phases
    for phase in PLANNING_PHASES:
        week_indices = [WEEK_INDEX[w] for w in phase["weeks"]]
        rows.append({"type": "task", "id": phase["id"], "name": phase["name"],
                      "week_start": min(week_indices), "week_end": max(week_indices),
                      "color": phase["color"]})
    rows.append({"type": "spacing"})

    for wp in WORK_PACKAGES:
        rows.append({"type": "header", "label": f"{wp['id']}: {wp['name']} [{wp['rbs_ref']}]", "color": wp["color"]})
        for task in wp["tasks"]:
            week_indices = [WEEK_INDEX[w] for w in task["weeks"]]
            rows.append({"type": "task", "id": task["id"], "name": task["name"],
                          "week_start": min(week_indices), "week_end": max(week_indices),
                          "color": wp["color"]})
        rows.append({"type": "spacing"})

    # Calculate heights
    header_h = 40
    task_h = 36
    spacing_h = 14
    total_row_height = 0
    for row in rows:
        if row["type"] == "header":
            total_row_height += header_h
        elif row["type"] == "task":
            total_row_height += task_h
        else:
            total_row_height += spacing_h

    height = top_margin + total_row_height + bottom_margin
    image = PILImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = get_font(38, bold=True)
    subtitle_font = get_font(19)
    header_font = get_font(19, bold=True)
    label_font = get_font(18)
    id_font = get_font(17, bold=True)
    bar_font = get_font(15)
    week_font = get_font(16)

    # Title
    draw.text((left_margin, 36), "HeartCompass Task-Level Gantt Chart", fill=hex_rgb(PALETTE["navy"]), font=title_font)
    draw.text((left_margin, 84), "Sub-task schedule from March 30, 2026 to June 14, 2026 (Weeks 1\u201311)", fill=hex_rgb(PALETTE["muted"]), font=subtitle_font)

    # Week columns
    week_col_w = chart_width / num_weeks
    for i, w in enumerate(WEEKS):
        x0 = left_margin + int(i * week_col_w)
        x1 = left_margin + int((i + 1) * week_col_w)
        if i % 2 == 0:
            draw.rectangle([x0, top_margin - 2, x1, height - bottom_margin], fill=hex_rgb(PALETTE["soft"]))
        draw.line([x0, top_margin - 2, x0, height - bottom_margin], fill=hex_rgb(PALETTE["grid"]), width=1)
        # Week label
        start_d = parse_date(w["start"])
        label = f'{w["id"]}\n{start_d.strftime("%b %d")}'
        draw.text((x0 + 4, top_margin - 50), w["id"], fill=hex_rgb(PALETTE["navy"]), font=id_font)
        draw.text((x0 + 4, top_margin - 30), start_d.strftime("%b %d"), fill=hex_rgb(PALETTE["muted"]), font=week_font)
    # Right edge
    draw.line([left_margin + chart_width, top_margin - 2, left_margin + chart_width, height - bottom_margin], fill=hex_rgb(PALETTE["grid"]), width=1)

    # Draw rows
    y_cursor = top_margin
    for row in rows:
        if row["type"] == "spacing":
            y_cursor += spacing_h
            continue
        if row["type"] == "header":
            # Draw header background
            draw.rectangle([20, y_cursor, width - right_margin, y_cursor + header_h - 2], fill=hex_rgb(PALETTE["pale"]))
            draw.text((28, y_cursor + 8), row["label"], fill=hex_rgb(row["color"]), font=header_font)
            y_cursor += header_h
            continue
        # Task row
        # Guide line
        draw.line([left_margin, y_cursor + task_h // 2, width - right_margin, y_cursor + task_h // 2], fill=hex_rgb(PALETTE["grid"]), width=1)
        # Task label
        draw.text((36, y_cursor + 6), row["id"], fill=hex_rgb(row["color"]), font=id_font)
        draw.text((110, y_cursor + 7), row["name"], fill=hex_rgb(PALETTE["text"]), font=label_font)
        # Bar
        bar_x0 = left_margin + int(row["week_start"] * week_col_w) + 4
        bar_x1 = left_margin + int((row["week_end"] + 1) * week_col_w) - 4
        bar_y0 = y_cursor + 5
        bar_y1 = y_cursor + task_h - 7
        draw.rounded_rectangle([bar_x0, bar_y0, bar_x1, bar_y1], radius=8, fill=hex_rgb(row["color"]), outline=hex_rgb(PALETTE["navy"]))
        # Bar text
        span_text = WEEKS[row["week_start"]]["id"]
        if row["week_end"] != row["week_start"]:
            span_text += f'\u2013{WEEKS[row["week_end"]]["id"]}'
        draw.text((bar_x0 + 8, bar_y0 + 4), span_text, fill=(255, 255, 255), font=bar_font)
        y_cursor += task_h

    # Legend
    legend_y = height - 70
    legends = [
        ("Planning / CLI / Deploy", PALETTE["navy"]),
        ("Foundation / Persona", PALETTE["teal"]),
        ("Conversation System", PALETTE["slate"]),
        ("Channel / Lark Bot", PALETTE["green"]),
    ]
    cursor_x = left_margin
    legend_font = get_font(17)
    for label, color in legends:
        draw.rounded_rectangle([cursor_x, legend_y, cursor_x + 26, legend_y + 20], radius=5, fill=hex_rgb(color))
        draw.text((cursor_x + 34, legend_y - 1), label, fill=hex_rgb(PALETTE["text"]), font=legend_font)
        cursor_x += 240

    image.save(path)


def trim_white_borders(path: Path, padding: int = 12) -> None:
    image = PILImage.open(path).convert("RGB")
    background = PILImage.new("RGB", image.size, "white")
    diff = ImageChops.difference(image, background)
    bbox = diff.getbbox()
    if not bbox:
        return
    left = max(0, bbox[0] - padding)
    top = max(0, bbox[1] - padding)
    right = min(image.width, bbox[2] + padding)
    bottom = min(image.height, bbox[3] + padding)
    image.crop((left, top, right, bottom)).save(path)


# ---------------------------------------------------------------------------
# RBS chart — numbered hierarchy
# ---------------------------------------------------------------------------
def create_rbs_chart(path: Path) -> None:
    dot_path = path.with_suffix(".dot")
    dot_path.write_text(
        f"""
digraph G {{
  graph [rankdir=TB, splines=polyline, nodesep=0.50, ranksep=0.85, bgcolor="white", pad=0.28];
  node [shape=box, style="rounded,filled", fontname="Helvetica", fontsize=14, color="{PALETTE["navy"]}", fillcolor="white", margin="0.20,0.12"];
  edge [color="{PALETTE["slate"]}", penwidth=1.25];

  Root [label="HeartCompass\\nRBS", fillcolor="{PALETTE["pale"]}", fontsize=16];

  R1 [label="R1: Foundation and\\nPersona Input", fillcolor="{PALETTE["soft"]}"];
  R2 [label="R2: Dialogue and\\nChannel", fillcolor="{PALETTE["soft"]}"];
  R3 [label="R3: CLI, Harness,\\nand Delivery", fillcolor="{PALETTE["soft"]}"];

  R11 [label="R1.1: User and Figure-Relation\\nData Foundation", fillcolor="{PALETTE["soft"]}"];
  R12 [label="R1.2: Persona Building\\nand Knowledge Mgmt", fillcolor="{PALETTE["soft"]}"];
  R21 [label="R2.1: Conversation and\\nDialogue System", fillcolor="{PALETTE["soft"]}"];
  R22 [label="R2.2: Channel Integration\\nand Lark Bot", fillcolor="{PALETTE["soft"]}"];
  R31 [label="R3.1: CLI, Harness,\\nPackaging, and Delivery", fillcolor="{PALETTE["soft"]}"];

  R111 [label="R1.1.1: User identity\\nand authentication"];
  R112 [label="R1.1.2: FigureAndRelation\\nand FigureRole mgmt"];
  R121 [label="R1.2.1: FRBuildingGraph\\nand feed extraction"];
  R122 [label="R1.2.2: Conflict detection\\nand embedding storage"];
  R211 [label="R2.1.1: ConversationGraph\\nand feed recall"];
  R212 [label="R2.1.2: Memory trim\\nand feed sync"];
  R221 [label="R2.2.1: Lark WebSocket\\nand message handling"];
  R222 [label="R2.2.2: Bot commands\\nand card templates"];
  R311 [label="R3.1.1: CLI framework,\\nlogs, and commands"];
  R312 [label="R3.1.2: Docker, PyPI,\\nHarness, and deployment"];

  Root -> R1; Root -> R2; Root -> R3;
  R1 -> R11; R1 -> R12;
  R2 -> R21; R2 -> R22;
  R3 -> R31;

  R11 -> R111; R11 -> R112;
  R12 -> R121; R12 -> R122;
  R21 -> R211; R21 -> R212;
  R22 -> R221; R22 -> R222;
  R31 -> R311; R31 -> R312;

  {{ rank=same; R1; R2; R3; }}
  {{ rank=same; R11; R12; R21; R22; R31; }}
}}
""".strip()
    )
    subprocess.run(
        ["dot", "-Tpng", "-Gdpi=220", str(dot_path), "-o", str(path)],
        check=True,
    )
    trim_white_borders(path, padding=8)


# ---------------------------------------------------------------------------
# WBS chart — task IDs with RBS references
# ---------------------------------------------------------------------------
def create_wbs_chart(path: Path) -> None:
    dot_path = path.with_suffix(".dot")

    cluster_blocks = []
    for idx, wp in enumerate(WORK_PACKAGES):
        wp_node = wp["id"]
        task_lines = []
        for task in wp["tasks"]:
            tid = task["id"].replace(".", "_")
            label = f'{task["id"]}\\n{task["name"][:30]}\\n{task["hours"]}h'
            task_lines.append(f'    {tid} [label="{label}"];')
            task_lines.append(f"    {wp_node} -> {tid};")
        cluster_blocks.append(
            f'  subgraph cluster_{wp_node} {{\n'
            f'    style=rounded; color="{PALETTE["grid"]}"; bgcolor="{PALETTE["soft"]}";\n'
            f'    label="";\n'
            f'    {wp_node} [label="{wp_node}: {wp["name"][:26]}\\n[{wp["rbs_ref"]}] {wp["hours"]}h", fillcolor="{PALETTE["pale"]}", fontsize=12];\n'
            + "\n".join(task_lines) + "\n"
            f"  }}"
        )

    dot_path.write_text(
        f"""
digraph G {{
  graph [rankdir=LR, splines=polyline, nodesep=0.25, ranksep=0.70, bgcolor="white", pad=0.25];
  node [shape=box, style="rounded,filled", fontname="Helvetica", fontsize=10, color="{PALETTE["navy"]}", fillcolor="white", margin="0.10,0.07"];
  edge [color="{PALETTE["slate"]}", penwidth=1.0];

  Root [label="HeartCompass\\nWBS", fillcolor="{PALETTE["pale"]}", fontsize=14];

{chr(10).join(cluster_blocks)}

  Root -> WP1;
  Root -> WP2;
  Root -> WP3;
  Root -> WP4;
  Root -> WP5;
}}
""".strip()
    )
    subprocess.run(
        ["dot", "-Tpng", "-Gdpi=200", str(dot_path), "-o", str(path)],
        check=True,
    )
    trim_white_borders(path, padding=8)


# ---------------------------------------------------------------------------
# Critical path diagram — with task references
# ---------------------------------------------------------------------------
def create_critical_path_diagram(path: Path) -> None:
    dot_path = path.with_suffix(".dot")
    dot_path.write_text(
        f"""
digraph G {{
  graph [rankdir=TB, splines=ortho, nodesep=0.65, ranksep=0.85, bgcolor="white", pad=0.40];
  node [shape=box, style="rounded,filled", fontname="Helvetica", fontsize=14, color="{PALETTE["navy"]}", fillcolor="{PALETTE["soft"]}", margin="0.20,0.14"];
  edge [fontname="Helvetica", fontsize=11, color="{PALETTE["muted"]}", arrowsize=0.8];

  Req [label="P0: Planning\\nW1 (Mar 30 - Apr 5)", fillcolor="{PALETTE["pale"]}"];
  Design [label="P1: System Design\\nW2 (Apr 6 - Apr 12)", fillcolor="{PALETTE["pale"]}"];
  WP1 [label="WP1: Data Foundation\\nW3-W4 (Apr 13 - Apr 26)", fillcolor="{PALETTE["pale"]}"];
  WP2 [label="WP2: Persona Building\\nW4-W6 (Apr 20 - May 10)", fillcolor="{PALETTE["pale"]}"];
  WP3 [label="WP3: Conversation System\\nW5-W8 (Apr 27 - May 24)"];
  WP4 [label="WP4: Channel, Lark Bot,\\nand Service Alignment\\nW6-W8 (May 4 - May 24)", fillcolor="{PALETTE["pale"]}"];
  WP5a [label="WP5.1-5.2: CLI and Logs\\nW7-W9 (May 11 - May 31)"];
  WP5b [label="WP5.3: Docker & DB\\nW9 (May 25 - May 31)", fillcolor="{PALETTE["pale"]}"];
  WP5c [label="WP5.4: System Testing\\nW9-W10 (May 25 - Jun 7)", fillcolor="{PALETTE["pale"]}"];
  WP5d [label="WP5.5: PyPI Publishing\\nW10 (Jun 1 - Jun 7)", fillcolor="{PALETTE["pale"]}"];
  WP5e [label="WP5.6: Docs, Harness,\\nand Deployment\\nW11 (Jun 8 - Jun 14)", fillcolor="{PALETTE["pale"]}"];

  Req -> Design [color="{PALETTE["navy"]}", penwidth=2.5];
  Design -> WP1 [color="{PALETTE["navy"]}", penwidth=2.5];
  WP1 -> WP2 [color="{PALETTE["navy"]}", penwidth=2.5];
  WP2 -> WP4 [color="{PALETTE["navy"]}", penwidth=2.5];
  WP4 -> WP5b [color="{PALETTE["navy"]}", penwidth=2.5];
  WP5b -> WP5c [color="{PALETTE["navy"]}", penwidth=2.5];
  WP5c -> WP5d [color="{PALETTE["navy"]}", penwidth=2.5];
  WP5d -> WP5e [color="{PALETTE["navy"]}", penwidth=2.5];

  WP2 -> WP3 [style=dashed, color="{PALETTE["muted"]}"];
  WP3 -> WP5b [style=dashed, color="{PALETTE["muted"]}"];
  WP1 -> WP5a [style=dashed, color="{PALETTE["muted"]}"];
  WP5a -> WP5b [style=dashed, color="{PALETTE["muted"]}"];

  {{ rank=same; WP3; WP4; WP5a; }}

  Legend [shape=note, label="Solid blue edges: critical path\\nDashed gray edges: parallel support flow", color="{PALETTE["grid"]}", fillcolor="white", fontsize=12];
  WP5e -> Legend [style=invis];
}}
""".strip()
    )
    subprocess.run(
        ["dot", "-Tpng", "-Gdpi=240", str(dot_path), "-o", str(path)],
        check=True,
    )
    trim_white_borders(path, padding=10)


# ---------------------------------------------------------------------------
# ReportLab PDF styles
# ---------------------------------------------------------------------------
def build_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CharterTitle",
            fontName="Helvetica-Bold",
            fontSize=22,
            textColor=colors.HexColor(PALETTE["navy"]),
            leading=28,
            spaceAfter=10,
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CharterSubtitle",
            fontName="Helvetica",
            fontSize=10.5,
            textColor=colors.HexColor(PALETTE["muted"]),
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionTitle",
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor(PALETTE["navy"]),
            spaceBefore=8,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Subhead",
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=15,
            textColor=colors.HexColor(PALETTE["text"]),
            spaceBefore=6,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor(PALETTE["text"]),
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14.5,
            leftIndent=14,
            firstLineIndent=-8,
            textColor=colors.HexColor(PALETTE["text"]),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Caption",
            fontName="Helvetica-Oblique",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor(PALETTE["muted"]),
            alignment=TA_CENTER,
            spaceBefore=4,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor(PALETTE["text"]),
            spaceAfter=3,
        )
    )
    return styles


def bullet_paragraph(text: str, style) -> Paragraph:
    return Paragraph(f"&bull; {escape(text)}", style)


def reportlab_image(path: Path, max_width_mm: float, max_height_mm: float) -> Image:
    with PILImage.open(path) as img:
        width_px, height_px = img.size
    ratio = width_px / height_px
    max_width = max_width_mm * mm
    max_height = max_height_mm * mm
    width = max_width
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    image = Image(str(path), width=width, height=height)
    image.hAlign = "CENTER"
    return image


def footer(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor(PALETTE["grid"]))
    canvas.line(doc.leftMargin, 18 * mm, A4[0] - doc.rightMargin, 18 * mm)
    canvas.setFont("Helvetica", 8.5)
    canvas.setFillColor(colors.HexColor(PALETTE["muted"]))
    canvas.drawString(doc.leftMargin, 12 * mm, PROJECT_TITLE)
    canvas.drawRightString(A4[0] - doc.rightMargin, 12 * mm, f"Page {doc.page}")
    canvas.restoreState()


# ---------------------------------------------------------------------------
# Build PDF
# ---------------------------------------------------------------------------
def build_pdf(rbs_chart_path: Path, wbs_chart_path: Path, gantt_path: Path, critical_path: Path) -> None:
    styles = build_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=22 * mm,
    )

    story = []

    # ---- Title and metadata ----
    story.append(Paragraph(PROJECT_TITLE, styles["CharterTitle"]))
    story.append(Paragraph("Course Project Charter for the HeartCompass (Digital Immortality) v1.2.2 platform", styles["CharterSubtitle"]))
    story.append(HRFlowable(color=colors.HexColor(PALETTE["grid"]), thickness=1.1, width="100%"))
    story.append(Spacer(1, 6))

    info_rows = [
        ["Project Name", PROJECT_TOPIC],
        ["Project Period", PROJECT_PERIOD],
        ["Project Type", "AI software course project"],
        ["Team Leader", TEAM_LEADER],
        ["Prepared By", ", ".join(TEAM_MEMBERS)],
        ["Approved By", APPROVED_BY],
        ["Estimated Workload", f"{TOTAL_HOURS} hours"],
        ["Document Date", "May 9, 2026"],
    ]
    info_table = Table(info_rows, colWidths=[36 * mm, 124 * mm])
    info_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor(PALETTE["text"])),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(PALETTE["soft"])),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor(PALETTE["grid"])),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(PALETTE["grid"])),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(info_table)
    story.append(Spacer(1, 10))

    # ---- Project Overview ----
    story.append(Paragraph("Project Overview", styles["SectionTitle"]))
    story.append(
        Paragraph(
            "HeartCompass (Digital Immortality) is planned as a persona-centric AI platform that will transform fragmented interpersonal materials into structured figure profiles and enable context-grounded dialogue. The current main branch has evolved into a v1.2.2 CLI-and-Lark delivery shape with service-backed graph workflows, explicit FRBuildingGraph concurrency control, and Harness documentation assets. The course project will deliver a persona building pipeline (FRBuildingGraph), a conversation pipeline (ConversationGraph), a CLI management tool, and a Lark Bot integration within a single integrated schedule.",
            styles["Body"],
        )
    )
    story.append(Paragraph("Primary Goal", styles["Subhead"]))
    story.append(
        Paragraph(
            "The project will build a platform that collects user and figure-relation data through service-layer contracts, organizes persona information into dimension-based FineGrainedFeed entries with vector embeddings, and uses LangGraph-based pipelines to support both persona construction and sustained dialogue.",
            styles["Body"],
        )
    )
    story.append(Paragraph("Objectives", styles["Subhead"]))
    story.append(Paragraph(OBJECTIVES_TEXT, styles["Body"]))
    story.append(Paragraph("Success Criteria", styles["Subhead"]))
    story.append(Paragraph(SUCCESS_CRITERIA_TEXT, styles["Body"]))
    story.append(Paragraph("Scope Anchors", styles["Subhead"]))
    story.append(Paragraph(SCOPE_ANCHORS_TEXT, styles["Body"]))
    story.append(Paragraph("Assumptions and Risks", styles["Subhead"]))
    story.append(Paragraph(ASSUMPTIONS_AND_RISKS_TEXT, styles["Body"]))
    story.append(
        Paragraph(
            "The document is intentionally written as a pre-execution project plan, even though the current repository already contains working implementations. All schedule language below therefore uses forward-looking, charter-style wording.",
            styles["Body"],
        )
    )

    # ---- RBS ----
    story.append(Spacer(1, 8))
    story.append(Paragraph("RBS", styles["SectionTitle"]))
    story.append(
        Paragraph(
            "The Responsibility Breakdown Structure (RBS) defines the major responsibility areas of the HeartCompass (Digital Immortality) course project using a numbered hierarchy. The top level identifies three responsibility groups: R1 (Foundation and Persona Input), R2 (Dialogue and Channel), and R3 (CLI, Harness, and Delivery). Each group is decomposed into responsibility streams (e.g., R1.1, R1.2), which are further divided into specific sub-areas (e.g., R1.1.1, R1.1.2). This numbering ensures clear traceability between responsibility definitions and the work packages in the WBS.",
            styles["Body"],
        )
    )
    for title, paragraph in RBS_CONTENT:
        story.append(Paragraph(title, styles["Subhead"]))
        story.append(Paragraph(paragraph, styles["Body"]))
        story.append(Spacer(1, 2))
    story.append(Spacer(1, 4))
    story.append(reportlab_image(rbs_chart_path, max_width_mm=170, max_height_mm=82))
    story.append(Paragraph("Figure 1. Responsibility Breakdown Structure with numbered hierarchy (R1-R3).", styles["Caption"]))

    # ---- WBS ----
    story.append(Spacer(1, 8))
    story.append(Paragraph("WBS", styles["SectionTitle"]))
    story.append(
        Paragraph(
            f"The Work Breakdown Structure (WBS) translates the RBS responsibility areas into five numbered work packages (WP1\u2013WP5), each containing individually tracked sub-tasks. The project estimates a total workload of {TOTAL_HOURS} hours distributed as follows.",
            styles["Body"],
        )
    )
    story.append(Paragraph(WORKLOAD_INTERPRETATION_TEXT, styles["Body"]))

    # Summary workload table
    story.append(Paragraph(WBS_TABLE_LEADIN_TEXT, styles["Body"]))
    workload_rows = [["Work Package", "RBS Ref", "Hours"]]
    for wp_id, wp_name, wp_hours in WORKLOAD_TABLE:
        rbs = next(wp["rbs_ref"] for wp in WORK_PACKAGES if wp["id"] == wp_id)
        workload_rows.append([f"{wp_id}: {wp_name}", rbs, str(wp_hours)])
    workload_rows.append(["Total", "", str(TOTAL_HOURS)])
    workload_table = Table(workload_rows, colWidths=[100 * mm, 20 * mm, 22 * mm])
    workload_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(PALETTE["navy"])),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("BACKGROUND", (0, 1), (-1, -2), colors.HexColor("#FBFCFE")),
                ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor(PALETTE["pale"])),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor(PALETTE["grid"])),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(PALETTE["grid"])),
                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(KeepTogether([workload_table, Spacer(1, 8)]))

    # RBS-to-WBS traceability
    story.append(Paragraph("RBS-to-WBS Traceability", styles["Subhead"]))
    story.append(Paragraph(WBS_TRACEABILITY_TEXT, styles["Body"]))
    story.append(Paragraph(WBS_INTRO_TEXT, styles["Body"]))

    # Detailed task table
    story.append(Paragraph("Detailed Task Schedule", styles["Subhead"]))
    detail_rows = [["Task ID", "Task Name", "RBS", "Hours", "Weeks"]]
    for task in ALL_TASKS:
        weeks_str = "\u2013".join([task["weeks"][0], task["weeks"][-1]]) if len(task["weeks"]) > 1 else task["weeks"][0]
        detail_rows.append([task["id"], task["name"], task["rbs_ref"], str(task["hours"]), weeks_str])
    detail_rows.append(["", "", "", str(TOTAL_HOURS), ""])
    detail_table = Table(detail_rows, colWidths=[18 * mm, 74 * mm, 14 * mm, 14 * mm, 18 * mm])
    detail_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(PALETTE["navy"])),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor(PALETTE["pale"])),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(PALETTE["grid"])),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor(PALETTE["grid"])),
                ("ALIGN", (2, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
                ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#FBFCFE")]),
            ]
        )
    )
    story.append(detail_table)
    story.append(Spacer(1, 8))

    # WP descriptions
    for title, paragraph in WBS_CONTENT:
        story.append(Paragraph(title, styles["Subhead"]))
        story.append(Paragraph(paragraph, styles["Body"]))
        story.append(Spacer(1, 2))

    story.append(PageBreak())
    story.append(reportlab_image(wbs_chart_path, max_width_mm=170, max_height_mm=220))
    story.append(Paragraph("Figure 2. Work Breakdown Structure with numbered sub-tasks and RBS references.", styles["Caption"]))

    # ---- Gantt Chart ----
    story.append(Spacer(1, 6))
    story.append(Paragraph("Gantt Chart", styles["SectionTitle"]))
    story.append(Paragraph(GANTT_INTRO_TEXT, styles["Body"]))
    story.append(reportlab_image(gantt_path, max_width_mm=174, max_height_mm=140))
    story.append(Paragraph("Figure 3. Task-level Gantt chart for HeartCompass (Weeks 1\u201311).", styles["Caption"]))
    story.append(Paragraph(GANTT_INTERPRETATION_TEXT, styles["Body"]))

    # Weekly milestone table
    story.append(Paragraph("Weekly Milestone Summary", styles["Subhead"]))
    story.append(Paragraph(GANTT_MILESTONE_TEXT, styles["Body"]))
    milestones = build_weekly_milestones()
    ms_rows = [["Week", "Dates", "Tasks Completing"]]
    for wid, date_range, tasks in milestones:
        ms_rows.append([wid, date_range, tasks])
    ms_table = Table(ms_rows, colWidths=[16 * mm, 34 * mm, 96 * mm])
    ms_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(PALETTE["navy"])),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(PALETTE["grid"])),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor(PALETTE["grid"])),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FBFCFE")]),
            ]
        )
    )
    story.append(ms_table)

    # ---- Critical Path ----
    story.append(Spacer(1, 10))
    story.append(Paragraph("Critical Path Diagram", styles["SectionTitle"]))
    story.append(Paragraph(CRITICAL_PATH_TEXT, styles["Body"]))
    story.append(reportlab_image(critical_path, max_width_mm=170, max_height_mm=115))
    story.append(Paragraph("Figure 4. Critical path diagram with work package and task references.", styles["Caption"]))
    story.append(Paragraph("Critical Path Interpretation", styles["Subhead"]))
    story.append(Paragraph(CRITICAL_PATH_INTERP_TEXT, styles["Body"]))
    story.append(Paragraph("Planned Deliverables", styles["Subhead"]))
    story.append(Paragraph(DELIVERABLES_TEXT, styles["Body"]))
    story.append(Paragraph("Project Governance Notes", styles["Subhead"]))
    story.append(Paragraph(GOVERNANCE_TEXT, styles["Body"]))
    story.append(Paragraph("Testing and Delivery Criteria", styles["Subhead"]))
    story.append(Paragraph(DELIVERY_CRITERIA_TEXT, styles["Body"]))

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_run_text(paragraph, text: str, bold: bool = False, size: float | None = None, color: str | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    run.font.name = "Times New Roman"
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    return run


def configure_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run_text(p, PROJECT_TITLE, size=8.5, color=PALETTE["muted"])
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(footer_para, "HeartCompass Project Charter", size=8.5, color=PALETTE["muted"])


def add_docx_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    add_run_text(p, title, bold=True, size=15, color=PALETTE["navy"])


def add_docx_subhead(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    add_run_text(p, text, bold=True, size=11.5, color=PALETTE["text"])


def add_docx_body(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = add_run_text(p, text, size=10.5, color=PALETTE["text"])
    if italic:
        run.italic = True


def add_docx_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.15
        add_run_text(p, item, size=10.5, color=PALETTE["text"])


def add_docx_picture_with_caption(doc: Document, image_path: Path, width_inches: float, caption: str) -> None:
    doc.add_picture(str(image_path), width=Inches(width_inches))
    picture_paragraph = doc.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(4)
    picture_paragraph.paragraph_format.space_after = Pt(2)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(6)
    add_run_text(caption_paragraph, caption, size=9, color=PALETTE["muted"])


# ---------------------------------------------------------------------------
# Build editable DOCX
# ---------------------------------------------------------------------------
def build_docx(rbs_chart_path: Path, wbs_chart_path: Path, gantt_path: Path, critical_path: Path) -> None:
    doc = Document()
    configure_docx_styles(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    add_run_text(title, PROJECT_TITLE, bold=True, size=20, color=PALETTE["navy"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    add_run_text(subtitle, "Course Project Charter for the HeartCompass (Digital Immortality) v1.2.2 platform", size=10, color=PALETTE["muted"])

    # Metadata table
    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    meta_rows = [
        ("Project Name", PROJECT_TOPIC),
        ("Project Period", PROJECT_PERIOD),
        ("Project Type", "AI software course project"),
        ("Team Leader", TEAM_LEADER),
        ("Prepared By", ", ".join(TEAM_MEMBERS)),
        ("Approved By", APPROVED_BY),
        ("Estimated Workload", f"{TOTAL_HOURS} hours"),
        ("Document Date", "May 9, 2026"),
    ]
    for r_idx, row in enumerate(meta_rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            if c_idx == 0:
                set_cell_shading(cell, PALETTE["soft"].replace("#", ""))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run_text(p, value, bold=c_idx == 0, size=10.3, color=PALETTE["text"])

    # Project Overview
    add_docx_section(doc, "Project Overview")
    add_docx_body(
        doc,
        "HeartCompass (Digital Immortality) is planned as a persona-centric AI platform that will transform fragmented interpersonal materials into structured figure profiles and enable context-grounded dialogue. The current main branch has evolved into a v1.2.2 CLI-and-Lark delivery shape with service-backed graph workflows, explicit FRBuildingGraph concurrency control, and Harness documentation assets. The course project will deliver a persona building pipeline (FRBuildingGraph), a conversation pipeline (ConversationGraph), a CLI management tool, and a Lark Bot integration within a single integrated schedule.",
    )
    add_docx_subhead(doc, "Primary Goal")
    add_docx_body(
        doc,
        "The project will build a platform that collects user and figure-relation data through service-layer contracts, organizes persona information into dimension-based FineGrainedFeed entries with vector embeddings, and uses LangGraph-based pipelines to support both persona construction and sustained dialogue.",
    )
    add_docx_subhead(doc, "Objectives")
    add_docx_body(doc, OBJECTIVES_TEXT)
    add_docx_subhead(doc, "Success Criteria")
    add_docx_body(doc, SUCCESS_CRITERIA_TEXT)
    add_docx_subhead(doc, "Scope Anchors")
    add_docx_body(doc, SCOPE_ANCHORS_TEXT)
    add_docx_subhead(doc, "Assumptions and Risks")
    add_docx_body(doc, ASSUMPTIONS_AND_RISKS_TEXT)

    # RBS
    add_docx_section(doc, "RBS")
    add_docx_body(
        doc,
        "The Responsibility Breakdown Structure (RBS) defines the major responsibility areas of the HeartCompass (Digital Immortality) course project using a numbered hierarchy. The top level identifies three responsibility groups: R1 (Foundation and Persona Input), R2 (Dialogue and Channel), and R3 (CLI, Harness, and Delivery). Each group is decomposed into responsibility streams (e.g., R1.1, R1.2), which are further divided into specific sub-areas (e.g., R1.1.1, R1.1.2). This numbering ensures clear traceability between responsibility definitions and the work packages in the WBS.",
    )
    for title_text, paragraph in RBS_CONTENT:
        add_docx_subhead(doc, title_text)
        add_docx_body(doc, paragraph)
    add_docx_picture_with_caption(
        doc, rbs_chart_path, 6.3,
        "Figure 1. Responsibility Breakdown Structure with numbered hierarchy (R1-R3).",
    )

    # WBS
    add_docx_section(doc, "WBS")
    add_docx_body(
        doc,
        f"The Work Breakdown Structure (WBS) translates the RBS responsibility areas into five numbered work packages (WP1\u2013WP5), each containing individually tracked sub-tasks. The project estimates a total workload of {TOTAL_HOURS} hours distributed as follows.",
    )
    add_docx_body(doc, WORKLOAD_INTERPRETATION_TEXT)
    add_docx_body(doc, WBS_TABLE_LEADIN_TEXT)

    # Summary workload table (DOCX)
    wl_table = doc.add_table(rows=len(WORKLOAD_TABLE) + 2, cols=3)
    wl_table.style = "Table Grid"
    wl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    wl_headers = ["Work Package", "RBS Ref", "Hours"]
    for idx, hdr in enumerate(wl_headers):
        cell = wl_table.cell(0, idx)
        cell.text = ""
        set_cell_shading(cell, PALETTE["navy"].replace("#", ""))
        p = cell.paragraphs[0]
        add_run_text(p, hdr, bold=True, size=10, color="#FFFFFF")
    for row_idx, (wp_id, wp_name, wp_hours) in enumerate(WORKLOAD_TABLE, start=1):
        rbs = next(wp["rbs_ref"] for wp in WORK_PACKAGES if wp["id"] == wp_id)
        for col_idx, value in enumerate((f"{wp_id}: {wp_name}", rbs, str(wp_hours))):
            cell = wl_table.cell(row_idx, col_idx)
            cell.text = ""
            if row_idx % 2 == 1:
                set_cell_shading(cell, "FBFCFE")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx >= 1 else WD_ALIGN_PARAGRAPH.LEFT
            add_run_text(p, value, size=10, color=PALETTE["text"])
    total_row_idx = len(WORKLOAD_TABLE) + 1
    for col_idx, value in enumerate(("Total", "", str(TOTAL_HOURS))):
        cell = wl_table.cell(total_row_idx, col_idx)
        cell.text = ""
        set_cell_shading(cell, PALETTE["pale"].replace("#", ""))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx >= 1 else WD_ALIGN_PARAGRAPH.LEFT
        add_run_text(p, value, bold=True, size=10, color=PALETTE["text"])

    add_docx_subhead(doc, "RBS-to-WBS Traceability")
    add_docx_body(doc, WBS_TRACEABILITY_TEXT)
    add_docx_body(doc, WBS_INTRO_TEXT)

    # Detailed task table (DOCX)
    add_docx_subhead(doc, "Detailed Task Schedule")
    dt_table = doc.add_table(rows=len(ALL_TASKS) + 2, cols=5)
    dt_table.style = "Table Grid"
    dt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dt_headers = ["Task ID", "Task Name", "RBS", "Hours", "Weeks"]
    for idx, hdr in enumerate(dt_headers):
        cell = dt_table.cell(0, idx)
        cell.text = ""
        set_cell_shading(cell, PALETTE["navy"].replace("#", ""))
        p = cell.paragraphs[0]
        add_run_text(p, hdr, bold=True, size=9, color="#FFFFFF")
    for row_idx, task in enumerate(ALL_TASKS, start=1):
        weeks_str = "\u2013".join([task["weeks"][0], task["weeks"][-1]]) if len(task["weeks"]) > 1 else task["weeks"][0]
        values = [task["id"], task["name"], task["rbs_ref"], str(task["hours"]), weeks_str]
        for col_idx, value in enumerate(values):
            cell = dt_table.cell(row_idx, col_idx)
            cell.text = ""
            if row_idx % 2 == 0:
                set_cell_shading(cell, "FBFCFE")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            add_run_text(p, value, size=9, color=PALETTE["text"])
    total_dt_row = len(ALL_TASKS) + 1
    for col_idx, value in enumerate(("", "", "", str(TOTAL_HOURS), "")):
        cell = dt_table.cell(total_dt_row, col_idx)
        cell.text = ""
        set_cell_shading(cell, PALETTE["pale"].replace("#", ""))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        add_run_text(p, value, bold=True, size=9, color=PALETTE["text"])

    # WP descriptions
    for title_text, paragraph in WBS_CONTENT:
        add_docx_subhead(doc, title_text)
        add_docx_body(doc, paragraph)

    doc.add_paragraph()
    add_docx_picture_with_caption(
        doc, wbs_chart_path, 6.7,
        "Figure 2. Work Breakdown Structure with numbered sub-tasks and RBS references.",
    )

    # Gantt Chart
    add_docx_section(doc, "Gantt Chart")
    add_docx_body(doc, GANTT_INTRO_TEXT)
    add_docx_picture_with_caption(
        doc, gantt_path, 6.7,
        "Figure 3. Task-level Gantt chart for HeartCompass (Weeks 1\u201311).",
    )
    add_docx_body(doc, GANTT_INTERPRETATION_TEXT)

    # Weekly milestone table (DOCX)
    add_docx_subhead(doc, "Weekly Milestone Summary")
    add_docx_body(doc, GANTT_MILESTONE_TEXT)
    milestones = build_weekly_milestones()
    ms_tbl = doc.add_table(rows=len(milestones) + 1, cols=3)
    ms_tbl.style = "Table Grid"
    ms_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, hdr in enumerate(["Week", "Dates", "Tasks Completing"]):
        cell = ms_tbl.cell(0, idx)
        cell.text = ""
        set_cell_shading(cell, PALETTE["navy"].replace("#", ""))
        p = cell.paragraphs[0]
        add_run_text(p, hdr, bold=True, size=9.5, color="#FFFFFF")
    for row_idx, (wid, date_range, tasks) in enumerate(milestones, start=1):
        for col_idx, value in enumerate((wid, date_range, tasks)):
            cell = ms_tbl.cell(row_idx, col_idx)
            cell.text = ""
            if row_idx % 2 == 0:
                set_cell_shading(cell, "FBFCFE")
            p = cell.paragraphs[0]
            add_run_text(p, value, size=9.5, color=PALETTE["text"])

    # Critical Path
    doc.add_paragraph()
    add_docx_section(doc, "Critical Path Diagram")
    add_docx_body(doc, CRITICAL_PATH_TEXT)
    add_docx_picture_with_caption(
        doc, critical_path, 6.5,
        "Figure 4. Critical path diagram with work package and task references.",
    )
    add_docx_subhead(doc, "Critical Path Interpretation")
    add_docx_body(doc, CRITICAL_PATH_INTERP_TEXT)
    add_docx_subhead(doc, "Planned Deliverables")
    add_docx_body(doc, DELIVERABLES_TEXT)
    add_docx_subhead(doc, "Project Governance Notes")
    add_docx_body(doc, GOVERNANCE_TEXT)
    add_docx_subhead(doc, "Testing and Delivery Criteria")
    add_docx_body(doc, DELIVERY_CRITERIA_TEXT)

    doc.save(str(EDITABLE_DOCX_PATH))


# ---------------------------------------------------------------------------
# Stable DOCX (embedded PDF page images)
# ---------------------------------------------------------------------------
def build_wps_stable_docx(page_paths: list[Path]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    usable_width = section.page_width - section.left_margin - section.right_margin

    for idx, page_path in enumerate(page_paths):
        if idx == 0 and doc.paragraphs:
            p = doc.paragraphs[0]
        else:
            p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str(page_path), width=usable_width)
        if idx < len(page_paths) - 1:
            doc.add_page_break()

    doc.save(str(DOCX_PATH))


def render_pdf_pages() -> None:
    subprocess.run(
        ["pdftoppm", "-png", str(PDF_PATH), str(PDF_RENDER_PREFIX)],
        check=True,
    )


def get_rendered_page_paths() -> list[Path]:
    return sorted(TMP_DIR.glob("pdf_render-*.png"))


def main() -> None:
    ensure_dirs()
    rbs_chart_path = ASSET_DIR / "heartcompass_rbs_chart.png"
    wbs_chart_path = ASSET_DIR / "heartcompass_wbs_chart.png"
    gantt_path = ASSET_DIR / "heartcompass_gantt.png"
    critical_path = ASSET_DIR / "heartcompass_critical_path.png"
    create_rbs_chart(rbs_chart_path)
    create_wbs_chart(wbs_chart_path)
    create_gantt_chart(gantt_path)
    create_critical_path_diagram(critical_path)
    build_pdf(rbs_chart_path, wbs_chart_path, gantt_path, critical_path)
    build_docx(rbs_chart_path, wbs_chart_path, gantt_path, critical_path)
    render_pdf_pages()
    build_wps_stable_docx(get_rendered_page_paths())
    print(f"Generated: {DOCX_PATH}")
    print(f"Generated: {EDITABLE_DOCX_PATH}")
    print(f"Generated: {PDF_PATH}")


if __name__ == "__main__":
    main()
