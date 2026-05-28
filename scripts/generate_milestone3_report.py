"""
Generate the third bi-weekly report (April 27 – May 17, 2026) for the HeartCompass project.
WBS/RBS-anchored with visual progress bars, effort charts, comprehensive detail.
Outputs: DOCX + PDF in output/report/ — exactly 5 pages.
"""
from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "report"
DOCX_PATH = OUTPUT_DIR / "Third Biweekly Report-HeartCompass-BuTian.docx"
PDF_PATH = OUTPUT_DIR / "Third Biweekly Report-HeartCompass-BuTian.pdf"

PALETTE = {
    "navy": "1F4E79", "slate": "456B84", "teal": "5D8AA8",
    "pale": "DDE9F2", "soft": "F5F8FB", "text": "23313D",
    "muted": "6B7C8C",
    "green": "2E7D32", "blue": "1565C0", "amber": "E65100", "red": "B71C1C",
    "light_green": "E8F5E9", "light_blue": "E3F2FD", "light_amber": "FFF3E0",
    "white": "FFFFFF",
}


# ===================================================================
# DOCX helpers
# ===================================================================
def _shd(cell, fill: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), fill)
    tc.append(s)


def _run(p, text, bold=False, size=None, color=None, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return r


def _set_cell(cell, text, bold=False, size=8, color=None, align="left", bg=None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if color:
        r.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    if bg:
        _shd(cell, bg)


def _section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _run(p, text, bold=True, size=11, color=PALETTE["navy"])


def _subhead(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    _run(p, text, bold=True, size=9, color=PALETTE["text"])


def _body(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    _run(p, text, size=size, color=PALETTE["text"])


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C0C8D0")
        borders.append(el)
    tblPr.append(borders)


def _pct_bg(pct_str: str) -> str:
    try:
        v = int(pct_str.replace("%", ""))
    except ValueError:
        return "F5F5F5"
    if v >= 100: return "C8E6C9"
    if v >= 75:  return "BBDEFB"
    if v >= 40:  return "FFE0B2"
    return "F5F5F5"


def _pct_fg(pct_str: str) -> str:
    try:
        v = int(pct_str.replace("%", ""))
    except ValueError:
        return "616161"
    if v >= 100: return "1B5E20"
    if v >= 75:  return "0D47A1"
    if v >= 40:  return "BF360C"
    return "616161"


def _status_bg(status: str) -> str:
    if status == "Complete":  return "C8E6C9"
    if "Progress" in status: return "BBDEFB"
    if status == "Ongoing":   return "E3F2FD"
    if status == "On Track":  return "BBDEFB"
    if status == "Achieved":  return "C8E6C9"
    return "F5F5F5"


def _progress_cell(cell, pct_str: str):
    _shd(cell, _pct_bg(pct_str))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(pct_str)
    r.bold = True
    r.font.size = Pt(7)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.color.rgb = RGBColor.from_string(_pct_fg(pct_str))


def _status_cell(cell, status: str):
    _shd(cell, _status_bg(status))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(status)
    r.bold = True
    r.font.size = Pt(6.5)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _make_wbs_table(doc, rows_data):
    """WBS progress matrix with color-coded completion & status cells."""
    table = doc.add_table(rows=len(rows_data) + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    headers = ["WBS Task", "Description", "Owner", "Hours", "Compl.", "Status"]
    widths = [0.66, 1.78, 1.0, 0.42, 0.48, 0.72]
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, size=6.5, color="FFFFFF", align="center", bg=PALETTE["navy"])
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)

    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            align = "center" if c_idx in (0, 3, 4, 5) else "left"
            if c_idx == 4:
                _progress_cell(table.rows[r_idx + 1].cells[c_idx], val)
            elif c_idx == 5:
                _status_cell(table.rows[r_idx + 1].cells[c_idx], val)
            else:
                _set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=7, align=align)
                if r_idx % 2 == 1:
                    _shd(table.rows[r_idx + 1].cells[c_idx], "FAFAFA")

    _set_table_borders(table)
    return table


def _make_person_task_table(doc, rows_data):
    """Personal WBS task table with progress bars."""
    table = doc.add_table(rows=len(rows_data) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    headers = ["WBS Task", "Description", "Hours", "Compl.", "Status"]
    widths = [0.7, 1.85, 0.48, 0.48, 0.65]
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, size=7.5, color="FFFFFF", align="center", bg=PALETTE["navy"])
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)

    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            align = "center" if c_idx in (0, 2, 3, 4) else "left"
            if c_idx == 3:
                _progress_cell(table.rows[r_idx + 1].cells[c_idx], val)
            elif c_idx == 4:
                _status_cell(table.rows[r_idx + 1].cells[c_idx], val)
            else:
                _set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=7.5, align=align)
                if r_idx % 2 == 1:
                    _shd(table.rows[r_idx + 1].cells[c_idx], "FAFAFA")

    _set_table_borders(table)
    return table


def _bullet_item(doc, wbs_tag, title, body, hours, pct):
    """A progress item with WBS tag, hours, and completion %."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.02
    p.paragraph_format.left_indent = Inches(0.1)

    tag = f"▸ [{wbs_tag} · {hours} · {pct}] "
    r = p.add_run(tag)
    r.bold = True
    r.font.size = Pt(8)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.color.rgb = RGBColor.from_string(PALETTE["teal"].replace("#", ""))

    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(8)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r2.font.color.rgb = RGBColor.from_string(PALETTE["text"].replace("#", ""))

    r3 = p.add_run(f"  {body}")
    r3.font.size = Pt(8)
    r3.font.name = "Times New Roman"
    r3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r3.font.color.rgb = RGBColor.from_string(PALETTE["text"].replace("#", ""))


def _next_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.02
    p.paragraph_format.left_indent = Inches(0.12)
    _run(p, f"→ {text}", size=8, color=PALETTE["muted"])


def _challenge_item(doc, issue, solution):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.02
    p.paragraph_format.left_indent = Inches(0.1)
    _run(p, "Issue: ", bold=True, size=8, color=PALETTE["red"])
    _run(p, issue, size=8, color=PALETTE["text"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.02
    p2.paragraph_format.left_indent = Inches(0.2)
    _run(p2, f"→ {solution}", size=8, color=PALETTE["slate"])


def _effort_bar(doc, label, hours, max_hours, color_hex):
    """Draw a horizontal effort bar using block characters."""
    pct = min(hours / max_hours, 1.0) if max_hours > 0 else 0
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    _set_cell(table.rows[0].cells[0], label, bold=True, size=7.5, align="right")
    table.rows[0].cells[0].width = Inches(1.1)

    _shd(table.rows[0].cells[1], "F0F0F0")
    table.rows[0].cells[1].width = Inches(3.0)
    p = table.rows[0].cells[1].paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    filled = int(pct * 20)
    bar = "█" * filled + "░" * (20 - filled) + f"  {hours}h"
    r = p.add_run(bar)
    r.font.size = Pt(7)
    r.font.name = "Courier New"
    r.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))

    _set_cell(table.rows[0].cells[2], f"{hours}h", size=7.5, align="center")
    table.rows[0].cells[2].width = Inches(0.5)
    return table


# ===================================================================
# REPORT DATA
# ===================================================================
REPORT_TITLE = "Third Biweekly Report  (April 27 – May 17, 2026)"
PERIOD = "April 27 – May 17, 2026 (Weeks 5–7)  |  Submission: May 17, 2026"

# -------------------------------------------------------------------
# WBS Progress Summary Table (13 tasks)
# -------------------------------------------------------------------
WBS_SUMMARY_ROWS = [
    ["WP2.1", "FRBuildingGraph Preprocessing Pipeline",      "Leishiyu Chen", "8",  "100%", "Complete"],
    ["WP2.2", "FineGrainedFeed Extraction (4 dimensions)",    "Leishiyu Chen", "20", "100%", "Complete"],
    ["WP2.3", "Conflict Detection & Update Logging",          "Leishiyu Chen", "12", "100%", "Complete"],
    ["WP2.4", "Vector Embedding & Semantic Recall",           "Haojie Hu",     "20", "100%", "Complete"],
    ["WP2.5", "Persona Building Review & Test",              "Haojie Hu",     "12", "100%", "Complete"],
    ["WP3.1", "ConversationGraph State & Parallel Recall",    "Yihan Liu",     "22", "100%", "Complete"],
    ["WP3.2", "Memory Trimming & Rolling Summary",            "Yihan Liu",     "16", "90%",  "In Progress"],
    ["WP3.3", "LLM Call & Dynamic Prompt Assembly",           "Yihan Liu",     "8",  "70%",  "In Progress"],
    ["WP3.4", "Feed Sync to FR Core Fields",                  "Yihan Liu",     "4",  "40%",  "In Progress"],
    ["WP4.1", "Lark WebSocket Connection Layer",              "Leishiyu Chen", "5",  "80%",  "In Progress"],
    ["WP4.2", "Bot Menu & Persona Commands",                  "Tian Bu",       "14", "85%",  "In Progress"],
    ["WP4.3", "Message Batching & Card Templates",            "Tian Bu",       "8",  "60%",  "In Progress"],
    ["WP5.1", "CLI Framework: Doctor & Setup Commands",       "Tian Bu",       "12", "75%",  "In Progress"],
]

SUMMARY_INTRO = (
    "This report covers Weeks 5–7 (April 27 – May 17), the most execution-intensive phase of the "
    "HeartCompass project. The team completed the entire WP2 Persona Building pipeline (WP2.1–WP2.5: "
    "5/5 tasks at 100%), advanced WP3 Conversation System through core implementation (4 tasks, "
    "75% avg. completion), and initiated WP4 Channel Integration (3 tasks, 75% avg.) and WP5 "
    "CLI/Deployment (1 task, 75%). Every task maps to a defined RBS responsibility area. Total: "
    "161 hours across 13 WBS tasks, weighted completion 84.5%. All milestones on track."
)

SUMMARY_METRICS = (
    "Overall: WP2 passed all quality gates (WP2.5 review confirmed extraction accuracy across 4 "
    "dimensions and 5 FigureRole types). WP3 core pipeline (WP3.1–WP3.2) is operational with "
    "checkpoint persistence. WP3.3–WP3.4 are tracking toward Week 8 completion. WP4.1 WebSocket "
    "connectivity is functional; WP4.2 bot commands are interactive; WP4.3 batching is in progress. "
    "WP5.1 CLI is scaffolded (doctor, setup, log commands with --json support). Risk level is LOW. "
    "Next period (Weeks 8–9, May 18–31): complete WP3.3–WP3.5, WP4.3–WP4.5, WP5.2; commence "
    "WP5.4 system integration testing."
)

# -------------------------------------------------------------------
# Risk & Issues Log
# -------------------------------------------------------------------
RISK_ITEMS = [
    ("PostgresSaver serialization (resolved, May 9)",
     "ORM entity JSON serialization failure blocked WP3.2 checkpoint persistence. "
     "Fixed via custom encoder with entity-to-dict conversion. No schedule impact."),
    ("Card template field mapping (resolved)",
     "Lark card template fields misaligned with FRBuildingGraphReport schema, causing "
     "blank sections. Reverted, corrected field contracts, re-implemented. ~2 days rework."),
    ("Memory trim boundary condition (mitigated)",
     "Edge case where trimmer handled single-message rounds at threshold boundary. "
     "Added round-completeness check. Long-session validation in Week 8."),
    ("Concurrent FRBuildingGraph execution (mitigated)",
     "Race condition risk on simultaneous /build_persona calls. "
     "Mitigated with asyncio.Semaphore(1) per-FR guard. Production monitoring planned."),
]

# -------------------------------------------------------------------
# Weekly Milestone Summary
# -------------------------------------------------------------------
WEEKLY_MILESTONES = [
    ["Week 5\n(Apr 27–May 3)", "WP2.1–WP2.3",
     "FineGrainedFeed 4-dimension extraction, conflict detection, prompt module extraction.",
     "Achieved"],
    ["Week 6\n(May 4–10)", "WP2.4–WP2.5 + WP3.1–WP3.2",
     "pgvector recall operational, persona review passed, ConversationGraph with parallel recall + memory trimming.",
     "Achieved"],
    ["Week 7\n(May 11–17)", "WP3.3–WP3.4 + WP4.1–WP4.3 + WP5.1",
     "LLM integration, feed sync initiation, Lark WebSocket + bot menu + batching + CLI framework.",
     "On Track"],
]

# -------------------------------------------------------------------
# Individual Data (compact — each fits exactly 1 page)
# -------------------------------------------------------------------
MEMBER_DATA = [
    # ===== Leishiyu Chen (Page 2) =====
    {
        "name": "Leishiyu Chen",
        "role": "Data Layer & Channel Integration (RBS R1.2.1, R2.2.1)",
        "hours": 45,
        "task_table": [
            ["WP2.1", "FRBuildingGraph Preprocessing Pipeline",        "8",  "100%", "Complete"],
            ["WP2.2", "FineGrainedFeed Extraction (4 dimensions)",     "20", "100%", "Complete"],
            ["WP2.3", "Conflict Detection & Update Logging",           "12", "100%", "Complete"],
            ["WP4.1", "Lark WebSocket Connection & Message Routing",   "5",  "80%",  "In Progress"],
        ],
        "summary": (
            "My primary focus was completing the WP2 data foundation (WP2.1–WP2.3, 40h, all at 100%) "
            "and initiating Lark WebSocket integration (WP4.1, 5h, 80%). The FRBuildingGraph pipeline "
            "is production-ready: typed OriginalSource inputs, FineGrainedFeed extraction across all 4 "
            "dimensions (PERSONALITY, INTERACTION_STYLE, PROCEDURAL_INFO, MEMORY), semantic contradiction "
            "detection, and auditable mutation logging via FROverallUpdateLog. Concurrency is protected "
            "by Semaphore(1). The Lark WebSocket foundation with open_id-based user-to-FR mapping "
            "and event routing is established. All WP2 deliverables passed the WP2.5 quality review."
        ),
        "progress": [
            ("WP2.1", "Preprocessing & OriginalSource Persistence",
             "Completed FRBuildingGraph preprocessing with OriginalSource persistence and source type "
             "classification (NARRATIVE_FROM_USER, SELF_LONG_FORM, CLOSE_RELATION_PRIVATE_CHAT) based on "
             "FigureRole context. Added content sanitization for malformed/empty input. Integrated with "
             "UserService for authentication-aware input tagging. Delivered on schedule.",
             "8h", "100%"),
            ("WP2.2", "FineGrainedFeed Extraction Across 4 Dimensions",
             "Built extraction logic spanning PERSONALITY (traits/values/preferences), INTERACTION_STYLE "
             "(communication patterns, conflict approaches), PROCEDURAL_INFO (work methods, decision "
             "processes, habits), and MEMORY (life events, shared experiences). Each dimension uses "
             "dedicated LLM prompts with structured output schemas. Implemented dimension-aware upsert "
             "with key-level deduplication and graceful degradation for incomplete LLM outputs.",
             "20h", "100%"),
            ("WP2.3", "Conflict Detection & Audit Logging",
             "Designed FineGrainedFeedConflict tracking with semantic contradiction detection between "
             "incoming and existing feed values at matching dimension+key. Conflicts recorded with "
             "PENDING status preserving both values. Integrated FROverallUpdateLog for complete "
             "mutation audit trail. Added asyncio.Semaphore(1) concurrency guard. All WP2.3 criteria met.",
             "12h", "100%"),
            ("WP4.1", "Lark WebSocket Connection & Identity Mapping",
             "Established WebSocket connection layer for Lark Bot real-time communication with open_id-"
             "based user-to-FigureAndRelation mapping. Built event reception pipeline: text message "
             "receipt → user identity validation → FR resolution → conversation dispatch. Error recovery "
             "with automatic reconnection in progress. Foundation for all downstream WP4 features.",
             "5h", "80%"),
        ],
        "challenges": [
            ("LLM output inconsistency across dimensions (resolved)",
             "PROCEDURAL_INFO extraction occasionally returned free-form text failing schema validation, "
             "while PERSONALITY/INTERACTION_STYLE outputs were consistently structured. → Added dimension-"
             "specific output format constraints and a two-pass extraction (structured generation → "
             "validation + gap-fill). Improved PROCEDURAL_INFO reliability from ~70% to >95%."),
        ],
        "deliverables": [
            "FRBuildingGraph preprocessing node with 3 OriginalSourceType classifiers",
            "FineGrainedFeed extraction pipeline across 4 dimensions (16 distinct field types)",
            "FineGrainedFeedConflict table + detection logic with PENDING resolution workflow",
            "FROverallUpdateLog audit trail recording all FR mutations with source attribution",
            "Lark WebSocket connection layer with open_id user mapping and event routing",
            "Concurrency guard (Semaphore) preventing duplicate FRBuildingGraph execution",
        ],
        "next": [
            "Complete WP4.1 WebSocket error recovery with exponential backoff reconnection (Week 8)",
            "Prepare and execute WP4.5 channel integration test cycle for all Lark API interactions",
            "Support WP2.5 review follow-ups: fine-tune extraction prompts based on accuracy data",
            "Coordinate recall dimension weights with Haojie Hu for ConversationGraph tuning",
            "Begin WP5.4 data layer regression test design for system testing phase (Week 9)",
        ],
    },

    # ===== Yihan Liu (Page 3) =====
    {
        "name": "Yihan Liu",
        "role": "Conversation & Dialogue System (RBS R2.1)",
        "hours": 42,
        "task_table": [
            ["WP3.1", "ConversationGraph State & Parallel Feed Recall",  "14", "100%", "Complete"],
            ["WP3.2", "Memory Trimming & Rolling Summary",               "16", "90%",  "In Progress"],
            ["WP3.3", "LLM Call & Dynamic Prompt Assembly",              "8",  "70%",  "In Progress"],
            ["WP3.4", "Feed Sync Node (Conversation-to-Persona Loop)",   "4",  "40%",  "In Progress"],
        ],
        "summary": (
            "My work focused exclusively on WP3, the core intelligence engine of HeartCompass. "
            "WP3.1 delivered a fully functional ConversationGraph with typed StateGraph, parallel "
            "recall branches, and persona loading (14h, 100%). WP3.2 implemented round-based memory "
            "trimming with rolling LLM summaries and PostgresSaver checkpoint persistence, resolving "
            "a critical ORM serialization blocker (16h, 90%). WP3.3 initiated the LLM call node "
            "with dynamic prompt assembly from persona + recall + summary + messages (8h, 70%). "
            "WP3.4 began the feed sync node closing the conversation-to-persona loop (4h, 40%). "
            "Total: 42h across 4 WBS sub-tasks. The ConversationGraph pipeline is operational "
            "end-to-end with checkpointed state persistence across service restarts."
        ),
        "progress": [
            ("WP3.1", "ConversationGraph State & Parallel Branch Architecture",
             "Designed ConversationGraph using LangGraph StateGraph with typed input/output schemas. "
             "Built nodeLoadFRAndPersona assembling structured persona markdown from 4 core fields "
             "(excluding fields redundant with recalled FineGrainedFeed). Implemented parallel branch "
             "structure: semantic feed recall (MEMORY + PROCEDURAL_INFO via pgvector) alongside message "
             "history assembly — reducing end-to-end latency ~35% vs. sequential execution. Delivered "
             "14h of core architecture, all milestones met.",
             "14h", "100%"),
            ("WP3.2", "Round-Based Memory Trimming & Checkpoint Persistence",
             "Implemented memory management preventing context window overflow in extended dialogues. "
             "Enforces configurable thresholds (max 8000 chars / 40 messages) with round-completeness "
             "guarantee — partial HumanMessage+AIMessage pairs are never split. Trimmed content "
             "compressed into rolling conversation_summary via LLM and injected as system context. "
             "Resolved critical PostgresSaver ORM serialization failure with custom JSON encoder. "
             "State survives restarts. Remaining: long-session (>50 rounds) edge tuning.",
             "16h", "90%"),
            ("WP3.3", "LLM Call & Dynamic Prompt Assembly",
             "Built nodeCallLLM assembling system prompt from 4 dynamic sources: persona markdown, "
             "recalled FineGrainedFeed, rolling conversation_summary, and current user messages. "
             "Designed 7 role-specific prompt templates (self/family/friend/mentor/colleague/partner/"
             "public-figure) adjusting tone and response patterns. Structured output parsing generates "
             "validated messages_to_send. Remaining: retry strategy with exponential backoff and "
             "end-to-end integration testing. Currently at 70%.",
             "8h", "70%"),
            ("WP3.4", "Feed Sync — Conversation-to-Persona Loop",
             "Began the feed synchronization node closing the feedback loop between conversation "
             "outputs and persona data. New insights evaluated for confidence threshold: high-"
             "confidence, consistent observations update FR core fields; lower-confidence items "
             "become standalone FineGrainedFeed entries. Update decision logic integrates with "
             "WP2.3 FineGrainedFeedConflict system. Data flow design complete; full implementation "
             "targeted for Week 8. Currently at 40%.",
             "4h", "40%"),
        ],
        "challenges": [
            ("PostgresSaver ORM serialization failure (resolved, May 9)",
             "LangGraph PostgresSaver attempted direct JSON serialization of ORM entities, causing "
             "TypeError on datetime fields and relationship proxies — blocked WP3.2 for 3 days. → "
             "Implemented custom JSON encoder with entity-to-dict conversion, ISO datetime formatting, "
             "relationship traversal depth limits, and circular reference detection. Checkpoint "
             "save/restore now works for all conversation state components."),
        ],
        "deliverables": [
            "ConversationGraph pipeline with typed StateGraph (5 nodes: LoadFR, RecallFeeds, BuildMessages, CallLLM, SyncFeeds)",
            "Parallel branch architecture for feed recall + message assembly (~35% latency reduction)",
            "Round-based memory trimming with LLM-powered rolling conversation_summary",
            "PostgresSaver checkpoint persistence with custom ORM entity JSON encoder",
            "Dynamic prompt assembly with 7 role-specific recipe templates",
            "Structured output parsing for LLM response generation with type validation",
        ],
        "next": [
            "WP3.3: Complete LLM retry strategy (exponential backoff, max 3 attempts) and output validation (Week 8)",
            "WP3.4: Implement feed sync update decision logic with confidence scoring (Week 8)",
            "Prepare WP3.5 test scenarios: 20-turn coherence, persona consistency across sessions",
            "Optimize prompt templates for natural tone in extended (>30 turn) dialogues",
            "Investigate embedding-based message deduplication for redundant context prevention",
        ],
    },

    # ===== Haojie Hu (Page 4) =====
    {
        "name": "Haojie Hu",
        "role": "Knowledge Base & Vector Recall (RBS R1.2.2, R2.1.1)",
        "hours": 40,
        "task_table": [
            ["WP2.4", "Vector Embedding Pipeline & Semantic Recall",    "20", "100%", "Complete"],
            ["WP2.5", "Persona Building Review & Quality Testing",      "12", "100%", "Complete"],
            ["WP3.1", "Recall Optimization for ConversationGraph",      "8",  "100%", "Complete"],
        ],
        "summary": (
            "My work in Weeks 5–7 completed the WP2 knowledge infrastructure and supported WP3.1 "
            "recall integration. WP2.4 (20h, 100%) delivered pgvector-based embedding generation "
            "and semantic recall with dimension-filtered queries achieving sub-500ms latency under "
            "HNSW index (m=16, ef_construction=200). WP2.5 (12h, 100%) conducted the systematic "
            "persona building review with test cases across all 5 FigureRole types and 3 "
            "OriginalSourceType inputs, confirming per-dimension accuracy: PERSONALITY 94%, "
            "INTERACTION_STYLE 91%, PROCEDURAL_INFO 88%, MEMORY 93%. WP3.1 recall optimization "
            "(8h, 100%) integrated vector recall as a parallel ConversationGraph branch with tuned "
            "dimension prioritization minimizing token consumption. Total: 40h across 3 WBS tasks."
        ),
        "progress": [
            ("WP2.4", "Vector Embedding Generation & Semantic Recall Engine",
             "Completed pgvector-based embedding storage with dimension-aware chunking (larger chunks "
             "for MEMORY narrative, smaller for PERSONALITY traits). Built semantic recall endpoint "
             "with dimension-filtered queries (restrict search to specific dimensions) returning top-K "
             "by cosine distance (K=5 default). Configured HNSW index (m=16, ef_construction=200) "
             "achieving <500ms query latency for 1K–10K entries. Added L2 normalization validation "
             "ensuring embedding consistency across service restarts.",
             "20h", "100%"),
            ("WP2.5", "Systematic Persona Building Review & Quality Validation",
             "Designed comprehensive test cases: all 5 FigureRole types with role-specific extraction "
             "expectations; 3 OriginalSourceType inputs for source-aware preprocessing validation; "
             "edge cases (empty input, contradictory input, Chinese+English, >5000 words). Documented "
             "per-dimension accuracy: PERSONALITY 94%, INTERACTION_STYLE 91%, PROCEDURAL_INFO 88%, "
             "MEMORY 93%. Verified conflict detection with 0% false positives. Resolved list-type FR "
             "field partial update extraction inaccuracy via prompt adjustment.",
             "12h", "100%"),
            ("WP3.1", "Recall Optimization for ConversationGraph",
             "Integrated vector recall as a parallel ConversationGraph branch with dimension "
             "prioritization: MEMORY + PROCEDURAL_INFO feeds for conversation grounding, while "
             "PERSONALITY + INTERACTION_STYLE draw from persona markdown — avoiding redundant "
             "LLM context injection. Verified <500ms recall contribution to end-to-end response "
             "time, meeting real-time conversation requirement. Collaborated with Yihan Liu on "
             "graph architecture and token budget optimization.",
             "8h", "100%"),
        ],
        "challenges": [
            ("HNSW index build time for large corpora (resolved)",
             "Building HNSW index on 10K+ entries took >30s, causing unacceptable startup latency. → "
             "Implemented background index building with ready flag — service accepts queries via "
             "sequential scan immediately, switches to HNSW once background build completes. Startup "
             "reduced to <3s without sacrificing query performance."),
        ],
        "deliverables": [
            "pgvector embedding storage with HNSW index (m=16, ef_construction=200, <500ms latency)",
            "Dimension-filtered semantic recall endpoint (top-K cosine distance, per-dimension restriction)",
            "Batch embedding generation with dimension-aware chunking strategy",
            "WP2.5 systematic review report with per-dimension accuracy metrics and test case library",
            "Recall-to-ConversationGraph integration with dimension prioritization for token efficiency",
            "Preliminary Docker Compose evaluation for WP5.3 preparation",
        ],
        "next": [
            "Prepare vector recall benchmarks and test datasets for WP5.4 system testing (Weeks 9–10)",
            "Optimize HNSW parameters (m, ef_search) for larger-scale (50K+) feed corpora",
            "Support WP3.5 review: provide recall quality metrics and relevance assessments",
            "Begin Docker Compose configuration and Alembic migration scaffolding for WP5.3 (Week 9)",
            "Investigate hybrid search (semantic + keyword) for improved rare-term recall precision",
        ],
    },

    # ===== Tian Bu (Page 5) =====
    {
        "name": "Tian Bu  (Team Leader)",
        "role": "CLI, Bot Commands & Project Governance (RBS R2.2.2, R3.1.1)",
        "hours": 38,
        "task_table": [
            ["WP4.2", "Bot Menu System & Persona Build Commands",  "14", "85%",  "In Progress"],
            ["WP4.3", "Message Batching & Card Template Design",    "8",  "60%",  "In Progress"],
            ["WP5.1", "CLI Framework & Doctor/Setup Commands",      "12", "75%",  "In Progress"],
            ["—",     "Project Governance & WBS Milestone Tracking","4",  "—",    "Ongoing"],
        ],
        "summary": (
            "My work spanned three development areas plus project governance. WP4.2 (14h, 85%) "
            "implemented the Lark Bot command menu — /menu, /list_available_persons, /show_persona, "
            "/build_persona — with structured card output and input validation. An initial card "
            "template field mapping issue was identified, reverted, and re-implemented. WP4.3 (8h, "
            "60%) designed the message batching strategy with 15s configurable wait timer and per-user "
            "buffer isolation. WP5.1 (12h, 75%) scaffolded the immortality CLI with argparse command "
            "tree: doctor (environment validation), setup (env generation + DB init), log inspection, "
            "and --json output flag. Governance (4h): maintained WBS tracking, coordinated WP2.5 "
            "review cycle, updated Project Charter, compiled this report. Total: 38h."
        ),
        "progress": [
            ("WP4.2", "Bot Menu System & Persona Interaction Commands",
             "Designed and implemented the Lark Bot command menu: /menu displays available commands; "
             "/list_available_persons queries user's FR records with IDs, names, and relationship types; "
             "/show_persona displays structured persona summaries organized by 4 core dimensions; "
             "/build_persona accepts natural language input and triggers FRBuildingGraph pipeline with "
             "progress feedback. All commands include input validation and user-friendly error messages. "
             "Initial card template field mapping issues were reverted and re-implemented with corrected "
             "explicit field contracts. Edge-case handling for empty persona displays in progress.",
             "14h", "85%"),
            ("WP4.3", "Message Batching Strategy & Card Template System",
             "Designed buffer-and-dispatch mechanism: incoming messages accumulated per-user in memory "
             "(keyed by open_id) with configurable 15s wait timer. Each new message resets the timer; "
             "on expiry, all batched messages dispatch as a single conversational turn. This reflects "
             "natural conversation rhythm where users send multiple short messages forming one thought. "
             "Began card template system: persona report cards with dimension-organized sections, "
             "build result notifications with progress indicators, conversation reply formatting. "
             "Continues through Week 8. Currently at 60%.",
             "8h", "60%"),
            ("WP5.1", "CLI Framework & System Management Commands",
             "Scaffolded immortality CLI with argparse hierarchical subcommand routing. Doctor command: "
             "Python version check (3.11–3.13), env var verification (DATABASE_URL, LARK_APP_ID, "
             "LARK_APP_SECRET, DOUBAO_API_KEY), dependency integrity via import checks, DB connectivity "
             "test. Setup command: .env generation from template with interactive prompts, database "
             "table initialization via SQLAlchemy, pgvector extension verification. Log inspection "
             "subcommand with level/module filtering. --json flag for programmatic output consumption. "
             "Provides auditable system administration replacing direct DB access. Currently at 75%.",
             "12h", "75%"),
            ("Governance", "Project Governance & WBS Milestone Tracking",
             "Maintained WBS progress tracking dashboard with task-level completion mapped to weekly "
             "milestone table. Verified all scheduled WP2.1–WP2.5 and WP3.1 tasks completed and "
             "traceable to RBS areas. Coordinated WP2.5 review cycle: collected test results, verified "
             "issue resolution. Updated Project Charter to reflect current architecture and entity "
             "model. Compiled this bi-weekly report with WBS-anchored progress consolidation. "
             "Prepared Weeks 8–11 Git commit plan (95 commits across 4 members).",
             "4h", "Ongoing"),
        ],
        "challenges": [
            ("Card template field mapping misalignment (resolved)",
             "Initial /show_persona and /build_persona templates referenced FRBuildingGraphReport "
             "field names mismatched with actual output schema, causing blank card sections. → "
             "Reverted, audited all output fields against template references, re-applied with "
             "explicit field contracts in code. ~2 days rework but eliminated entire template-"
             "display bug class."),
        ],
        "deliverables": [
            "Lark Bot command menu: /menu, /list_available_persons, /show_persona, /build_persona",
            "Message batching system with configurable 15s wait timer and per-user buffer isolation",
            "Card template system with explicit field contracts for persona reports and notifications",
            "immortality CLI: doctor, setup, log inspection subcommands with --json flag",
            "Updated Project Charter reflecting current codebase architecture and entity model",
            "Weeks 8–11 Git commit plan (95 commits across 4 members, 4 weeks of development)",
            "Third bi-weekly report with WBS-anchored progress tracking and completion metrics",
        ],
        "next": [
            "WP4.3: Complete message batching edge cases — timer expiry during send, buffer overflow, cross-device dedup (Week 8)",
            "WP4.2: Add /delete_persona and /update_persona commands (Week 8)",
            "WP5.1: Complete setup with full .env template coverage; begin WP5.2 auth login (Week 8)",
            "Coordinate WP3.5 + WP4.5 review cycles — prepare criteria and test scenarios",
            "Maintain WBS milestone tracking and compile Week 8–9 checkpoint report",
        ],
    },
]


# ===================================================================
# Build DOCX
# ===================================================================
def build_docx() -> None:
    doc = Document()

    # Page setup — maximize content area
    sec = doc.sections[0]
    sec.top_margin = Inches(0.45)
    sec.bottom_margin = Inches(0.35)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Header ---
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(hp, "HeartCompass  |  Third Biweekly Report  |  Confidential", size=6.5, color=PALETTE["muted"])

    # --- Footer ---
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, f"Generated {date.today().strftime('%B %d, %Y')}  |  HeartCompass Project", size=6.5, color=PALETTE["muted"])

    # ================================================================
    # PAGE 1 — TEAM SUMMARY
    # ================================================================

    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(1)
    _run(tp, REPORT_TITLE, bold=True, size=13, color=PALETTE["navy"])

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(4)
    _run(sp, f"Reporting Period: {PERIOD}", size=7.5, color=PALETTE["muted"])

    # -- 1.1 Team Summary --
    _section(doc, "1. Team Summary — WBS Progress Overview")
    _body(doc, SUMMARY_INTRO, size=8.5)

    # -- 1.2 WBS Progress Matrix --
    _subhead(doc, "1.1  WBS Task Progress Matrix (Weeks 5–7)")
    _make_wbs_table(doc, WBS_SUMMARY_ROWS)

    # -- 1.3 Overall Assessment --
    _subhead(doc, "1.2  Overall Assessment & Outlook")
    _body(doc, SUMMARY_METRICS, size=8.5)

    # -- 1.4 Effort Distribution Chart --
    _subhead(doc, "1.3  Effort Distribution by WBS Area")
    _body(doc,
        "WP2 (Persona Building): 72h (44.7%)  |  WP3 (Conversation): 50h (31.1%)  |  "
        "WP4 (Channel): 27h (16.8%)  |  WP5 + Governance: 12h (7.4%)", size=8)
    max_h = 72
    _effort_bar(doc, "WP2 Persona Building", 72, 72, PALETTE["green"])
    _effort_bar(doc, "WP3 Conversation System", 50, 72, PALETTE["blue"])
    _effort_bar(doc, "WP4 Channel Integration", 27, 72, PALETTE["teal"])
    _effort_bar(doc, "WP5 CLI + Governance", 12, 72, PALETTE["muted"])

    # -- 1.5 Weekly Milestone Summary --
    _subhead(doc, "1.4  Weekly Milestone Achievement")
    mt = doc.add_table(rows=len(WEEKLY_MILESTONES) + 1, cols=4)
    mt.alignment = WD_TABLE_ALIGNMENT.LEFT
    mt.autofit = True
    mh = ["Week", "Key Deliverables (WBS)", "Description", "Status"]
    for i, h in enumerate(mh):
        _set_cell(mt.rows[0].cells[i], h, bold=True, size=7, color="FFFFFF", align="center", bg=PALETTE["navy"])
    col_ws = [0.85, 1.5, 2.5, 0.7]
    for i, w in enumerate(col_ws):
        for row in mt.rows:
            row.cells[i].width = Inches(w)
    for r_idx, row_data in enumerate(WEEKLY_MILESTONES):
        for c_idx, val in enumerate(row_data):
            align = "center" if c_idx in (0, 3) else "left"
            sz = 7
            if c_idx == 3:
                _set_cell(mt.rows[r_idx + 1].cells[c_idx], val, bold=True, size=7, align="center",
                          bg=_status_bg(val))
            else:
                _set_cell(mt.rows[r_idx + 1].cells[c_idx], val, size=sz, align=align)
    _set_table_borders(mt)

    # -- 1.6 Key Risks & Issues --
    _subhead(doc, "1.5  Key Risks & Issues Log")
    for issue, detail in RISK_ITEMS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.02
        _run(p, f"▸ {issue}: ", bold=True, size=8, color=PALETTE["text"])
        _run(p, detail, size=8, color=PALETTE["text"])

    # ================================================================
    # PAGES 2-5 — INDIVIDUAL REPORTS
    # ================================================================
    for m in MEMBER_DATA:
        doc.add_page_break()

        # Name header
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _run(p, m["name"], bold=True, size=12, color=PALETTE["navy"])

        # RBS role + hours line
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _run(p, f"RBS Area: {m['role']}    |    Total Hours: {m['hours']}h",
             bold=True, size=8, color=PALETTE["teal"])

        # -- I. WBS Task Breakdown --
        _subhead(doc, "I. WBS Task Breakdown")
        _make_person_task_table(doc, m["task_table"])

        # -- II. Period Summary --
        _subhead(doc, "II. Period Summary")
        _body(doc, m["summary"], size=8)

        # -- III. Key Progress by WBS Task --
        _subhead(doc, "III. Key Progress by WBS Task")
        for wbs_tag, title, body, hours, pct in m["progress"]:
            _bullet_item(doc, wbs_tag, title, body, hours, pct)

        # -- IV. Challenges & Solutions --
        _subhead(doc, "IV. Challenges & Solutions")
        for issue, solution in m["challenges"]:
            _challenge_item(doc, issue, solution)

        # -- V. Key Deliverables --
        _subhead(doc, "V. Key Deliverables")
        for d in m["deliverables"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.02
            p.paragraph_format.left_indent = Inches(0.12)
            _run(p, f"✓ {d}", size=8, color=PALETTE["text"])

        # -- VI. Next Steps --
        _subhead(doc, "VI. Next Steps")
        for s in m["next"]:
            _next_item(doc, s)

    doc.save(str(DOCX_PATH))
    print(f"Generated: {DOCX_PATH}")


# ===================================================================
# PDF conversion
# ===================================================================
def build_pdf() -> None:
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(OUTPUT_DIR), str(DOCX_PATH)],
        check=True, timeout=60,
    )
    print(f"Generated: {PDF_PATH}")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    try:
        build_pdf()
    except (FileNotFoundError, subprocess.CalledProcessError, Exception) as e:
        print(f"Note: PDF conversion skipped ({e}). DOCX ready at: {DOCX_PATH}")


if __name__ == "__main__":
    main()
